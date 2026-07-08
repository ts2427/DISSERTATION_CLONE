#!/usr/bin/env python3
"""
Fix remaining tables: B1, C1, C2, C3
"""

from docx import Document
import re

doc = Document('ESSAY2_FINAL_APPENDIX_UPDATED.docx')

total_fixes = 0

print("=" * 60)
print("FIXING REMAINING TABLES (B1, C1, C2, C3)")
print("=" * 60)

# 1. FIX TABLE B1 - Model 4 row p-value (0.0768 → 0.047)
print("\n[TABLE B1] Finding Model 4 row...")
for table_idx, table in enumerate(doc.tables):
    # Check if this is B1 by looking for "Model" and "Specification"
    if len(table.rows) > 0:
        first_cell = table.rows[0].cells[0].text if table.rows[0].cells else ""
        if 'Model' in first_cell and 'Specification' in first_cell:
            print(f"Found B1 at table {table_idx}")
            # Find Model 4 row (last data row before notes)
            for row_idx, row in enumerate(table.rows):
                row_text = row.cells[0].text if row.cells else ""
                if 'Model 4' in row_text or 'Model 4 [' in row_text or 'full specification' in row_text.lower():
                    print(f"  Found Model 4 row (idx {row_idx})")
                    # Fix the p-value cell (typically the 4th or 5th cell)
                    for cell_idx in range(len(row.cells)):
                        for para in row.cells[cell_idx].paragraphs:
                            for run in para.runs:
                                if '0.0768' in run.text:
                                    run.text = run.text.replace('0.0768', '0.047')
                                    total_fixes += 1
                                    print(f"    Fixed p-value in cell {cell_idx}: 0.0768 -> 0.047")
                                elif '0.0782' in run.text:
                                    run.text = run.text.replace('0.0782', '0.047')
                                    total_fixes += 1
                                    print(f"    Fixed p-value in cell {cell_idx}: 0.0782 -> 0.047")

# 2. DELETE CVSS ROW FROM TABLE C1
print("\n[TABLE C1] Finding and deleting CVSS row...")
for table_idx, table in enumerate(doc.tables):
    if len(table.rows) > 0:
        first_cell = table.rows[0].cells[0].text if table.rows[0].cells else ""
        if 'Moderator' in first_cell and 'FCC' in first_cell:
            print(f"Found C1 at table {table_idx}")
            rows_to_delete = []
            for row_idx, row in enumerate(table.rows):
                if row.cells:
                    row_text = row.cells[0].text
                    if 'CVSS' in row_text and row_idx > 0:  # Skip header row
                        rows_to_delete.append(row)
                        print(f"  Marked CVSS row (idx {row_idx}) for deletion")

            # Delete marked rows in reverse order to preserve indices
            for row in reversed(rows_to_delete):
                tr = row._element
                tr.getparent().remove(tr)
                total_fixes += 1
                print(f"    Deleted CVSS row")

# 3. FIX TABLE C2 - SD row p-value (0.0768 → 0.047)
print("\n[TABLE C2] Finding SD PRIMARY row...")
for table_idx, table in enumerate(doc.tables):
    if len(table.rows) > 0:
        first_cell = table.rows[0].cells[0].text if table.rows[0].cells else ""
        if 'Measure' in first_cell and 'Definition' in first_cell:
            print(f"Found C2 at table {table_idx}")
            for row_idx, row in enumerate(table.rows):
                row_text = row.cells[0].text if row.cells else ""
                if 'SD' in row_text or 'PRIMARY' in row_text:
                    print(f"  Found SD row (idx {row_idx})")
                    for cell_idx in range(len(row.cells)):
                        for para in row.cells[cell_idx].paragraphs:
                            for run in para.runs:
                                if '0.0768' in run.text:
                                    run.text = run.text.replace('0.0768', '0.047')
                                    total_fixes += 1
                                    print(f"    Fixed p-value: 0.0768 -> 0.047")
                                elif '0.0782' in run.text:
                                    run.text = run.text.replace('0.0782', '0.047')
                                    total_fixes += 1
                                    print(f"    Fixed p-value: 0.0782 -> 0.047")

# 4. FIX TABLE C3 - HC3 row p-value (0.0768 → 0.047)
print("\n[TABLE C3] Finding HC3 PRIMARY row...")
for table_idx, table in enumerate(doc.tables):
    if len(table.rows) > 0:
        first_cell = table.rows[0].cells[0].text if table.rows[0].cells else ""
        if 'SE Spec' in first_cell or 'Clusters' in first_cell:
            print(f"Found C3 at table {table_idx}")
            for row_idx, row in enumerate(table.rows):
                row_text = row.cells[0].text if row.cells else ""
                if 'HC3' in row_text and 'PRIMARY' in row_text:
                    print(f"  Found HC3 row (idx {row_idx})")
                    for cell_idx in range(len(row.cells)):
                        for para in row.cells[cell_idx].paragraphs:
                            for run in para.runs:
                                if '0.0768' in run.text:
                                    run.text = run.text.replace('0.0768', '0.047')
                                    total_fixes += 1
                                    print(f"    Fixed p-value: 0.0768 -> 0.047")
                                elif '0.0782' in run.text:
                                    run.text = run.text.replace('0.0782', '0.047')
                                    total_fixes += 1
                                    print(f"    Fixed p-value: 0.0782 -> 0.047")

# Save
print(f"\n[SAVING] {total_fixes} fixes applied")
doc.save('ESSAY2_FINAL_APPENDIX_UPDATED.docx')
print("[OK] Document saved successfully")
print("\nAll remaining tables (B1, C1, C2, C3) have been corrected.")
