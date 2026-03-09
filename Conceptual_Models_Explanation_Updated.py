#!/usr/bin/env python3
"""
Create explanation guide for conceptual models - UPDATED WITH MECHANISM ANALYSIS
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Conceptual Models: Explanation Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Brief Talking Points for Proposal & Defense', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

date_para = doc.add_paragraph('March 2026 - UPDATED WITH MECHANISM ANALYSIS')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.runs[0].font.italic = True

doc.add_paragraph()

# Overview
doc.add_heading('Overview', level=1)
doc.add_paragraph(
    'These four conceptual models illustrate how mandatory data breach disclosure '
    'regulation (specifically the FCC Rule 37.3, implemented in 2007) affects firms '
    'through three distinct and independent mechanisms. The models are designed to be '
    'included in your dissertation proposal and defense presentation. This updated guide '
    'incorporates heterogeneous mechanism analysis identifying firm size as the key moderator.'
)

doc.add_paragraph()

# MODEL 1
doc.add_heading('MODEL 1: Essay 1 - Market Reactions (Stock Price Impact)', level=1)

doc.add_heading('Visual Layout', level=2)
doc.add_paragraph(
    'A left-to-right flow showing: Data Breach Disclosure → Timing Decision → Information Quality → Market Valuation (CAR -0.74%)'
)

doc.add_heading('What This Model Shows', level=2)
doc.add_paragraph(
    'This model illustrates Essay 1\'s research question: "Does disclosure timing affect how markets value firms '
    'after breaches?" It depicts a causal chain where the timing of breach disclosure leads to different information '
    'quality, which theoretically should affect market valuation.'
)

doc.add_heading('Your Talking Points (30-45 seconds)', level=2)
talking_points = [
    'We hypothesized that faster disclosure would lead to more favorable market reactions because it demonstrates '
    'management transparency and reduces uncertainty.',
    'However, we find a surprising null result: disclosure timing does NOT affect stock price reactions (coefficient +0.57%, p=0.539).',
    'Instead, what DOES matter: FCC regulatory status (-2.20%, highly significant), health breach complexity (-2.51%), '
    'and prior breach history (-0.22% per breach—the strongest effect).',
    'Key insight: Markets appear to price breach severity and firm characteristics, NOT disclosure speed. This challenges '
    'the regulatory assumption that "faster is better."'
]

for i, point in enumerate(talking_points, 1):
    doc.add_paragraph(point, style='List Number')

doc.add_paragraph()

# MODEL 2
doc.add_heading('MODEL 2: Essay 2 - Information Asymmetry (Cost of Capital)', level=1)

doc.add_heading('Visual Layout', level=2)
doc.add_paragraph(
    'Left-to-right flow: Data Breach → Timing Requirement (FCC: 7 days) → Investigation Incompleteness → Market Uncertainty (Volatility). '
    'Add branching beneath for firm size heterogeneity.'
)

doc.add_heading('What This Model Shows', level=2)
doc.add_paragraph(
    'This model illustrates Essay 2\'s paradoxical finding: mandatory timing requirements can INCREASE market uncertainty. '
    'The logic: forced early disclosure prevents thorough investigation, leaving questions unanswered, which increases volatility. '
    'Additionally, the mechanism analysis reveals this effect is NOT uniform—it depends critically on firm size and compliance capacity.'
)

doc.add_heading('Your Talking Points (30-45 seconds)', level=2)
talking_points = [
    'While Essay 1 found markets don\'t reward disclosure speed, Essay 2 discovers that mandatory timing creates a different cost: '
    'increased information asymmetry and cost of capital.',
    'When regulators mandate 7-day disclosure (FCC) or 4-day disclosure (SEC 2023 rules), companies face a time constraint that may '
    'force incomplete disclosure. This leaves investors uncertain.',
    'Our finding: FCC regulation increases volatility by 1.83%, equivalent to a 0.0137% increase in cost of capital. For a $50B firm, '
    'this translates to $10M+ in additional annual financing costs.',
    'The paradox: disclosure requirements achieve their goal of faster information BUT at the cost of higher information asymmetry '
    'and higher cost of capital for affected firms.'
]

for i, point in enumerate(talking_points, 1):
    doc.add_paragraph(point, style='List Number')

doc.add_paragraph()

# NEW: MECHANISM ANALYSIS SECTION
doc.add_heading('MODEL 2A: Essay 2 Mechanism Analysis - Firm Size Heterogeneity', level=1)

doc.add_heading('Visual Layout', level=2)
doc.add_paragraph(
    'Visual showing FCC Rule 37.3 branching into two paths: Small Firm Path (with resource constraints → incomplete disclosure → high volatility +7.31%) '
    'and Large Firm Path (with resources → complete disclosure → low/reversed volatility -3.39%).'
)

doc.add_heading('What This Model Shows', level=2)
doc.add_paragraph(
    'The heterogeneous mechanism analysis reveals that the FCC volatility effect depends critically on firm size. Small firms face '
    'tighter information processing capacity constraints and thus experience larger volatility increases. Large firms with dedicated breach response '
    'teams can investigate thoroughly within the deadline and experience no penalty (or even reversed effects). This is explained by Tushman & Nadler (1978) '
    'information processing theory: organizations have capacity limits that constrain how fast they can produce complete information.'
)

doc.add_heading('Your Talking Points (45-60 seconds)', level=2)
talking_points = [
    'The mechanism analysis reveals a critical insight: the FCC volatility penalty is NOT uniform across firms. It depends on firm size and compliance capacity.',
    'For the smallest firms (Q1): FCC regulation increases volatility by 7.31% (highly significant, p=0.003). For the largest firms (Q4): it actually DECREASES volatility by 3.39% (p=0.015). That\'s an 11 percentage point difference.',
    'Why? The key mechanism is information processing capacity. Small firms with limited compliance staff cannot compress investigation timelines. The FCC 7-day deadline forces them to choose: miss the deadline or disclose incompletely.',
    'Large firms with dedicated breach response teams can investigate thoroughly within 7 days. They face no penalty.',
    'Theory: Tushman & Nadler (1978) predicted that when information demands exceed organizational capacity, forced speed produces incomplete output that markets punish. That\'s exactly what we find.',
    'Policy implication: Uniform timing mandates create differential regulatory burden. Evidence-based policy would allow flexibility based on firm size and investigation progress.'
]

for i, point in enumerate(talking_points, 1):
    doc.add_paragraph(point, style='List Number')

doc.add_paragraph()

# MODEL 3
doc.add_heading('MODEL 3: Essay 3 - Governance Response (Executive Turnover)', level=1)

doc.add_heading('Visual Layout', level=2)
doc.add_paragraph(
    'Left-to-right flow: Data Breach Event → Mandatory Disclosure (Public announcement) → Stakeholder Activation (Pressure) → '
    'Governance Response (Executive Turnover with temporal decay shown)'
)

doc.add_heading('What This Model Shows', level=2)
doc.add_paragraph(
    'This model shows how mandatory disclosure functions as a stakeholder activation mechanism. Public announcement of a breach '
    'triggers investor and regulator scrutiny, which pressures boards to take governance action. However, this effect is transient—it peaks at 30 days then decays.'
)

doc.add_heading('Your Talking Points (30-45 seconds)', level=2)
talking_points = [
    'Essays 1 and 2 focused on MARKET OUTCOMES. Essay 3 asks: does mandatory disclosure affect ORGANIZATIONAL OUTCOMES, specifically '
    'executive governance changes?',
    'We find that disclosure timing activates executive turnover: FCC-regulated firms experience 5.3 percentage points higher '
    'executive departure rates within 30 days of breach disclosure.',
    'Mechanism: Public announcement makes the breach salient to board members, investors, and stakeholders. This creates pressure for '
    'board action (firing executives accountable for security failures).',
    'However, this effect is TRANSIENT: it peaks at 30 days then decays. This suggests crisis response rather than systematic reform—firms '
    'react urgently but don\'t embed lasting governance changes.'
]

for i, point in enumerate(talking_points, 1):
    doc.add_paragraph(point, style='List Number')

doc.add_paragraph()

# MODEL 4
doc.add_heading('MODEL 4: Dissertation Framework - Three Parallel Mechanisms', level=1)

doc.add_heading('Visual Layout', level=2)
doc.add_paragraph(
    'Central stimulus (FCC Rule 37.3 natural experiment, 2007 shock) branching into three parallel outcomes: '
    'Market Valuation (Essay 1, null result), Information Asymmetry (Essay 2, +1.83% volatility, with firm-size heterogeneity), '
    'and Governance Response (Essay 3, +5.3pp executive turnover). Bottom unifying theory box highlighting Tushman & Nadler information processing capacity constraint.'
)

doc.add_heading('What This Model Shows', level=2)
doc.add_paragraph(
    'This is the master framework showing how all three essays interconnect. It demonstrates that disclosure regulation does NOT '
    'work through a single mechanism, but rather through multiple independent channels with different outcomes and dynamics. '
    'The heterogeneous mechanism analysis adds a critical layer: firm-size heterogeneity in how regulatory constraints bind.'
)

doc.add_heading('Your Talking Points (1-2 minutes)', level=2)
talking_points = [
    '***Central Finding***: Disclosure regulation works through THREE DISTINCT MECHANISMS, not one.',
    '**Essay 1 (Market Valuation)**: Timing does not affect stock price reactions. Instead, regulatory status and breach characteristics matter. '
    'Markets efficiently price breach severity regardless of disclosure speed.',
    '**Essay 2 (Information Asymmetry)**: Timing DOES increase cost of capital through volatility (information asymmetry). Mandatory timing can '
    'force incomplete disclosure, which markets punish. Critical finding: this effect is concentrated in SMALL FIRMS (7.31% volatility increase) '
    'and reverses in LARGE FIRMS (-3.39%). Root cause: Information processing capacity constraints.',
    '**Essay 3 (Governance Response)**: Timing activates stakeholder pressure, which drives executive turnover. But the effect is crisis-driven and transient—not embedded reform.',
    '**Integration**: The three mechanisms operate independently. Timing doesn\'t help valuations, can harm information asymmetry (especially for small firms), '
    'but does activate governance response.',
    '**Theoretical foundation**: All three findings are consistent with information asymmetry theory (Akerlof, 1970), signaling theory (Myers & Majluf, 1984), '
    'and organizational information processing theory (Tushman & Nadler, 1978). Tushman & Nadler especially predicts firm-size heterogeneity in compliance burden.'
]

for i, point in enumerate(talking_points, 1):
    doc.add_paragraph(point, style='List Number')

doc.add_paragraph()

# POLICY IMPLICATIONS
doc.add_heading('Policy Implications from Mechanism Analysis', level=1)

doc.add_heading('Key Insight: Firm-Size Heterogeneity in Regulatory Burden', level=2)

policy_points = [
    'Current timing mandates apply uniformly regardless of firm size. But compliance capacity is NOT uniform.',
    'Small firms (Q1): FCC volatility effect = +7.31%*** → ~$37M additional annual cost of capital per breach',
    'Large firms (Q4): FCC volatility effect = -3.39%** → No cost, or possible benefit',
    'Regulatory burden differential: 11 percentage points between smallest and largest firms',
    'This represents a significant policy inequity for smaller public companies and smaller telecom firms.',
]

for i, point in enumerate(policy_points, 1):
    doc.add_paragraph(point, style='List Bullet')

doc.add_paragraph()

doc.add_heading('Recommended Policy Changes', level=2)

policy_recs = [
    'Replace uniform timing mandates with flexible timelines based on investigation progress',
    'Allow firms to disclose initial findings within 7 days, then supplement with complete findings as investigation progresses',
    'Prioritize INFORMATION QUALITY standards over SPEED metrics',
    'Scale timeline requirements by firm size to account for compliance capacity heterogeneity',
    'Provide regulatory safe harbor for good-faith disclosure that documents investigation process',
]

for i, point in enumerate(policy_recs, 1):
    doc.add_paragraph(point, style='List Bullet')

doc.add_paragraph()

# FINAL TALKING POINTS
doc.add_heading('Final Talking Points: Integrating All Models (2-3 minutes)', level=1)

final_points = [
    'This dissertation examines a simple policy question: Is faster disclosure always better? The answer is no—it depends on information quality, firm size, and market mechanism.',
    'Essay 1 shows: Timing doesn\'t affect valuations. Markets price information environment.',
    'Essay 2 shows: Mandatory timing can INCREASE market uncertainty through forced incompleteness. This effect is strongest for small firms with limited compliance capacity.',
    'Essay 3 shows: Timing activates stakeholder pressure and governance response, but only temporarily.',
    'Mechanism analysis reveals: The key driver is firm-size heterogeneity in information processing capacity. Small firms face binding constraints that large firms don\'t face.',
    'Policy implication: Evidence-based policy should move from "faster is always better" to "quality is the constraint." This requires flexibility in timing, investment in small-firm compliance capacity, and regulatory focus on information quality rather than speed compliance.',
    'The research demonstrates that one-size-fits-all timing mandates can create unintended consequences and impose disproportionate burdens on smaller firms. Evidence-based alternatives could reduce compliance costs while improving information quality and market efficiency.',
]

for i, point in enumerate(final_points, 1):
    doc.add_paragraph(point, style='List Number')

doc.add_paragraph()

# Save
doc.save('Conceptual_Models_Explanation_Updated.docx')
print("[OK] Saved: Conceptual_Models_Explanation_Updated.docx")
print("\nKey updates made:")
print("  - Added MODEL 2A: Firm Size Heterogeneity mechanism analysis")
print("  - Updated all talking points to reference mechanism findings")
print("  - Added policy implications section")
print("  - Integrated Tushman & Nadler information processing theory")
print("  - Added final integrated talking points")
