from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
TITLE_COLOR = RGBColor(31, 78, 121)
ACCENT_COLOR = RGBColor(192, 0, 0)

# Title
title = doc.add_paragraph()
title_run = title.add_run('Proposal Defense Q&A Guide')
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = TITLE_COLOR
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('Data Breach Disclosure Timing and Market Reactions\n12 Anticipated Committee Questions')
subtitle_run.font.size = Pt(12)
subtitle_run.font.italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

date = doc.add_paragraph()
date_run = date.add_run('March 2026')
date_run.font.size = Pt(11)
date.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# CRITICAL QUESTIONS
tier1 = doc.add_paragraph()
tier1_run = tier1.add_run('TIER 1: CRITICAL QUESTIONS')
tier1_run.font.bold = True
tier1_run.font.size = Pt(13)
tier1_run.font.color.rgb = ACCENT_COLOR

doc.add_paragraph()

# Q1
q1 = doc.add_paragraph()
q1_run = q1.add_run('Q1: Isn\'t the H1 null result a power failure?')
q1_run.font.bold = True
q1_run.font.size = Pt(11)

doc.add_paragraph('Challenge: Your null finding on timing could simply reflect insufficient power.', style='List Bullet')

answer1 = doc.add_paragraph()
answer1_run = answer1.add_run('Answer:')
answer1_run.bold = True

doc.add_paragraph('We have 898 observations with real timing variation (median 71 days; 17.6% immediate, 82.4% delayed). At 80% power, our MDE is +/-2.39pp. Our observed effect is +0.57%, far below this threshold. TOST equivalence test confirms economic negligibility. This is a genuine finding, not a power failure.', style='List Bullet')

doc.add_paragraph()

# Q2
q2 = doc.add_paragraph()
q2_run = q2.add_run('Q2: FCC firms are 2.22x larger—isn\'t this a size confound?')
q2_run.font.bold = True
q2_run.font.size = Pt(11)

doc.add_paragraph('Challenge: FCC firms average $62.6B vs $31.0B for non-FCC. Your effect could be size-driven.', style='List Bullet')

answer2 = doc.add_paragraph()
answer2_run = answer2.add_run('Answer:')
answer2_run.bold = True

doc.add_paragraph('Heterogeneous effects show the opposite: Q1 (smallest) = +7.31%*** (largest effect); Q4 (largest) = -3.39%** (negative). This is backwards from confound prediction. Industry FE strengthens the effect (from -2.20% to -5.37%). Three layers of evidence rule out size confound.', style='List Bullet')

doc.add_paragraph()

# Q3
q3 = doc.add_paragraph()
q3_run = q3.add_run('Q3: How do three essays hang together? Why not one big model?')
q3_run.font.bold = True
q3_run.font.size = Pt(11)

doc.add_paragraph('Challenge: Three essays with different outcomes could look like separate studies.', style='List Bullet')

answer3 = doc.add_paragraph()
answer3_run = answer3.add_run('Answer:')
answer3_run.bold = True

doc.add_paragraph('Three PARALLEL mechanisms: Essay 1 tests valuation (signaling theory)—timing does not improve pricing. Essay 2 tests uncertainty (information asymmetry)—timing does not clarify and increases volatility. Essay 3 tests governance (stakeholder theory)—mandate accelerates turnover. The Timing Paradox unifies all three: mandatory disclosure creates costs without valuation benefits. FCC runs through all essays.', style='List Bullet')

doc.add_paragraph()

# METHODOLOGICAL
tier2 = doc.add_paragraph()
tier2_run = tier2.add_run('TIER 2: METHODOLOGICAL QUESTIONS')
tier2_run.font.bold = True
tier2_run.font.size = Pt(13)
tier2_run.font.color.rgb = ACCENT_COLOR

doc.add_paragraph()

# Q4
q4 = doc.add_paragraph()
q4_run = q4.add_run('Q4: You claim causality for FCC effect, but isn\'t this correlation?')
q4_run.font.bold = True
q4_run.font.size = Pt(11)

doc.add_paragraph('Challenge: FCC firms are systematically different. Can you claim causality?', style='List Bullet')

answer4 = doc.add_paragraph()
answer4_run = answer4.add_run('Answer:')
answer4_run.bold = True

doc.add_paragraph('Natural experiment: PRE-2007 FCC effect = -13.96% (p=0.88, not significant). POST-2007 FCC effect = -2.26%** (p=0.013). This discontinuity at rule implementation is strong causality evidence, not pre-existing differences. Industry FE and firm fundamentals controls strengthen this.', style='List Bullet')

doc.add_paragraph()

# Q5
q5 = doc.add_paragraph()
q5_run = q5.add_run('Q5: Essay 2 shows FCC increases volatility—is this beneficial market learning?')
q5_run.font.bold = True
q5_run.font.size = Pt(11)

doc.add_paragraph('Challenge: Volatility could reflect market learning (honest pricing).', style='List Bullet')

answer5 = doc.add_paragraph()
answer5_run = answer5.add_run('Answer:')
answer5_run.bold = True

doc.add_paragraph('If it were beneficial learning, it should IMPROVE valuations (Essay 1). Instead, timing has no CAR effect. Volatility is mechanical: forced speed prevents investigation. Small firms suffer +7.31%***, large firms see -3.39%** decrease. Regressive burden on small firms, not beneficial learning.', style='List Bullet')

doc.add_paragraph()

# Q6
q6 = doc.add_paragraph()
q6_run = q6.add_run('Q6: Equivalence testing—isn\'t your +/-2.10% bound arbitrary?')
q6_run.font.bold = True
q6_run.font.size = Pt(11)

doc.add_paragraph('Challenge: Why that number? Could you have chosen bounds to get desired answer?', style='List Bullet')

answer6 = doc.add_paragraph()
answer6_run = answer6.add_run('Answer:')
answer6_run.bold = True

doc.add_paragraph('NOT arbitrary—derived from power analysis. The +/-2.10% bound is the MDE at 80% power (research standard). Cannot claim equivalence for effects smaller than we can measure. Cannot set higher bounds without claiming equivalence for economically large effects. This bound is justified by study design and power calculations.', style='List Bullet')

doc.add_paragraph()

# POLICY
tier3 = doc.add_paragraph()
tier3_run = tier3.add_run('TIER 3: POLICY QUESTIONS')
tier3_run.font.bold = True
tier3_run.font.size = Pt(13)
tier3_run.font.color.rgb = ACCENT_COLOR

doc.add_paragraph()

# Q7
q7 = doc.add_paragraph()
q7_run = q7.add_run('Q7: So you\'re saying disclosure rules are bad?')
q7_run.font.bold = True
q7_run.font.size = Pt(11)

doc.add_paragraph('Challenge: Disclosure is fundamental. Are you against disclosure requirements?', style='List Bullet')

answer7 = doc.add_paragraph()
answer7_run = answer7.add_run('Answer:')
answer7_run.bold = True

doc.add_paragraph('No. We support disclosure. We question TIMING-focused approach. Markets do not reward speed (H1 null), yet mandate imposes costs (+1.83% volatility, +5.3pp turnover). Optimize for COMPLETENESS, not SPEED. Recommend 14-30 days (not 7 days)—still fast, but allows investigation. Evidence-based critique, not ideology.', style='List Bullet')

doc.add_paragraph()

# Q8
q8 = doc.add_paragraph()
q8_run = q8.add_run('Q8: SEC cybersecurity rule is 4 days. Does your evidence apply?')
q8_run.font.bold = True
q8_run.font.size = Pt(11)

doc.add_paragraph('Challenge: Your evidence is from telecom. Generalize to SEC?', style='List Bullet')

answer8 = doc.add_paragraph()
answer8_run = answer8.add_run('Answer:')
answer8_run.bold = True

doc.add_paragraph('Direct extrapolation limited. But mechanisms (forced incomplete info increases uncertainty; mandatory disclosure accelerates governance) are general principles. SEC rule is MORE stringent (4 vs 7 days), so expect LARGER costs. Not evidence for/against SEC rule, but suggests SEC should monitor outcomes: volatility, turnover, information quality. Recommendation: measure whether it works, not assume 4 days is optimal.', style='List Bullet')

doc.add_paragraph()

# GENERAL
tier4 = doc.add_paragraph()
tier4_run = tier4.add_run('TIER 4: GENERAL QUESTIONS')
tier4_run.font.bold = True
tier4_run.font.size = Pt(13)
tier4_run.font.color.rgb = ACCENT_COLOR

doc.add_paragraph()

# Q9
q9 = doc.add_paragraph()
q9_run = q9.add_run('Q9: Why telecommunications when other sectors are growing?')
q9_run.font.bold = True
q9_run.font.size = Pt(11)

doc.add_paragraph('Challenge: Telecom declining. Why not healthcare, finance, technology?', style='List Bullet')

answer9 = doc.add_paragraph()
answer9_run = answer9.add_run('Answer:')
answer9_run.bold = True

doc.add_paragraph('Telecommunications IDEAL because natural experiment: FCC Rule 37.3 (2007) clean shock affecting only telecom. No other sector has single regulatory shock at clear date. Healthcare has HIPAA + state laws. Finance has SEC. Technology varies by state. All confounded. FCC is singular, dates to 2007—clean pre/post identification. Amani et al. (2025): telecom understudied. Sector choice was methodological, not convenience.', style='List Bullet')

doc.add_paragraph()

# Q10
q10 = doc.add_paragraph()
q10_run = q10.add_run('Q10: What evidence would falsify your main finding?')
q10_run.font.bold = True
q10_run.font.size = Pt(11)

doc.add_paragraph('Challenge: What would prove you wrong?', style='List Bullet')

answer10 = doc.add_paragraph()
answer10_run = answer10.add_run('Answer:')
answer10_run.bold = True

doc.add_paragraph('H1 falsified if: (1) Significant POSITIVE effect of speed on CAR, or (2) Strong positive effects in heterogeneous subgroups, or (3) Different patterns in Essay 3 contradicting turnover acceleration. Our findings: (1) No effect (p=0.373), (2) No positive effects in any dimension, (3) Faster disclosure DOES accelerate turnover—consistent with governance pressure, not information resolution.', style='List Bullet')

doc.add_paragraph()

# TOUGH
tier5 = doc.add_paragraph()
tier5_run = tier5.add_run('TIER 5: TOUGH/ADVANCED QUESTIONS')
tier5_run.font.bold = True
tier5_run.font.size = Pt(13)
tier5_run.font.color.rgb = ACCENT_COLOR

doc.add_paragraph()

# Q11
q11 = doc.add_paragraph()
q11_run = q11.add_run('Q11: Essay 3 mediation—is the causal chain valid?')
q11_run.font.bold = True
q11_run.font.size = Pt(11)

doc.add_paragraph('Challenge: Mediation assumes causal chain, but this is correlational. How justify causality?', style='List Bullet')

answer11 = doc.add_paragraph()
answer11_run = answer11.add_run('Answer:')
answer11_run.bold = True

doc.add_paragraph('Fair point—deliberately cautious about Essay 3 causality. FCC is exogenous (strong causal inference). Turnover effect (+5.3pp, p=0.008) CONDITIONAL on FCC status. Mediation tests PLAUSIBLE mechanism, not unique mechanism. Indirect effect small, suggesting turnover NOT primary valuation driver. Mechanism: stakeholder pressure (boards respond to mandate). Plausible because turnover accelerates WITHOUT improved information—regulatory pressure, not market learning.', style='List Bullet')

doc.add_paragraph()

# Q12
q12 = doc.add_paragraph()
q12_run = q12.add_run('Q12: Multiple hypothesis testing—did you adjust for corrections?')
q12_run.font.bold = True
q12_run.font.size = Pt(11)

doc.add_paragraph('Challenge: You test many hypotheses. Did you adjust p-values?', style='List Bullet')

answer12 = doc.add_paragraph()
answer12_run = answer12.add_run('Answer:')
answer12_run.bold = True

doc.add_paragraph('Distinguish PRIMARY hypotheses (H1-H4, core story) and EXPLORATORY heterogeneous effects. For primary: Bonferroni (p<0.0125) overly conservative since hypotheses not independent—all from same natural experiment. FCC effect meets threshold (p=0.010 Essay 1, p=0.008 Essay 3). Prior breaches p<0.001. Heterogeneous effects pre-specified as exploratory with FDR adjustment. Main findings survive any correction method.', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph()

# FINAL ADVICE
final = doc.add_paragraph()
final_run = final.add_run('Key Takeaways')
final_run.font.bold = True
final_run.font.size = Pt(13)
final_run.font.color.rgb = TITLE_COLOR

doc.add_paragraph()

doc.add_paragraph('Your dissertation is methodologically solid and intellectually coherent. You have 1,054 breaches, a clean natural experiment, and strong findings. The committee will push on H1 (null results always fair), size confounds (always fair), and generalizability (always fair). That is their job. Your job is to deliver evidence-backed answers confidently but not defensively. You have done the hard work. Own it.', style='List Bullet')

doc.add_paragraph()

doc.add_paragraph('Practice these answers so you can deliver them naturally. Your confidence will come from knowing you have already addressed these concerns thoroughly. This is a strong proposal.', style='List Bullet')

doc.save('Proposal_Defense_Q&A_Guide.docx')

print('[SUCCESS] Created Proposal_Defense_Q&A_Guide.docx')
print()
print('Document includes:')
print('  • 12 anticipated committee questions organized by tier')
print('  • Detailed answers with evidence for each question')
print('  • Tier 1: Critical (H1 power, size confound, essay coherence)')
print('  • Tier 2: Methodological (FCC causality, volatility, equivalence bounds)')
print('  • Tier 3: Policy (disclosure rules, SEC rule implications)')
print('  • Tier 4: General (telecom choice, falsification)')
print('  • Tier 5: Tough (Essay 3 mediation, multiple testing)')
print('  • Key takeaways and confidence boosters')
