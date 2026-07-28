"""
Update table references in ESSAY1_APPENDIX_TABLES_648.docx
Old Table 9-13 → New Table 10-14
"""

from docx import Document
import re

# Load document
doc = Document('outputs/ESSAY1_APPENDIX_TABLES_648.docx')

# Mapping of old → new table numbers
replacements = {
    'Table 13': 'Table 14',
    'Table 12': 'Table 13',
    'Table 11': 'Table 12',
    'Table 10': 'Table 11',
    'Table 9': 'Table 10',
}

# Track replacements
count = 0

# Process all paragraphs
for paragraph in doc.paragraphs:
    original = paragraph.text
    updated = paragraph.text

    # Apply replacements (in reverse order to avoid double-replacement)
    for old, new in replacements.items():
        # Use word boundaries to match whole table references
        pattern = r'\b' + re.escape(old) + r'\b'
        updated = re.sub(pattern, new, updated)

    # If text changed, update the paragraph
    if updated != original:
        count += 1
        # Clear existing runs and add new text
        for run in paragraph.runs:
            run.text = ''
        if paragraph.runs:
            paragraph.runs[0].text = updated
        else:
            paragraph.add_run(updated)
        print(f"Updated: {original} → {updated}")

# Save updated document
doc.save('outputs/ESSAY1_APPENDIX_TABLES_648.docx')

print(f"\nCompleted: {count} paragraphs updated")
print("File saved: outputs/ESSAY1_APPENDIX_TABLES_648.docx")
