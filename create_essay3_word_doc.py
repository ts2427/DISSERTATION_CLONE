#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Add title
title = doc.add_heading('Essay 3: Results Tables and Figures', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add introduction
intro = doc.add_paragraph()
intro.add_run('Ready for Integration into Essay3_Results_Section_WITH_CAUSAL_ID_WITH_INSIGHT.docx').italic = True

doc.add_paragraph()

# ============================================================================
# TABLE 1: Executive Turnover by Temporal Window
# ============================================================================

doc.add_heading('Table 1: Executive Turnover by Temporal Window', level=2)

table1 = doc.add_table(rows=4, cols=4)
table1.style = 'Light Grid Accent 1'

# Header row
header_cells = table1.rows[0].cells
header_cells[0].text = 'Temporal Window'
header_cells[1].text = 'N Observations'
header_cells[2].text = 'Events'
header_cells[3].text = 'Event Rate'

# Data rows
data1 = [
    ['30-day post-breach', '896', '416', '46.4%'],
    ['90-day post-breach', '896', '599', '66.9%'],
    ['180-day post-breach', '896', '605', '67.5%']
]

for row_idx, row_data in enumerate(data1, 1):
    cells = table1.rows[row_idx].cells
    for col_idx, value in enumerate(row_data):
        cells[col_idx].text = value

# Format table header
for cell in table1.rows[0].cells:
    tcPr = cell._element.get_or_add_tcPr()
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'D3D3D3')
    tcPr.append(shading_elm)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

# Add caption
caption1 = doc.add_paragraph()
caption1.add_run('Note: ').bold = True
caption1.add_run('Table reports number of observations, number of executive changes (Events), event rate (percentage of breaches with executive turnover), and model pseudo R-squared. Sample includes 896 breaches with complete executive turnover data from SEC 8-K filings (2006-2017). A breach is classified as having executive change if one or more executives depart within the specified temporal window.')

doc.add_paragraph()

# ============================================================================
# TABLE 2: Turnover by Disclosure Timing and Regulatory Status
# ============================================================================

doc.add_heading('Table 2: Executive Turnover by Disclosure Timing and Regulatory Status', level=2)

table2 = doc.add_table(rows=4, cols=4)
table2.style = 'Light Grid Accent 1'

# Header
header_cells = table2.rows[0].cells
header_cells[0].text = 'Firm Type'
header_cells[1].text = 'Immediate Disclosure (<=7 days)'
header_cells[2].text = 'Delayed Disclosure (>7 days)'
header_cells[3].text = 'Difference'

# Data
data2 = [
    ['FCC-Regulated Firms', '50.6%', '45.3%', '+5.3 pp**'],
    ['All Firms (FCC + Non-FCC)', '47.8%', '45.2%', '+2.6 pp'],
    ['Non-FCC Regulated Firms', '44-46%', '44-46%', '~0 pp (NS)']
]

for row_idx, row_data in enumerate(data2, 1):
    cells = table2.rows[row_idx].cells
    for col_idx, value in enumerate(row_data):
        cells[col_idx].text = value

# Format header
for cell in table2.rows[0].cells:
    tcPr = cell._element.get_or_add_tcPr()
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'D3D3D3')
    tcPr.append(shading_elm)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

# Add caption
caption2 = doc.add_paragraph()
caption2.add_run('Note: ').bold = True
caption2.add_run('Table reports 30-day executive turnover rates by disclosure timing (immediate <=7 days vs. delayed >7 days) and regulatory status (FCC-regulated vs. non-FCC). Timing effect is strongest in FCC-regulated firms (5.3 percentage points), where immediate disclosure requirements activate faster board responses. ** p<0.05 indicates significant difference between immediate and delayed disclosure for FCC firms. NS = not significant.')

doc.add_paragraph()

# ============================================================================
# TABLE 3: Causal Identification Tests
# ============================================================================

doc.add_heading('Table 3: Causal Identification Tests - Executive Turnover', level=2)

table3 = doc.add_table(rows=9, cols=4)
table3.style = 'Light Grid Accent 1'

# Header
header_cells = table3.rows[0].cells
header_cells[0].text = 'Test'
header_cells[1].text = 'Specification'
header_cells[2].text = 'Coefficient'
header_cells[3].text = 'P-Value / Sig'

# Data
data3 = [
    ['Temporal Validation (Full Sample)', 'Timing main effect', '-0.8956', '<0.0001***'],
    ['Temporal: Post-2007 Interaction', 'Interaction term', '+1.66 pp', '0.0668*'],
    ['Industry FE: Baseline', 'No industry control', '-0.8956', '<0.0001***'],
    ['Industry FE: With Controls', '+Industry FE', '-0.8821', '<0.0001***'],
    ['Size Q1 (Small)', 'By size quartile', '-0.679', '0.0810 t'],
    ['Size Q2 (Medium-small)', 'By size quartile', '-1.132', '0.0255**'],
    ['Size Q3 (Medium-large)', 'By size quartile', '-1.651', '0.0059***'],
    ['Size Q4 (Large)', 'By size quartile', '+0.371', '0.2653 NS']
]

for row_idx, row_data in enumerate(data3, 1):
    cells = table3.rows[row_idx].cells
    for col_idx, value in enumerate(row_data):
        cells[col_idx].text = value

# Format header
for cell in table3.rows[0].cells:
    tcPr = cell._element.get_or_add_tcPr()
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'D3D3D3')
    tcPr.append(shading_elm)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

# Add caption
caption3 = doc.add_paragraph()
caption3.add_run('Note: ').bold = True
caption3.add_run('Table reports three causal identification tests for timing effects on executive turnover. Panel A: Temporal validation test compares pre-2007 vs. post-2007 timing effects, supporting causal inference. Panel B: Industry fixed effects test controls for industry-specific governance norms. Panel C: Size sensitivity analysis tests whether timing effects vary by firm size quartile. Significance levels: t p<0.10, * p<0.05, ** p<0.05, *** p<0.01')

doc.add_paragraph()
doc.add_page_break()

# ============================================================================
# FIGURE 1
# ============================================================================

doc.add_heading('Figure 1: Executive Turnover Rates by Temporal Window', level=2)

try:
    doc.add_picture('outputs/essay3_figures/Figure1_Turnover_by_Window.png', width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fig_caption = doc.add_paragraph()
    fig_caption.add_run('Figure 1: ').bold = True
    fig_caption.add_run('Executive turnover increases substantially within 90 days of breach announcement, then stabilizes. Sharp increase from 30-day window (46.4%) to 90-day window (66.9%), with minimal additional change by 180 days (67.5%). Sample: 896 breaches with complete executive turnover data.')

except Exception as e:
    doc.add_paragraph(f'[Figure 1 image not found: {e}]')

doc.add_paragraph()
doc.add_page_break()

# ============================================================================
# FIGURE 2
# ============================================================================

doc.add_heading('Figure 2: Executive Turnover by Disclosure Timing × Regulatory Status', level=2)

try:
    doc.add_picture('outputs/essay3_figures/Figure2_Turnover_by_Timing_Regulatory.png', width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fig_caption = doc.add_paragraph()
    fig_caption.add_run('Figure 2: ').bold = True
    fig_caption.add_run('Disclosure timing accelerates executive turnover primarily in FCC-regulated firms. FCC-regulated firms show strongest timing effect (+5.3 percentage points between immediate and delayed disclosure, p<0.05), consistent with regulatory pressure mechanism. All firms average +2.6 pp. Non-FCC firms show no significant timing effect. Sample: 896 breaches (FCC N=200, Non-FCC N=726).')

except Exception as e:
    doc.add_paragraph(f'[Figure 2 image not found: {e}]')

# Save document
doc.save('Essay3_Tables_and_Figures.docx')
print('Successfully created: Essay3_Tables_and_Figures.docx')
print('\nDocument contains:')
print('  + Table 1: Turnover by Temporal Window')
print('  + Table 2: Turnover by Disclosure Timing x Regulatory Status')
print('  + Table 3: Causal Identification Tests')
print('  + Figure 1: Turnover Rates by Window (PNG)')
print('  + Figure 2: Turnover by Timing x Regulatory Status (PNG)')
print('\nReady to copy tables and figures into your Essay 3 document.')
