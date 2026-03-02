"""
Create Updated Speaker Notes with Policy Alternatives Section
Expanded speaker notes incorporating the 3 policy frameworks from comprehensive proposal
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt(14 if level == 1 else 12)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)

def add_para(doc, text, size=11, indent=0.5, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.size = Pt(size)
    return p

# Create document
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Title
p = doc.add_paragraph("SPEAKER NOTES: Dissertation Committee Presentation")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(14)
    run.font.bold = True

add_para(doc, "Data Breach Disclosure Timing and Market Reactions", size=12)
add_para(doc, f"Generated {datetime.now().strftime('%B %Y')}", size=10)

doc.add_paragraph()

# OPENING REMARKS
add_heading(doc, "Opening Remarks (Slide 1)", level=1)
add_para(doc, "Thank you for taking time to review this dissertation work. This research examines a fundamental question about regulation and market outcomes: does faster disclosure always produce better results? The FCC, SEC, and HHS all assume the answer is yes, requiring 7-day, 4-day, and 60-day breach disclosures respectively. But what if faster disclosure creates unintended costs that outweigh benefits?")

add_para(doc, "Using a natural experiment with the FCC Rule 37.3 (2007) as an exogenous shock, we study 1,054 publicly-traded firm data breaches from 2004-2025 to test this assumption. The evidence suggests that regulatory timing requirements achieve governance benefits but at significant market uncertainty costs, with estimated annual shareholder losses of approximately $9.9 billion in the telecommunications sector alone.")

# RESEARCH QUESTION
add_heading(doc, "Research Question & Motivation (Slides 2-3)", level=1)
add_para(doc, "The core research question asks: How do mandatory disclosure timing requirements and regulatory status affect firm valuation, market uncertainty, and governance response following data breaches?")

add_para(doc, "The motivation comes from three gaps in existing literature: First, no study exploits a clean regulatory natural experiment to isolate timing effects from firm quality signals. Second, the 'faster is better' assumption embedded in regulations has never been empirically tested at scale. Third, policymakers lack evidence on whether coordination versus differentiation across regulatory agencies (FCC 7-day, SEC 4-day, HHS 60-day) serves shareholders or creates inefficiency.")

add_para(doc, "The stakes are significant. If the FCC 7-day rule alone costs shareholders $9.9 billion annually, and similar rules exist across multiple regulators, the aggregate cost could exceed $20 billion per year. These costs are borne entirely by shareholders despite the regulation being justified on stakeholder protection grounds.")

# LITERATURE REVIEW OVERVIEW
add_heading(doc, "Literature Review (Slides 4-7)", level=1)
add_para(doc, "The literature review identifies four major streams: (1) Market Reactions to Breaches, (2) Mandatory Disclosure Paradoxes, (3) Information Asymmetry & Signaling Theory, and (4) Organizational Governance Response.")

add_para(doc, "Stream 1 shows that data breaches consistently produce negative returns (-0.3% to -2.1%), but the mechanisms underlying these returns are heterogeneous across firm and breach characteristics. Stream 2 reveals a critical paradox: mandatory disclosure laws sometimes increase crash risk and bad news hoarding, suggesting that forced disclosure creates unintended costs.")

add_para(doc, "Stream 3 grounds the analysis in Myers & Majluf (1984) signaling theory, predicting that mandatory timing requirements reduce signal quality by forcing disclosure before investigation completion. Stream 4 applies stakeholder theory to show how disclosure timing restructures regulatory stakeholder prioritization, activating governance response mechanisms.")

add_para(doc, "The synthesis suggests a testable prediction: mandatory timing requirements create a paradox where governance activation (positive) is offset by increased market uncertainty and reduced disclosure quality (negative).")

# HYPOTHESES
add_heading(doc, "Hypotheses & Model Specification (Slides 8-9)", level=1)
add_para(doc, "Six formal hypotheses structure the analysis:")

add_para(doc, "H1 (Essay 1): Disclosure timing has no significant effect on cumulative abnormal returns (CAR), supporting the view that markets eventually price breaches accurately regardless of timing.", size=10)

add_para(doc, "H2 (Essay 1): FCC-regulated firms experience more negative CAR following breaches, supporting the regulatory burden hypothesis.", size=10)

add_para(doc, "H3 (Essay 1): Prior breach history amplifies CAR effects, supporting reputation/signaling mechanisms.", size=10)

add_para(doc, "H4 (Essay 1): Health information breaches produce more negative CAR due to liability complexity.", size=10)

add_para(doc, "H5 (Essay 2): Mandatory disclosure timing increases return volatility relative to flexible timeline firms.", size=10)

add_para(doc, "H6 (Essay 3): Mandatory disclosure accelerates executive turnover relative to voluntary disclosure firms.", size=10)

add_para(doc, "The model specification uses Fama-French 3-factor adjusted returns for the main outcome (CAR), with controls for firm size, leverage, ROA, HHI, and industry/year fixed effects. All findings are tested with firm-clustered standard errors and robustness checks across alternative event windows, timing thresholds, and estimation approaches.")

# METHODS
add_heading(doc, "Methods & Causal Identification (Slides 10-11)", level=1)
add_para(doc, "The causal identification strategy exploits the FCC Rule 37.3 (2007) as a natural experiment. Telecommunications firms (SIC 4813, 4899, 4841) became subject to mandatory 7-day disclosure requirements, while other publicly-traded firms remained unregulated. This creates a difference-in-differences structure.")

add_para(doc, "The identification test shows that FCC coefficient estimates are essentially zero in the pre-2007 period and significant only post-2007, supporting the assumption that regulation, not selection, drives the effect. Industry fixed effects do not eliminate the FCC effect, suggesting the effect is not driven by industry-level selection into telecommunications.")

add_para(doc, "Event study methodology follows Brown & Warner (1985) and Fama & French (1993), with daily stock returns measured in a 30-day event window centered on breach disclosure date. We control for pre-event beta drift and use multiple model specifications (market model, Fama-French, Carhart) as sensitivity tests.")

# RESULTS OVERVIEW
add_heading(doc, "Key Findings (Slides 12-16)", level=1)
add_para(doc, "Essay 1 (Market Valuation): We find no significant effect of disclosure timing on final CAR (H1 not supported), but significant FCC effect of -2.91% initial CAR with +2.49% recovery, indicating temporary market confusion from regulatory burden rather than permanent valuation effect (H2 supported). Prior breach history amplifies effects by 1.5x (H3 supported). Health data breaches produce an additional -1.23% penalty (H4 supported).")

add_para(doc, "Essay 2 (Market Uncertainty): FCC-regulated firms show 34% higher return volatility in the 20-day post-breach window, and this increased volatility persists for up to 60 days post-breach. This supports H5: mandatory timing requirements increase market uncertainty even though final valuations converge.")

add_para(doc, "Essay 3 (Governance Response): Firms subject to FCC disclosure requirements show 2.8x higher executive turnover probability within 30 days of breach disclosure, compared to non-FCC firms. This supports H6: mandatory disclosure accelerates governance response. The mechanism appears to operate through stakeholder salience activation, with regulators becoming definitive stakeholders under mandatory disclosure regimes.")

# ROBUSTNESS
add_heading(doc, "Robustness & Limitations (Slides 17-18)", level=1)
add_para(doc, "We test robustness across: (1) alternative event windows (5, 10, 20, 30, 60-day), (2) different timing thresholds (same-day disclosure, 3-day, 7-day, 14-day cutoffs), (3) alternative standard errors (HC3, firm-clustered, industry-clustered), (4) fixed effects specifications (no FE, industry FE, industry-year FE), and (5) sample restrictions (excluding media-reported breaches, excluding pre-2010 observations, excluding large/small outliers).")

add_para(doc, "Limitations include: (1) FCC firms are larger on average, requiring size control interpretation; (2) causal chain from regulation → disclosure → market outcome relies on parallel trends assumption; (3) executive turnover data limited to documented cases in SEC filings; (4) results limited to publicly-traded firms, may not generalize to private companies.")

add_para(doc, "Despite these limitations, the natural experiment structure and multiple robustness checks support causal inference that regulatory timing requirements, not firm quality differences, drive the observed market and governance effects.")

# POLICY ALTERNATIVES
add_heading(doc, "Policy Alternatives & Recommendations (Slides 19-20)", level=1)
add_para(doc, "The empirical evidence suggests that current time-based disclosure mandates achieve governance benefits but at substantial market uncertainty costs. Three alternative policy frameworks merit consideration:")

add_heading(doc, "Alternative 1: Staged Disclosure Framework", level=2)
add_para(doc, "Require firms to disclose breach occurrence within 7 days, but allow extended investigation windows for final assessment disclosures. This preserves governance activation benefits (stakeholders become aware immediately) while allowing investigation completeness benefits (final disclosure contains thorough root cause analysis).")

add_para(doc, "Implementation: Modify FCC Rule 37.3 to distinguish between 'notice of breach' (7 days) and 'final assessment' (upon completion, with progress updates at 30/60-day intervals). Expected effect: reduced market uncertainty during investigation period while maintaining governance activation.")

add_para(doc, "Estimated savings: 40-50% reduction in temporary volatility costs, preserving governance response benefits. Shareholder cost savings: $4-5 billion annually in telecommunications sector.")

add_heading(doc, "Alternative 2: Quality Standards Framework", level=2)
add_para(doc, "Replace time-based mandates with quality standards: require disclosure within 7 days OR within 3 days of completing root cause investigation, whichever is later. This aligns disclosure timing with investigation completeness, eliminating penalties for time-bound incomplete disclosure.")

add_para(doc, "Implementation: Regulatory language shift from 'disclose within 7 days' to 'disclose completed root cause analysis, with required initial notice within 7 days.' Requires coordination across FCC, SEC, HHS but allows each agency to define 'completion' standards appropriate to their regulatory domain.")

add_para(doc, "Expected effect: firms naturally disclose earlier when investigation is quick, but avoid forced incomplete disclosure for complex breaches. Stakeholder trust improves because disclosure quality is guaranteed, not timing. Shareholder cost savings: $5-6 billion annually.")

add_heading(doc, "Alternative 3: Safe Harbor Framework", level=2)
add_para(doc, "Introduce liability protections (safe harbors) for firms that disclose early with preliminary information, protecting them from shareholder derivative suits or regulatory penalties for information gaps later clarified. This incentivizes early voluntary disclosure while eliminating penalties for disclosure incompleteness.")

add_para(doc, "Implementation: Statutory safe harbor (similar to Litigation Reform Act forward-looking statement protections) that protects preliminary breach disclosures from class action liability provided firm commits to updating disclosures. Requires Congressional action but does not require regulatory coordination.")

add_para(doc, "Expected effect: firms have incentive to disclose early (preserving governance benefits) without penalty for incompleteness (reducing market uncertainty). This leverages market mechanisms rather than regulatory mandates. Shareholder cost savings: $6-7 billion annually.")

add_para(doc, "Recommendation: Pursue quality standards framework as primary approach (maximum governance benefit with minimum regulatory disruption), with safe harbor as complementary instrument to incentivize voluntary early disclosure. Staged disclosure can serve as interim bridge framework while quality standards are implemented.")

add_heading(doc, "Closing Remarks", level=2)
add_para(doc, "This research demonstrates that regulatory assumptions about faster disclosure are empirically incorrect. Markets eventually price breach information accurately, but regulatory timing mandates create uncertainty costs that dwarf valuation benefits. The governance benefits of mandatory disclosure are real but modest compared to these costs.")

add_para(doc, "The three alternative frameworks offer policymakers evidence-based approaches to preserve governance benefits while eliminating unnecessary market uncertainty. Implementation would require coordination across FCC, SEC, and HHS, but the $9+ billion annual shareholder savings would justify this coordination investment.")

doc.save(r'C:\Users\mcobp\BA798_TIM\SPEAKER_NOTES_Updated.docx')
print("[OK] Updated speaker notes with policy alternatives created")
print("[OK] File: SPEAKER_NOTES_Updated.docx")
print("[OK] Added detailed talking points for 3 policy frameworks")
