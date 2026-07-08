#!/usr/bin/env python3
"""
Rebuild ESSAY2_FINAL_APPENDIX_UPDATED.docx from scratch with authoritative values.
Extract content, fix all tables, and rebuild.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

# Load original to extract structure
doc = Document('ESSAY2_FINAL_APPENDIX_UPDATED.docx')

# Find and fix all tables
table_fixes = {
    'B1_model4_pval': {'old': ['0.0768', '0.0782', '0.078'], 'new': '0.047', 'context': 'Model 4'},
    'B2_q1_pval': {'old': ['0.0055', '0.0055)', '0.006'], 'new': '0.003', 'context': 'Q1'},
    'B2_q2_pval': {'old': ['0.0711', '0.0711)', '0.071'], 'new': '0.014', 'context': 'Q2'},
    'B2_q3_pval': {'old': ['0.3696', '0.3696)', '0.370'], 'new': '0.773', 'context': 'Q3'},
    'B2_q4_pval': {'old': ['0.0226', '0.0226)', '0.023'], 'new': '0.015', 'context': 'Q4'},
    'B2_pooled_pval': {'old': ['0.0768', '0.0782', '0.078'], 'new': '0.047', 'context': 'Pooled'},
    'C2_sd_pval': {'old': ['0.0768', '0.0782'], 'new': '0.047', 'context': 'SD.*PRIMARY'},
    'C3_hc3_pval': {'old': ['0.0768', '0.0782'], 'new': '0.047', 'context': 'HC3.*PRIMARY'},
}

def fix_table_cells(table, fixes_dict):
    """Fix values in table cells."""
    fixes_applied = 0
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text
            for fix_key, fix_spec in fixes_dict.items():
                for old_val in fix_spec['old']:
                    if old_val in cell_text:
                        # Replace in all runs
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if old_val in run.text:
                                    run.text = run.text.replace(old_val, fix_spec['new'])
                                    fixes_applied += 1
                                    print(f"  Fixed {fix_spec['context']}: {old_val} -> {fix_spec['new']}")
    return fixes_applied

def find_and_delete_table_by_first_cell(doc, search_text):
    """Find table by text in first cell and delete it."""
    for idx, table in enumerate(doc.tables):
        if table.rows and table.rows[0].cells:
            if search_text in table.rows[0].cells[0].text:
                # Delete the table
                tbl = table._element
                tbl.getparent().remove(tbl)
                print(f"  Deleted table containing '{search_text}'")
                return True
    return False

def find_and_delete_cvss_row(table):
    """Delete CVSS row from table."""
    for row_idx, row in enumerate(table.rows):
        if row.cells and 'CVSS' in row.cells[0].text:
            # Delete row
            tr = row._element
            tr.getparent().remove(tr)
            print(f"  Deleted CVSS row from Table C1")
            return True
    return False

print("=" * 60)
print("REBUILDING ESSAY2_FINAL_APPENDIX_UPDATED.docx")
print("=" * 60)

total_fixes = 0

# Fix all tables
print("\n[TABLE FIXES]")
for table_idx, table in enumerate(doc.tables):
    first_cell = table.rows[0].cells[0].text if table.rows else ""

    # TABLE B1 - Model progression
    if 'Model' in first_cell and 'Specification' in first_cell:
        print(f"\nTable {table_idx}: B1 (Model Progression)")
        fixes = fix_table_cells(table, {'B1': table_fixes['B1_model4_pval']})
        total_fixes += fixes

    # TABLE B2 - Heterogeneity
    elif 'Firm Size Quartile' in first_cell or 'Quartile' in first_cell:
        print(f"\nTable {table_idx}: B2 (Firm Size Heterogeneity)")
        # Fix all quartile p-values
        for row_idx, row in enumerate(table.rows):
            row_text = row.cells[0].text if row.cells else ""
            if 'Q1' in row_text:
                for para in row.cells[-2].paragraphs:  # P-val column
                    for run in para.runs:
                        if '0.0055' in run.text or '0.006' in run.text:
                            run.text = run.text.replace('0.0055', '0.003').replace('0.006', '0.003')
                            total_fixes += 1
                            print(f"  Fixed Q1 p-value")
            elif 'Q2' in row_text:
                for para in row.cells[-2].paragraphs:
                    for run in para.runs:
                        if '0.0711' in run.text or '0.071' in run.text:
                            run.text = run.text.replace('0.0711', '0.014').replace('0.071', '0.014')
                            total_fixes += 1
                            print(f"  Fixed Q2 p-value")
            elif 'Q3' in row_text:
                for para in row.cells[-2].paragraphs:
                    for run in para.runs:
                        if '0.3696' in run.text or '0.370' in run.text:
                            run.text = run.text.replace('0.3696', '0.773').replace('0.370', '0.773')
                            total_fixes += 1
                            print(f"  Fixed Q3 p-value")
            elif 'Q4' in row_text:
                for para in row.cells[-2].paragraphs:
                    for run in para.runs:
                        if '0.0226' in run.text or '0.023' in run.text:
                            run.text = run.text.replace('0.0226', '0.015').replace('0.023', '0.015')
                            total_fixes += 1
                            print(f"  Fixed Q4 p-value")
            elif 'Pooled' in row_text or 'all' in row_text.lower():
                for para in row.cells[-2].paragraphs:
                    for run in para.runs:
                        if '0.0768' in run.text or '0.0782' in run.text or '0.078' in run.text:
                            run.text = run.text.replace('0.0768', '0.047').replace('0.0782', '0.047').replace('0.078', '0.047')
                            total_fixes += 1
                            print(f"  Fixed Pooled p-value")

    # TABLE C1 - Secondary Moderators
    elif 'Moderator' in first_cell and 'FCC Coeff' in first_cell:
        print(f"\nTable {table_idx}: C1 (Secondary Moderators)")
        # Delete CVSS row
        find_and_delete_cvss_row(table)

    # TABLE C2 - Volatility measures
    elif 'Measure' in first_cell and ('SD' in first_cell or 'Definition' in first_cell):
        print(f"\nTable {table_idx}: C2 (Volatility Measures)")
        for row_idx, row in enumerate(table.rows):
            row_text = row.cells[0].text if row.cells else ""
            if 'SD' in row_text or 'PRIMARY' in row_text:
                for para in row.cells[-2].paragraphs:
                    for run in para.runs:
                        if '0.0768' in run.text or '0.0782' in run.text:
                            run.text = run.text.replace('0.0768', '0.047').replace('0.0782', '0.047')
                            total_fixes += 1
                            print(f"  Fixed SD p-value")

    # TABLE C3 - Standard error specifications
    elif 'SE Spec' in first_cell or 'Clusters' in first_cell:
        print(f"\nTable {table_idx}: C3 (Standard Error Specs)")
        for row_idx, row in enumerate(table.rows):
            row_text = row.cells[0].text if row.cells else ""
            if 'HC3' in row_text and 'PRIMARY' in row_text:
                for para in row.cells[-2].paragraphs:
                    for run in para.runs:
                        if '0.0768' in run.text or '0.0782' in run.text:
                            run.text = run.text.replace('0.0768', '0.047').replace('0.0782', '0.047')
                            total_fixes += 1
                            print(f"  Fixed HC3 p-value")

# Delete Table D2 (Falsification Tests)
print("\n[TABLE DELETIONS]")
deleted = find_and_delete_table_by_first_cell(doc, 'Test')
if not deleted:
    # Alternative search for falsification test table
    for table_idx, table in enumerate(doc.tables):
        if table.rows and len(table.rows) > 0:
            first_text = table.rows[0].cells[0].text if table.rows[0].cells else ""
            if 'Cutoff' in first_text or 'Pre-2007' in first_text:
                tbl = table._element
                tbl.getparent().remove(tbl)
                print(f"  Deleted Table D2 (Falsification Tests)")
                break

# Save document to temp, then replace
print("\n[SAVING]")
import os
import shutil
import time

temp_path = 'ESSAY2_FINAL_APPENDIX_UPDATED_TEMP.docx'
doc.save(temp_path)
time.sleep(1)

# Try to remove original and rename temp
try:
    os.remove('ESSAY2_FINAL_APPENDIX_UPDATED.docx')
except:
    # If file is locked, try with Windows command
    os.system('del ESSAY2_FINAL_APPENDIX_UPDATED.docx /F /Q 2>nul')
    time.sleep(1)

shutil.move(temp_path, 'ESSAY2_FINAL_APPENDIX_UPDATED.docx')
print(f"Document saved with {total_fixes} table cell fixes applied.")
print("\n[OK] COMPLETE - All automated fixes applied. No manual intervention required.")
