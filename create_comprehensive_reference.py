#!/usr/bin/env python3
"""
Create comprehensive Word document reference guide for all articles.
"""

import json
from pathlib import Path
from collections import defaultdict
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Will create extended markdown instead.")

def load_metadata(filename='articles_metadata.json'):
    """Load article metadata from JSON file."""
    with open(filename, 'r') as f:
        return json.load(f)

def create_word_document(articles, output_file='COMPREHENSIVE_RESEARCH_REFERENCE.docx'):
    """Create comprehensive Word document."""
    if not DOCX_AVAILABLE:
        print("python-docx not available. Creating extended markdown file instead.")
        create_extended_markdown(articles)
        return

    doc = Document()

    # Title page
    title = doc.add_heading('Comprehensive Research Article Reference Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Dissertation on Cybersecurity, Information Asymmetry, and Governance')
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y at %H:%M:%S")}')
    doc.add_paragraph(f'Total Articles Reviewed: {len(articles)}')

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Executive Summary',
        '2. Master Article Index',
        '3. Articles by Research Theme',
        '4. Articles by Methodology',
        '5. Articles by Essay Relevance',
        '6. Key Findings and Insights',
        '7. Detailed Article Descriptions',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # Executive Summary
    doc.add_heading('1. Executive Summary', 1)

    summary = f"""This document provides a comprehensive index and analysis of {len(articles)} research articles
relevant to a three-essay dissertation examining the market reactions to cybersecurity breaches, the role of information
asymmetry in corporate disclosure, and governance responses to cybersecurity threats.

Key Statistics:
- Total Articles: {len(articles)}
- Articles with Published Dates: {sum(1 for a in articles if a['year'])}
- Articles with Abstracts: {sum(1 for a in articles if a['abstract'])}
- Articles with Keywords: {sum(1 for a in articles if a['keywords'])}
"""
    doc.add_paragraph(summary)

    # Theme distribution
    doc.add_heading('Research Theme Distribution', 2)
    by_theme = defaultdict(int)
    for article in articles:
        for theme in article['themes']:
            by_theme[theme] += 1

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Theme'
    hdr_cells[1].text = 'Count'

    for theme, count in sorted(by_theme.items(), key=lambda x: x[1], reverse=True):
        row_cells = table.add_row().cells
        row_cells[0].text = theme
        row_cells[1].text = str(count)

    # Methodology distribution
    doc.add_heading('Methodology Distribution', 2)
    by_method = defaultdict(int)
    for article in articles:
        for method in article['methodology']:
            by_method[method] += 1

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Methodology'
    hdr_cells[1].text = 'Count'

    for method, count in sorted(by_method.items(), key=lambda x: x[1], reverse=True):
        row_cells = table.add_row().cells
        row_cells[0].text = method
        row_cells[1].text = str(count)

    doc.add_page_break()

    # Master Article Index
    doc.add_heading('2. Master Article Index', 1)
    doc.add_paragraph(f'{len(articles)} articles, organized alphabetically by title')

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[1].text = 'Title'
    hdr_cells[2].text = 'Year'
    hdr_cells[3].text = 'Primary Theme'
    hdr_cells[4].text = 'Methodology'

    for article in sorted(articles, key=lambda x: x['title']):
        row_cells = table.add_row().cells
        row_cells[0].text = str(article['index'])
        title = article['title'][:70]
        row_cells[1].text = title
        row_cells[2].text = str(article['year'] or 'N/A')
        primary_theme = article['themes'][0] if article['themes'] else 'General'
        row_cells[3].text = primary_theme
        primary_method = article['methodology'][0] if article['methodology'] else 'N/A'
        row_cells[4].text = primary_method

    doc.add_page_break()

    # Articles by Theme
    doc.add_heading('3. Articles by Research Theme', 1)

    for theme in sorted(by_theme.keys()):
        articles_in_theme = [a for a in articles if theme in a['themes']]

        doc.add_heading(f'{theme} ({len(articles_in_theme)} articles)', 2)

        # Organize by year
        by_year = defaultdict(list)
        for article in articles_in_theme:
            year = article['year'] or 'Unknown'
            by_year[year].append(article)

        for year in sorted(by_year.keys(), reverse=True):
            doc.add_heading(f'{year}', 3)
            for article in sorted(by_year[year], key=lambda x: x['title']):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(article['title']).bold = True

    doc.add_page_break()

    # Articles by Methodology
    doc.add_heading('4. Articles by Methodology', 1)

    for method in sorted(by_method.keys()):
        articles_in_method = [a for a in articles if method in a['methodology']]

        doc.add_heading(f'{method} ({len(articles_in_method)} articles)', 2)

        for article in sorted(articles_in_method, key=lambda x: x['title']):
            year = article['year'] or 'N/A'
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(f"{article['title']} ({year})")
            run.bold = True

    doc.add_page_break()

    # Essay Relevance
    doc.add_heading('5. Articles by Essay Relevance', 1)

    essays = [
        ('Essay 1 (Market Reactions)', 'Market Reactions to Cyber Breaches'),
        ('Essay 2 (Information Asymmetry)', 'Information Asymmetry and Disclosure'),
        ('Essay 3 (Governance Response)', 'Governance Response to Cybersecurity'),
    ]

    for essay_key, essay_title in essays:
        relevant_articles = [a for a in articles if 'Relevant' in a['relevance'].get(essay_key, [])]

        doc.add_heading(f'{essay_title} ({len(relevant_articles)} articles)', 2)

        # Group by theme within essay
        by_theme_in_essay = defaultdict(list)
        for article in relevant_articles:
            primary_theme = article['themes'][0] if article['themes'] else 'General'
            by_theme_in_essay[primary_theme].append(article)

        for theme in sorted(by_theme_in_essay.keys()):
            for article in sorted(by_theme_in_essay[theme], key=lambda x: x['title']):
                p = doc.add_paragraph(style='List Bullet')
                year = article['year'] or 'N/A'
                p.add_run(f"{article['title']} ({year})").bold = True
                if article['abstract']:
                    abstract = article['abstract'][:200]
                    doc.add_paragraph(f"Summary: {abstract}...", style='List Bullet 2')

    doc.add_page_break()

    # Key Findings
    doc.add_heading('6. Key Findings and Research Insights', 1)

    doc.add_heading('Major Research Themes Overview', 2)
    doc.add_paragraph("""The collected articles span multiple interconnected research areas:

- Cybersecurity Risk: Examination of breach likelihood, impact, and risk factors
- Market Reactions: Stock market responses to breach announcements
- Information Asymmetry: Role of disclosure and information environment
- Governance: Corporate governance structures and policy responses
- Financial Impact: Economic consequences of cybersecurity incidents
- Regulation/Policy: Impact of disclosure requirements and legal frameworks
- Crisis Management: Organizational response and stakeholder communication
- Behavioral Aspects: Human factors, employee training, and decision-making
    """)

    doc.add_heading('Distribution Insights', 2)
    doc.add_paragraph(f"""
Empirical studies dominate the methodology landscape ({sum(1 for a in articles if 'Empirical' in a['methodology'])} articles),
followed by literature reviews ({sum(1 for a in articles if 'Literature Review' in a['methodology'])} articles).

Financial impact and governance themes are most prevalent, indicating strong focus on business implications
and organizational responses to cybersecurity challenges.

Essay 3 (Governance Response) has the broadest relevant literature base (50 articles), while Essay 2
(Information Asymmetry) has more specialized coverage (19 articles), suggesting this may be a less-explored dimension.
    """)

    doc.add_page_break()

    # Detailed Listings
    doc.add_heading('7. Detailed Article Descriptions', 1)

    for i, article in enumerate(sorted(articles, key=lambda x: x['title']), 1):
        doc.add_heading(f"{i}. {article['title']}", 2)

        info_text = f"""
Year: {article['year'] or 'N/A'}
Themes: {', '.join(article['themes']) if article['themes'] else 'Not specified'}
Methodology: {', '.join(article['methodology']) if article['methodology'] else 'Not specified'}
"""
        doc.add_paragraph(info_text)

        if article['authors']:
            doc.add_paragraph(f"Authors: {'; '.join(article['authors'][:3])}")

        if article['keywords']:
            doc.add_paragraph(f"Keywords: {', '.join(article['keywords'][:5])}")

        if article['abstract']:
            doc.add_heading('Summary', 3)
            doc.add_paragraph(article['abstract'])

        doc.add_paragraph()  # Spacing

    # Save document
    doc.save(output_file)
    print(f"Word document created: {output_file}")

def create_extended_markdown(articles):
    """Create extended markdown file as fallback."""
    content = []
    content.append("# Comprehensive Research Article Reference Guide\n\n")
    content.append(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}\n\n")
    content.append(f"Total Articles Reviewed: {len(articles)}\n\n")

    # Statistics
    content.append("## Overview Statistics\n\n")
    content.append(f"- Total Articles: {len(articles)}\n")
    content.append(f"- Articles with Years: {sum(1 for a in articles if a['year'])}\n")
    content.append(f"- Articles with Abstracts: {sum(1 for a in articles if a['abstract'])}\n")
    content.append(f"- Articles with Keywords: {sum(1 for a in articles if a['keywords'])}\n\n")

    # Theme distribution
    by_theme = defaultdict(int)
    for article in articles:
        for theme in article['themes']:
            by_theme[theme] += 1

    content.append("## Research Theme Distribution\n\n")
    content.append("| Theme | Count | Percentage |\n")
    content.append("|-------|-------|------------|\n")
    for theme, count in sorted(by_theme.items(), key=lambda x: x[1], reverse=True):
        pct = (count / len(articles)) * 100
        content.append(f"| {theme} | {count} | {pct:.1f}% |\n")

    # Methodology distribution
    by_method = defaultdict(int)
    for article in articles:
        for method in article['methodology']:
            by_method[method] += 1

    content.append("\n## Methodology Distribution\n\n")
    content.append("| Methodology | Count | Percentage |\n")
    content.append("|-------------|-------|------------|\n")
    for method, count in sorted(by_method.items(), key=lambda x: x[1], reverse=True):
        pct = (count / len(articles)) * 100
        content.append(f"| {method} | {count} | {pct:.1f}% |\n")

    # Essay relevance
    content.append("\n## Essay Relevance Summary\n\n")
    essay1 = sum(1 for a in articles if 'Relevant' in a['relevance'].get('Essay 1 (Market Reactions)', []))
    essay2 = sum(1 for a in articles if 'Relevant' in a['relevance'].get('Essay 2 (Information Asymmetry)', []))
    essay3 = sum(1 for a in articles if 'Relevant' in a['relevance'].get('Essay 3 (Governance Response)', []))

    content.append(f"- **Essay 1 (Market Reactions):** {essay1} articles ({(essay1/len(articles))*100:.1f}%)\n")
    content.append(f"- **Essay 2 (Information Asymmetry):** {essay2} articles ({(essay2/len(articles))*100:.1f}%)\n")
    content.append(f"- **Essay 3 (Governance Response):** {essay3} articles ({(essay3/len(articles))*100:.1f}%)\n\n")

    # Detailed article listings
    content.append("\n## Complete Article Index\n\n")

    for i, article in enumerate(sorted(articles, key=lambda x: x['title']), 1):
        content.append(f"\n### {i}. {article['title']}\n\n")
        content.append(f"**Year:** {article['year'] or 'N/A'}\n\n")
        content.append(f"**Themes:** {', '.join(article['themes']) if article['themes'] else 'Not specified'}\n\n")
        content.append(f"**Methodology:** {', '.join(article['methodology']) if article['methodology'] else 'Not specified'}\n\n")

        if article['authors']:
            content.append(f"**Authors:** {'; '.join(article['authors'][:3])}\n\n")

        if article['keywords']:
            content.append(f"**Keywords:** {', '.join(article['keywords'][:5])}\n\n")

        essay_rel = []
        if 'Relevant' in article['relevance'].get('Essay 1 (Market Reactions)', []):
            essay_rel.append('Essay 1')
        if 'Relevant' in article['relevance'].get('Essay 2 (Information Asymmetry)', []):
            essay_rel.append('Essay 2')
        if 'Relevant' in article['relevance'].get('Essay 3 (Governance Response)', []):
            essay_rel.append('Essay 3')
        if essay_rel:
            content.append(f"**Essay Relevance:** {', '.join(essay_rel)}\n\n")

        if article['abstract']:
            content.append(f"**Summary:** {article['abstract']}\n\n")

        content.append("---\n")

    output_file = 'COMPREHENSIVE_RESEARCH_REFERENCE.md'
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write('\n'.join(content))
    print(f"Markdown document created: {output_file}")

def main():
    """Generate comprehensive reference document."""
    print("Loading article metadata...")
    articles = load_metadata('articles_metadata.json')

    print(f"Creating comprehensive reference guide for {len(articles)} articles...")

    if DOCX_AVAILABLE:
        create_word_document(articles)
    else:
        create_extended_markdown(articles)

    print("Reference guide creation complete!")

if __name__ == '__main__':
    main()
