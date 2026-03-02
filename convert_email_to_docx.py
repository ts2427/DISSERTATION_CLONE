"""
Convert WORK_SUMMARY_EMAIL.md to professional Word document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_email_docx():
    """Generate Word document from email markdown."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('Dissertation Research Update', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Complete Presentation & Proposal Package')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.italic = True

    doc.add_paragraph()

    # Subject
    p = doc.add_paragraph()
    p.add_run('Subject: ').bold = True
    p.add_run('Dissertation Proposal Materials Complete - Comprehensive Research on Data Breach Disclosure Timing')

    doc.add_paragraph()

    # Overview
    doc.add_heading('Overview', 1)
    doc.add_paragraph(
        'This email summarizes the completion of a comprehensive dissertation proposal package, including a '
        'detailed written proposal, 28-slide presentation with speaker notes, and integration with an existing '
        'interactive dashboard. The research analyzes how regulatory disclosure timing requirements affect firm '
        'outcomes using a natural experiment approach.'
    )

    # What This Research Achieves
    doc.add_heading('What This Research Achieves', 1)

    doc.add_heading('Central Question', 2)
    p = doc.add_paragraph()
    p.add_run('"Is there any benefit to disclosing data breaches immediately, or should companies delay until '
              'investigation is complete?"').bold = True
    doc.add_paragraph(
        'This research tests a core assumption underlying all major cybersecurity disclosure regulations '
        '(FCC 7-day rule, SEC 4-day rule, HIPAA 60-day rule): that faster disclosure is better.'
    )

    doc.add_heading('The Finding', 2)
    p = doc.add_paragraph()
    p.add_run('Markets punish WHO YOU ARE and WHAT WAS BREACHED — not WHEN YOU TALK.').bold = True
    doc.add_paragraph('Rather than a single answer, the research reveals three independent mechanisms:')

    doc.add_paragraph('Valuation (Essay 1): Disclosure timing does NOT affect shareholder returns. Markets price '
                     'firm characteristics (prior breach history, regulatory status, data type), not announcement speed.',
                     style='List Bullet')
    doc.add_paragraph('Uncertainty (Essay 2): Mandatory disclosure timing INCREASES market volatility. Forced speed '
                     'prevents investigation completion, creating incomplete disclosures that markets recognize as quality '
                     'problems, raising uncertainty.',
                     style='List Bullet')
    doc.add_paragraph('Governance (Essay 3): Mandatory disclosure ACCELERATES executive turnover. Immediate public '
                     'disclosure triggers stakeholder pressure, forcing boards to respond with governance changes faster.',
                     style='List Bullet')

    doc.add_heading('Scope', 2)
    doc.add_paragraph('1,054 data breaches (2006-2025)', style='List Bullet')
    doc.add_paragraph('926 breaches with complete stock price data (87.9% match rate)', style='List Bullet')
    doc.add_paragraph('Natural experiment: FCC Rule 37.3 (effective January 1, 2007) creates regulatory discontinuity',
                     style='List Bullet')
    doc.add_paragraph('Treatment group: 200 FCC-regulated firms (telecom)', style='List Bullet')
    doc.add_paragraph('Control group: 854 firms in other industries', style='List Bullet')
    doc.add_paragraph('Economic impact: $0.76B in aggregate shareholder losses from FCC regulation alone',
                     style='List Bullet')

    # What Was Delivered
    doc.add_heading('What Was Delivered Today', 1)

    doc.add_heading('1. Dissertation_Proposal_Comprehensive.docx (58 KB, 25-35 pages)', 2)
    p = doc.add_paragraph()
    p.add_run('Purpose: ').bold = True
    p.add_run('Committee pre-reading document')

    doc.add_heading('Contents:', 3)
    doc.add_paragraph('Section 1 (Introduction): Problem statement, literature gap, research questions',
                     style='List Bullet')
    doc.add_paragraph('Section 2 (Literature Review): Four thematic streams with 30+ citations',
                     style='List Bullet')
    p = doc.add_paragraph('Stream 1: Empirical evidence on breach-driven market reactions',
                         style='List Bullet 2')
    p = doc.add_paragraph('Stream 2: Mandatory disclosure paradoxes and unintended effects',
                         style='List Bullet 2')
    p = doc.add_paragraph('Stream 3: Signaling theory and information asymmetry under time pressure',
                         style='List Bullet 2')
    p = doc.add_paragraph('Stream 4: Stakeholder theory and governance response mechanisms',
                         style='List Bullet 2')
    doc.add_paragraph('Section 3 (Hypotheses): Six testable hypotheses (H1-H6) with theoretical rationale',
                     style='List Bullet')
    doc.add_paragraph('Section 4 (Methods): Data sources and sample composition (1,054 breaches, detailed breakdown)',
                     style='List Bullet')
    doc.add_paragraph('Section 5 (Findings): All three essays\' actual results with effect sizes and significance levels',
                     style='List Bullet')
    doc.add_paragraph('Section 6 (Implications): Policy recommendations for FCC, SEC, HIPAA; firm strategy',
                     style='List Bullet')

    doc.add_heading('Key Features:', 3)
    doc.add_paragraph('Based entirely on actual research findings', style='List Bullet')
    doc.add_paragraph('All citations in APA format', style='List Bullet')
    doc.add_paragraph('Double-spaced, professional formatting', style='List Bullet')
    doc.add_paragraph('Explicitly addresses all PR review feedback', style='List Bullet')

    doc.add_heading('2. Dissertation_Presentation.pptx (64 KB, 28 slides)', 2)
    p = doc.add_paragraph()
    p.add_run('Purpose: ').bold = True
    p.add_run('In-person/Zoom presentation with visual storytelling')

    doc.add_heading('Slide Structure:', 3)
    doc.add_paragraph('Slides 1-4: Introduction (title, problem, why it matters, literature foundation)',
                     style='List Bullet')
    doc.add_paragraph('Slides 5-7: Research design (sample, natural experiment, validation tests)',
                     style='List Bullet')
    doc.add_paragraph('Slides 8-16: Hypotheses and findings (H1-H6, each with mechanism explanation)',
                     style='List Bullet')
    doc.add_paragraph('Slides 17-19: Integration (three independent mechanisms, heterogeneous effects, robustness)',
                     style='List Bullet')
    doc.add_paragraph('Slides 20-24: Economic significance and policy (FCC, SEC, optimal timeline, business implications)',
                     style='List Bullet')
    doc.add_paragraph('Slides 25-28: Limitations, takeaways, dashboard preview, Q&A',
                     style='List Bullet')

    doc.add_heading('Design:', 3)
    doc.add_paragraph('Zoom-friendly (large text 22-24pt minimum, clear hierarchy)', style='List Bullet')
    doc.add_paragraph('~1.5-2 minutes per slide (45-50 minute presentation)', style='List Bullet')
    doc.add_paragraph('Professional color scheme (dark blue headers, light blue metrics, red/green accents)',
                     style='List Bullet')
    doc.add_paragraph('Each finding on its own slide with mechanism explanation', style='List Bullet')

    doc.add_heading('3. Dissertation_Speaker_Notes.txt (50 KB, 1,240 lines)', 2)
    p = doc.add_paragraph()
    p.add_run('Purpose: ').bold = True
    p.add_run('Talking points for presenting all 28 slides')

    doc.add_heading('Contents:', 3)
    doc.add_paragraph('Per slide: Specific language to use, key points to emphasize, timing estimates, transitions',
                     style='List Bullet')
    doc.add_paragraph('Examples: Concrete numbers, stories, analogies', style='List Bullet')
    doc.add_paragraph('Contingencies: Q&A guidance for likely questions', style='List Bullet')
    doc.add_paragraph('General tips: Pacing advice (45-50 minutes), emphasis on storytelling, transitions',
                     style='List Bullet')
    doc.add_paragraph('Summary: Quick reference table for what to reference when', style='List Bullet')

    doc.add_heading('How to Use:', 3)
    doc.add_paragraph('Print or have on laptop during presentation', style='List Bullet')
    doc.add_paragraph('Glance while talking (audience sees slides, not notes)', style='List Bullet')
    doc.add_paragraph('Use as talking points, not verbatim script', style='List Bullet')
    doc.add_paragraph('Reference provided for deeper questions', style='List Bullet')

    # Addressing PR Review
    doc.add_heading('Addressing PR Review Feedback', 1)
    doc.add_paragraph('All five critical items from the code review have been explicitly addressed:')
    doc.add_paragraph()

    # Table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Item'
    header_cells[1].text = 'How Addressed'

    rows_data = [
        ('H1 null framing',
         'Proposal p.287-288 explains equivalence testing (TOST); Slide 9 interpretation; Speaker notes lead with "meaningful contribution"'),
        ('FCC causal ID strengthening',
         'Proposal p.178-198 covers pre/post-2007 test, industry FE strengthening effect, size heterogeneity Q1-Q4; Slide 7 visualizes three validation tests'),
        ('Clustered SEs',
         'Proposal p.232 specifies HC3 + firm-clustered; Speaker notes detail 6 SE specifications; Robustness script results referenced'),
        ('Essay 2 vs 3 tension',
         'Proposal p.34-49 & 402-415 frames "Timing Paradox"; Slide 17 explains three independent mechanisms; Speaker notes clarify distinction'),
        ('Policy bounds',
         'Proposal p.303-309 explicitly limits to "shareholder returns only"; Slides 21-25 bound implications; Speaker notes note what\'s OUT OF SCOPE'),
    ]

    for i, (item, addressed) in enumerate(rows_data, 1):
        cells = table.rows[i].cells
        cells[0].text = item
        cells[1].text = addressed

    doc.add_paragraph()

    # Integration
    doc.add_heading('Integration with Existing Work', 1)
    doc.add_paragraph('These new materials complement, not duplicate:')
    doc.add_paragraph('Proposal = Comprehensive written argument (25-35 pages)', style='List Bullet')
    doc.add_paragraph('Presentation = Visual storytelling (28 slides, 45-50 minutes)', style='List Bullet')
    doc.add_paragraph('Speaker Notes = Talking points for presenter', style='List Bullet')
    doc.add_paragraph('Dashboard = Interactive evidence explorer (15 pages, live filtering, detailed tables)',
                     style='List Bullet')

    doc.add_heading('Workflow for defense:', 2)
    doc.add_paragraph('Committee reads proposal beforehand', style='List Number')
    doc.add_paragraph('Presenter delivers 28-slide presentation (45-50 min) using speaker notes', style='List Number')
    doc.add_paragraph('Committee asks questions', style='List Number')
    doc.add_paragraph('Team explores Streamlit dashboard together (10-15 min) with filters by FCC status, firm size, breach type',
                     style='List Number')
    doc.add_paragraph('Q&A and synthesis (10-15 min)', style='List Number')

    # Key Findings
    doc.add_heading('Key Findings Summary', 1)

    doc.add_heading('What Drives Market Reactions (Essay 1):', 2)
    doc.add_paragraph('Prior breach history: -0.22% CAR per prior breach (p<0.001) — STRONGEST EFFECT',
                     style='List Number')
    doc.add_paragraph('FCC regulation: -2.20% CAR (p=0.010) — $0.76B aggregate cost', style='List Number')
    doc.add_paragraph('Health data: -2.51% CAR (p=0.004) — regulatory complexity premium', style='List Number')
    doc.add_paragraph('Disclosure timing: +0.57% CAR (p=0.539) — NOT SIGNIFICANT', style='List Number')

    doc.add_heading('What Increases Uncertainty (Essay 2):', 2)
    doc.add_paragraph('FCC regulation: +1.68% to +5.02% volatility increase', style='List Bullet')
    doc.add_paragraph('Mechanism: Forced speed prevents investigation → incomplete disclosure → higher uncertainty',
                     style='List Bullet')
    doc.add_paragraph('Size heterogeneity: Small firms +7.31% volatility (p<0.001); large firms -3.39% (p=0.024)',
                     style='List Bullet')
    doc.add_paragraph('Interpretation: Capacity constraints, not inherent information quality issues', style='List Bullet')

    doc.add_heading('What Triggers Governance Response (Essay 3):', 2)
    doc.add_paragraph('46.4% of breaches see 30-day executive turnover', style='List Bullet')
    doc.add_paragraph('Immediate disclosure accelerates this by +5.3 percentage points', style='List Bullet')
    doc.add_paragraph('Mechanism: Stakeholder pressure, not information resolution', style='List Bullet')
    doc.add_paragraph('Self-governance (46%) >> regulatory enforcement (0.6%)', style='List Bullet')

    doc.add_heading('Policy Implications:', 2)
    doc.add_paragraph('FCC 7-day rule: Creates $0.76B costs without valuation benefits; recommend 14-21 day extension',
                     style='List Bullet')
    doc.add_paragraph('SEC 4-day rule: Faces identical paradoxes; recommend 18-24 month evaluation post-implementation',
                     style='List Bullet')
    doc.add_paragraph('Optimal timeline: 45-60 days balances speed with investigation completeness (HIPAA model)',
                     style='List Bullet')
    doc.add_paragraph('For companies: Early disclosure accelerates governance response (good for accountability), doesn\'t protect stock price',
                     style='List Bullet')

    # Project Statistics
    doc.add_heading('Project Statistics', 1)
    doc.add_paragraph('Breaches analyzed: 1,054 (2006-2025)', style='List Bullet')
    doc.add_paragraph('CRSP-matched: 926 (87.9% match rate)', style='List Bullet')
    doc.add_paragraph('Regression specifications tested: 27+', style='List Bullet')
    doc.add_paragraph('Robustness methods: 6 different SE approaches', style='List Bullet')
    doc.add_paragraph('Subsamples: 8 (FCC, non-FCC, health, financial, repeat offenders, etc.)', style='List Bullet')
    doc.add_paragraph('Validation tests: 3 causal ID tests (pre/post, industry FE, size sensitivity)', style='List Bullet')
    doc.add_paragraph('Economic analysis scenarios: 5 (small firm, median, large, S&P 500, aggregate)',
                     style='List Bullet')
    doc.add_paragraph('Heterogeneous effect splits: Size quartiles (Q1-Q4)', style='List Bullet')
    doc.add_paragraph('Dashboard pages: 15 interactive pages', style='List Bullet')
    doc.add_paragraph('Presentation slides: 28', style='List Bullet')
    doc.add_paragraph('Speaker notes lines: 1,240', style='List Bullet')
    doc.add_paragraph('Proposal pages: 25-35', style='List Bullet')

    # Files in Repository
    doc.add_heading('Files in Repository', 1)
    doc.add_paragraph('All materials now in /DISSERTATION_CLONE/:')
    doc.add_paragraph()

    doc.add_heading('Main Documents:', 2)
    doc.add_paragraph('Dissertation_Proposal_Comprehensive.docx (58 KB) — Primary proposal document', style='List Bullet')
    doc.add_paragraph('Dissertation_Presentation.pptx (64 KB) — 28-slide presentation', style='List Bullet')
    doc.add_paragraph('Dissertation_Speaker_Notes.txt (50 KB) — Talking points for all slides', style='List Bullet')

    doc.add_heading('Supporting Scripts:', 2)
    doc.add_paragraph('create_proposal_comprehensive.py — Generates proposal (version-controlled)', style='List Bullet')
    doc.add_paragraph('create_presentation_expanded.py — Generates presentation (version-controlled)', style='List Bullet')

    doc.add_heading('Existing Dashboard:', 2)
    doc.add_paragraph('Dashboard/app.py — 15-page Streamlit application with interactive filtering', style='List Bullet')

    doc.add_heading('Data & Analysis:', 2)
    doc.add_paragraph('All underlying analysis in scripts/, Data/, and outputs/ directories', style='List Bullet')
    doc.add_paragraph('Complete audit trail with sample attrition documentation', style='List Bullet')

    # Next Steps
    doc.add_heading('Next Steps for Proposal Meeting', 1)

    doc.add_heading('Before Meeting:', 2)
    doc.add_paragraph('Committee receives Dissertation_Proposal_Comprehensive.docx for pre-reading', style='List Number')
    doc.add_paragraph('Presenter reviews Dissertation_Speaker_Notes.txt and practices pacing', style='List Number')
    doc.add_paragraph('Test Streamlit dashboard locally to ensure smooth filtering during Q&A', style='List Number')

    doc.add_heading('During Meeting:', 2)
    doc.add_paragraph('Present 28 slides (45-50 minutes) following speaker notes', style='List Number')
    doc.add_paragraph('Use slides as visual reference while talking (don\'t read them)', style='List Number')
    doc.add_paragraph('When FCC causality questioned → reference Slide 7 + Proposal p.178-198', style='List Number')
    doc.add_paragraph('When H1 null questioned → reference Slide 9 + Proposal p.287-288', style='List Number')
    doc.add_paragraph('When policy implications questioned → reference Slide 25 + Proposal p.303-309', style='List Number')
    doc.add_paragraph('For deep dives → launch dashboard and filter by FCC/size/breach type', style='List Number')

    doc.add_heading('Key Talking Points to Emphasize:', 2)
    doc.add_paragraph('"Markets don\'t reward speed — they reward completeness"', style='List Bullet')
    doc.add_paragraph('"The paradox: regulation forces speed but harms information quality"', style='List Bullet')
    doc.add_paragraph('"Three independent mechanisms: valuations don\'t respond to timing, but uncertainty and governance do"',
                     style='List Bullet')
    doc.add_paragraph('"The practical puzzle: why don\'t firms disclose early if it\'s beneficial? Answer: because markets don\'t reward it"',
                     style='List Bullet')

    # What Makes This Research Strong
    doc.add_heading('What Makes This Research Strong', 1)
    doc.add_paragraph('Scale: 1,054 breaches, 19 years of data, 926 firms — substantial dataset', style='List Number')
    doc.add_paragraph('Causal ID: FCC Rule 37.3 natural experiment with three validation tests (temporal, controls, heterogeneity)',
                     style='List Number')
    doc.add_paragraph('Robustness: 27+ specifications, 6 SE methods, multiple subsamples — not fragile to specification choices',
                     style='List Number')
    doc.add_paragraph('Mechanisms: Three independent essays test three different theoretical predictions about how regulation affects outcomes',
                     style='List Number')
    doc.add_paragraph('Economic significance: Findings translate to $0.76B-$1.94B real economic impact', style='List Number')
    doc.add_paragraph('Transparency: Complete audit trail documenting sample composition, validation, and robustness',
                     style='List Number')
    doc.add_paragraph('Presentation: Comprehensive proposal + visual presentation + speaker notes + interactive dashboard',
                     style='List Number')

    # Contact
    doc.add_heading('Contact & Questions', 1)
    doc.add_paragraph(
        'All files are version-controlled and reproducible. The scripts/ directory contains all analysis code, '
        'and the Dashboard/ directory contains the interactive exploration tool.'
    )
    doc.add_paragraph()
    doc.add_paragraph('For questions about any section:')
    doc.add_paragraph('Proposal content → See Dissertation_Proposal_Comprehensive.docx sections 1-6', style='List Bullet')
    doc.add_paragraph('Presentation flow → See Dissertation_Presentation.pptx slides 1-28', style='List Bullet')
    doc.add_paragraph('Talking points → See Dissertation_Speaker_Notes.txt (1,240 lines of detailed guidance)',
                     style='List Bullet')
    doc.add_paragraph('Interactive evidence → See Dashboard/app.py (15-page Streamlit dashboard)', style='List Bullet')
    doc.add_paragraph('Raw analysis → See scripts/80_essay1_regressions.py, scripts/81_essay2_volatility.py, scripts/82_essay3_governance.py',
                     style='List Bullet')

    # Summary
    doc.add_heading('Summary', 1)
    doc.add_paragraph(
        'This package represents a complete, polished, proposal-ready dissertation research project. The research '
        'is methodologically sound (natural experiment causal ID with validation), economically significant ($0.76B+ '
        'measured impact), theoretically grounded (integrating signaling theory, information asymmetry theory, and '
        'stakeholder theory), and clearly communicated (comprehensive proposal, visual presentation, interactive dashboard).'
    )
    doc.add_paragraph()
    doc.add_paragraph('You\'re ready for the proposal defense.')

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('Generated: February 27, 2026\n').italic = True
    footer.add_run('Research: Data Breach Disclosure Timing and Market Outcomes (1,054 breaches, 2006-2025)\n').italic = True
    footer.add_run('Contact: Timothy D. Spivey, University of South Alabama').italic = True

    # Save
    output_path = r'C:\Users\mcobp\BA798_TIM\Dissertation_Research_Update_Email.docx'
    doc.save(output_path)
    print(f"Email converted to Word document: {output_path}")


if __name__ == '__main__':
    create_email_docx()
