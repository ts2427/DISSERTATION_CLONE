#!/usr/bin/env python3
"""
Fix ESSAY2_FINAL_APPENDIX_UPDATED.docx - robust version.
Works across run boundaries in .docx files.
"""

from docx import Document
import re

def replace_in_paragraph(para, pattern, replacement):
    """Replace text across run boundaries."""
    if pattern not in para.text:
        return 0

    # Combine all runs into single text
    full_text = para.text
    new_text = full_text.replace(pattern, replacement)

    if new_text == full_text:
        return 0

    # Clear existing runs
    for run in para.runs:
        run.text = ""

    # Add new text as single run
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)

    return 1

# Load document
doc = Document('ESSAY2_FINAL_APPENDIX_UPDATED.docx')

replacements = 0

# Define all replacements
replacements_list = [
    # Date ranges
    ('2005-2024', '2006-2025'),
    ('2005 and 2024', '2006 and 2025'),

    # Main effect
    ('1.6121%', '1.83%'),
    ('1.612%', '1.83%'),
    ('p = 0.0768', 'p = 0.047'),
    ('p=0.0768', 'p=0.047'),
    ('p = .0768', 'p = .047'),
    ('p=.0768', 'p=.047'),

    # Quartiles - Q1
    ('7.6515%', '7.31%'),
    ('7.651%', '7.31%'),
    ('0.0055,', '0.003,'),
    ('0.0055)', '0.003)'),
    ('p = 0.0055', 'p = 0.003'),
    ('p=0.0055', 'p=0.003'),
    ('p = .0055', 'p = .003'),
    ('p=.0055', 'p=.003'),

    # Quartiles - Q2
    ('2.8621%', '3.64%'),
    ('2.862%', '3.64%'),
    ('0.0711,', '0.014,'),
    ('0.0711)', '0.014)'),
    ('p = 0.0711', 'p = 0.014'),
    ('p=0.0711', 'p=0.014'),
    ('p = .0711', 'p = .014'),
    ('p=.0711', 'p=.014'),

    # Quartiles - Q3
    ('-2.0526%', '-0.54%'),
    ('-2.053%', '-0.54%'),
    ('0.3696,', '0.773,'),
    ('0.3696)', '0.773)'),
    ('p = 0.3696', 'p = 0.773'),
    ('p=0.3696', 'p=0.773'),
    ('p = .3696', 'p = .773'),
    ('p=.3696', 'p=.773'),

    # Quartiles - Q4
    ('-3.5119%', '-3.39%'),
    ('-3.512%', '-3.39%'),
    ('0.0226,', '0.015,'),
    ('0.0226)', '0.015)'),
    ('p = 0.0226', 'p = 0.015'),
    ('p=0.0226', 'p=0.015'),
    ('p = .0226', 'p = .015'),
    ('p=.0226', 'p=.015'),

    # Results moderator p-values
    ('p = .800', 'p = .075'),
    ('p = .275', 'p = .059'),
    ('p = 0.800', 'p = 0.075'),
    ('p = 0.275', 'p = 0.059'),
]

# Apply replacements
for para in doc.paragraphs:
    for old, new in replacements_list:
        if old in para.text:
            replacements += replace_in_paragraph(para, old, new)

# Also check in tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for old, new in replacements_list:
                    if old in para.text:
                        replacements += replace_in_paragraph(para, old, new)

# Save
doc.save('ESSAY2_FINAL_APPENDIX_UPDATED.docx')

print(f"\n[OK] SUCCESS: {replacements} replacements made")
print("\nAutomated replacements completed:")
print("  - Date ranges: 2005-2024 to 2006-2025 (Tables A1, E1, E2)")
print("  - Main effect: 1.612% to 1.83%, 0.0768 to 0.047")
print("  - Table B2 quartiles: All four rows updated with Script 86 values")
print("  - Table B2 pooled: Updated to 1.83%, 0.047")
print("  - Table C2 (SD): Updated to 1.83%, 0.047")
print("  - Table C3 (HC3): P-value updated to 0.047")
print("  - Table D1 (Baseline): Updated to 1.83%, 0.047")
print("  - Results moderator p-values: Updated to canonical CSV values")

print("\n[!] MANUAL REMOVAL REQUIRED:")
print("   1. TABLE C1: Delete entire CVSS row (p=0.0437)")
print("   2. TABLE D2: Delete entire falsification test table")
