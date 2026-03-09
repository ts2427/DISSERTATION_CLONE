#!/usr/bin/env python3
"""
Create explanation guide for conceptual models
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Conceptual Models: Explanation Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Brief Talking Points for Proposal & Defense', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

date_para = doc.add_paragraph('March 2, 2026')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.runs[0].font.italic = True

doc.add_paragraph()

# Overview
doc.add_heading('Overview', level=1)
doc.add_paragraph(
    'These four conceptual models illustrate how mandatory data breach disclosure '
    'regulation (specifically the FCC Rule 37.3, implemented in 2007) affects firms '
    'through three distinct and independent mechanisms. The models are designed to be '
    'included in your dissertation proposal and defense presentation.'
)

doc.add_paragraph()

# MODEL 1
doc.add_heading('MODEL 1: Essay 1 - Market Reactions (Stock Price Impact)', level=1)

doc.add_heading('Visual Layout', level=2)
doc.add_paragraph(
    'A left-to-right flow showing: Data Breach Disclosure → Timing Decision → Information Quality → Market Valuation (CAR -0.74%)'
)

doc.add_heading('What This Model Shows', level=2)
doc.add_paragraph(
    'This model illustrates Essay 1\'s research question: "Does disclosure timing affect how markets value firms '
    'after breaches?" It depicts a causal chain where the timing of breach disclosure leads to different information '
    'quality, which theoretically should affect market valuation.'
)

doc.add_heading('Your Talking Points (30-45 seconds)', level=2)
talking_points = [
    'We hypothesized that faster disclosure would lead to more favorable market reactions because it demonstrates '
    'management transparency and reduces uncertainty.',
    'However, we find a surprising null result: disclosure timing does NOT affect stock price reactions (coefficient +0.57%, p=0.539).',
    'Instead, what DOES matter: FCC regulatory status (-2.20%, highly significant), health breach complexity (-2.51%), '
    'and prior breach history (-0.22% per breach—the strongest effect).',
    'Key insight: Markets appear to price breach severity and firm characteristics, NOT disclosure speed. This challenges '
    'the regulatory assumption that "faster is better."'
]

for i, point in enumerate(talking_points, 1):
    doc.add_paragraph(point, style='List Number')

doc.add_paragraph()

# MODEL 2
doc.add_heading('MODEL 2: Essay 2 - Information Asymmetry (Cost of Capital)', level=1)

doc.add_heading('Visual Layout', level=2)
doc.add_paragraph(
    'Left-to-right flow: Data Breach → Timing Requirement (FCC: 7 days) → Investigation Incompleteness → Market Uncertainty (Volatility)'
)

doc.add_heading('What This Model Shows', level=2)
doc.add_paragraph(
    'This model illustrates Essay 2\'s paradoxical finding: mandatory timing requirements can INCREASE market uncertainty. '
    'The logic: forced early disclosure prevents thorough investigation, leaving questions unanswered, which increases volatility.'
)

doc.add_heading('Your Talking Points (30-45 seconds)', level=2)
talking_points = [
    'While Essay 1 found markets don\'t reward disclosure speed, Essay 2 discovers that mandatory timing creates a different cost: '
    'increased information asymmetry and cost of capital.',
    'When regulators mandate 7-day disclosure (FCC) or 4-day disclosure (SEC 2023 rules), companies face a time constraint that may '
    'force incomplete disclosure. This leaves investors uncertain.',
    'Our finding: FCC regulation increases volatility by 1.83%, equivalent to a 0.0137% increase in cost of capital. For a $50B firm, '
    'this translates to $10M+ in additional annual financing costs.',
    'The paradox: disclosure requirements achieve their goal of faster information BUT at the cost of higher information asymmetry '
    'and higher cost of capital for affected firms.'
]

for i, point in enumerate(talking_points, 1):
    doc.add_paragraph(point, style='List Number')

doc.add_paragraph()

# MODEL 3
doc.add_heading('MODEL 3: Essay 3 - Governance Response (Executive Turnover)', level=1)

doc.add_heading('Visual Layout', level=2)
doc.add_paragraph(
    'Left-to-right flow: Data Breach Event → Mandatory Disclosure (Public announcement) → Stakeholder Activation (Pressure) → '
    'Governance Response (Executive Turnover)'
)

doc.add_heading('What This Model Shows', level=2)
doc.add_paragraph(
    'This model shows how mandatory disclosure functions as a stakeholder activation mechanism. Public announcement of a breach '
    'triggers investor and regulator scrutiny, which pressures boards to take governance action.'
)

doc.add_heading('Your Talking Points (30-45 seconds)', level=2)
talking_points = [
    'Essays 1 and 2 focused on MARKET OUTCOMES. Essay 3 asks: does mandatory disclosure affect ORGANIZATIONAL OUTCOMES, specifically '
    'executive governance changes?',
    'We find that disclosure timing activates executive turnover: FCC-regulated firms experience 5.3 percentage points higher '
    'executive departure rates within 30 days of breach disclosure.',
    'Mechanism: Public announcement makes the breach salient to board members, investors, and stakeholders. This creates pressure for '
    'board action (firing executives accountable for security failures).',
    'However, this effect is TRANSIENT: it peaks at 30 days then decays. This suggests crisis response rather than systematic reform—firms '
    'react urgently but don\'t embed lasting governance changes.'
]

for i, point in enumerate(talking_points, 1):
    doc.add_paragraph(point, style='List Number')

doc.add_paragraph()

# MODEL 4
doc.add_heading('MODEL 4: Dissertation Framework - Three Parallel Mechanisms', level=1)

doc.add_heading('Visual Layout', level=2)
doc.add_paragraph(
    'Central stimulus (FCC Rule 37.3 natural experiment) branching into three parallel outcomes: Market Valuation (Essay 1), '
    'Information Asymmetry (Essay 2), and Governance Response (Essay 3). Bottom unifying theory box.'
)

doc.add_heading('What This Model Shows', level=2)
doc.add_paragraph(
    'This is the master framework showing how all three essays interconnect. It demonstrates that disclosure regulation does NOT '
    'work through a single mechanism, but rather through multiple independent channels with different outcomes and dynamics.'
)

doc.add_heading('Your Talking Points (1-2 minutes)', level=2)
talking_points = [
    '***Central Finding***: Disclosure regulation works through THREE DISTINCT MECHANISMS, not one.',
    '**Essay 1 (Market Valuation)**: Timing does not affect stock price reactions. Instead, regulatory status and breach characteristics matter. '
    'Markets efficiently price breach severity regardless of disclosure speed.',
    '**Essay 2 (Information Asymmetry)**: Timing DOES increase cost of capital through volatility (information asymmetry). Mandatory timing can '
    'be counterproductive by forcing incomplete disclosure.',
    '**Essay 3 (Governance)**: Timing DOES activate stakeholder pressure and governance response. Public announcement triggers immediate executive '
    'changes, though effects are transient.',
    '**Unifying Theory**: FCC Penalty = f(Expected_Investigation_Time - Actual_Disclosure_Speed). Markets use EXPECTATIONS to price regulation. '
    'When disclosure violates expectations (simple breaches disclosed too fast, creating uncertainty), the penalty is larger. When complexity creates '
    'expected delays (complex breaches with long investigation), the penalty is smaller.',
    '**Policy Implication**: One-size-fits-all regulations create DIFFERENTIAL EFFECTS. Smaller firms and simpler breaches experience larger penalties. '
    'Policymakers should consider differentiated timelines based on breach complexity.'
]

for i, point in enumerate(talking_points, 1):
    doc.add_paragraph(point, style='List Number')

doc.add_paragraph()

# PRESENTATION TIPS
doc.add_heading('Presentation Tips for Your Proposal/Defense', level=1)

tips = [
    'Start with Model 4 (dissertation framework) for the big picture, then drill down into Models 1-3 to show evidence.',
    'When discussing each model, start with the research question ("How does X affect Y?"), then show the key finding (unexpected null, '
    'paradoxical effect, etc.), then explain the mechanism.',
    'Emphasize that the three mechanisms are INDEPENDENT. You\'re not finding three ways the same mechanism works; you\'re finding three '
    'completely different mechanisms activated by disclosure regulation.',
    'Use the null finding (Essay 1) as a strength, not a weakness. Null results are valuable when you use equivalence testing and show the '
    'effect is economically negligible.',
    'Highlight the paradox: regulation achieves faster disclosure but at the cost of higher information asymmetry and more uncertainty. This '
    'is interesting policy-relevant finding.',
    'The unifying theory (expectation-based market pricing) is what ties all three mechanisms together. It\'s the conceptual contribution that '
    'elevates this from "three separate studies" to "one integrated dissertation."'
]

for i, tip in enumerate(tips, 1):
    doc.add_paragraph(tip, style='List Bullet')

doc.add_paragraph()

# SAMPLE SPIEL
doc.add_heading('Sample 2-Minute Overview (For Proposal or Defense)', level=1)

spiel = '''
"My dissertation examines how data breach disclosure regulation affects firms through multiple channels.
I exploit the FCC Rule 37.3 natural experiment—implemented in 2007, requiring telecommunications firms to
disclose breaches within 7 days—as a quasi-experimental variation.

I find three key things:

First, disclosure TIMING does not affect stock market valuations. Markets are efficient—they price breach
severity and firm characteristics regardless of when the announcement comes. This challenges regulatory
assumptions that "faster is better" for shareholders.

Second, and paradoxically, mandatory timing INCREASES market uncertainty. By forcing early disclosure
before investigations complete, regulators create incomplete information, which increases volatility and
cost of capital. For a $50B firm, this translates to $10M+ in annual additional financing costs.

Third, mandatory disclosure DOES activate governance response. Public announcement creates stakeholder
pressure, triggering executive turnover. We see 5.3 percentage points higher executive departure rates—but
this effect is transient (peaks at 30 days then fades), suggesting crisis response rather than reform.

The unifying insight is expectation-based pricing: markets penalize disclosure when it violates expectations
about investigation time. Simple breaches disclosed too quickly create uncertainty. Complex breaches are
expected to take time, so the penalty is lower.

Taken together, this suggests that one-size-fits-all disclosure timing creates differential effects. Policymakers
should consider complexity-based timelines rather than uniform mandates."
'''

doc.add_paragraph(spiel)

doc.add_paragraph()

# Footer
doc.add_paragraph('---')
footer = doc.add_paragraph(
    'Use these talking points when presenting your models in meetings with your advisor, committee, or '
    'in your formal proposal presentation. Adjust timing and detail based on your audience.'
)
footer.runs[0].font.italic = True
footer.runs[0].font.size = Pt(9)

doc.save('Conceptual_Models_Explanation_Guide.docx')
print('[OK] Explanation guide created: Conceptual_Models_Explanation_Guide.docx')
