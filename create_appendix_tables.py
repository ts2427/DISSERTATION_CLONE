from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Set up margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ============================================================================
# TABLE B8
# ============================================================================

title_b8 = doc.add_heading('TABLE B8: POST-2007 INTERACTION TEST', level=2)
title_b8.alignment = WD_ALIGN_PARAGRAPH.LEFT

subtitle_b8 = doc.add_paragraph('Isolating Regulatory Effect from Industry Effect')
subtitle_b8_format = subtitle_b8.paragraph_format
subtitle_b8_format.left_indent = Inches(0)
for run in subtitle_b8.runs:
    run.italic = True
    run.font.size = Pt(11)

doc.add_paragraph()

# Create TABLE B8
table_b8 = doc.add_table(rows=4, cols=6)
table_b8.style = 'Light Grid Accent 1'

# Header row
header_cells_b8 = table_b8.rows[0].cells
headers_b8 = ['Model', 'N', 'FCC Coefficient', 'Std Error', 'P-Value', 'R²']
for i, header in enumerate(headers_b8):
    header_cells_b8[i].text = header

# Data rows
data_b8 = [
    ['Model 1: Full Sample (2004-2025)', '898', '-2.2994%', '0.8935', '0.0101', '0.0201'],
    ['Model 3: Post-2007 (2007+)', '894', '-2.2557%', '0.9034', '0.0125', '0.0204'],
    ['FCC x Post-2007 Interaction', '', '-2.2294%', '(Implied)', '0.0125', '0.0226']
]

for row_idx, row_data in enumerate(data_b8, 1):
    cells = table_b8.rows[row_idx].cells
    for col_idx, value in enumerate(row_data):
        cells[col_idx].text = value

# Format header
for cell in table_b8.rows[0].cells:
    tcPr = cell._element.get_or_add_tcPr()
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'D3D3D3')
    tcPr.append(shading_elm)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

# Add notes for B8
doc.add_paragraph()
notes_b8 = doc.add_paragraph()
notes_b8.add_run('Notes: ').bold = True
notes_b8_text = 'Dependent variable is 30-day cumulative abnormal returns (CAR). FCC regulation (47 CFR § 64.2011) became effective in 2007. The FCC main effect (pre-2007) is not statistically significant (p=0.8818), indicating FCC-regulated firms were not different from non-FCC firms before the rule took effect. The post-2007 FCC effect emerges as significantly negative (p=0.0125), with an implied coefficient of -2.2294%, indicating that the regulatory penalty materializes after the rule was implemented. This temporal pattern supports a causal interpretation: the FCC penalty results from regulatory constraints, not from pre-existing industry characteristics or selection effects. Standard errors are heteroskedasticity-consistent (HC3). Significance levels: * p<0.10, ** p<0.05, *** p<0.01.'
notes_b8.add_run(notes_b8_text)
notes_b8_format = notes_b8.paragraph_format
notes_b8_format.left_indent = Inches(0.25)
for run in notes_b8.runs:
    run.font.size = Pt(10)

doc.add_paragraph()
doc.add_page_break()

# ============================================================================
# TABLE B9
# ============================================================================

title_b9 = doc.add_heading('TABLE B9: STANDARD ERROR SPECIFICATION COMPARISON', level=2)
title_b9.alignment = WD_ALIGN_PARAGRAPH.LEFT

subtitle_b9 = doc.add_paragraph('HC3 Heteroskedasticity-Consistent vs. Firm-Level Clustered Standard Errors')
subtitle_b9_format = subtitle_b9.paragraph_format
subtitle_b9_format.left_indent = Inches(0)
for run in subtitle_b9.runs:
    run.italic = True
    run.font.size = Pt(11)

doc.add_paragraph()

# Create TABLE B9
table_b9 = doc.add_table(rows=9, cols=8)
table_b9.style = 'Light Grid Accent 1'

# Header row
header_cells_b9 = table_b9.rows[0].cells
headers_b9 = ['Variable', 'Model A (HC3) Coef', 'Model A SE', 'Model A P-Value',
              'Model B (Clustered) Coef', 'Model B SE', 'Model B P-Value', 'Sig Change?']
for i, header in enumerate(headers_b9):
    header_cells_b9[i].text = header

# Data rows
data_b9 = [
    ['immediate_disclosure', '1.1059', '0.8428', '0.1895', '1.1059', '0.9430', '0.2409', 'NO'],
    ['fcc_reportable', '-2.7568', '0.9051', '0.0023***', '-2.7568', '1.5693', '0.0790*', 'YES'],
    ['prior_breaches_1yr', '-0.2156', '0.0399', '0.0000***', '-0.2156', '0.0681', '0.0015***', 'NO'],
    ['health_breach', '-1.6693', '0.8103', '0.0394**', '-1.6693', '0.9080', '0.0660*', 'YES'],
    ['firm_size_log', '0.4198', '0.2891', '0.1465', '0.4198', '0.4731', '0.3749', 'NO'],
    ['leverage', '1.7788', '1.0371', '0.0863*', '1.7788', '1.3095', '0.1743', 'YES'],
    ['roa', '19.5074', '7.2224', '0.0069***', '19.5074', '7.9940', '0.0147**', 'YES'],
]

for row_idx, row_data in enumerate(data_b9, 1):
    cells = table_b9.rows[row_idx].cells
    for col_idx, value in enumerate(row_data):
        cells[col_idx].text = value

# Format header
for cell in table_b9.rows[0].cells:
    tcPr = cell._element.get_or_add_tcPr()
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'D3D3D3')
    tcPr.append(shading_elm)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(9)

# Add info row
info_row = table_b9.rows[8]
info_cell = info_row.cells[0]
info_cell.text = 'R-squared'
info_row.cells[1].text = '0.0501'

# Add notes for B9
doc.add_paragraph()
notes_b9 = doc.add_paragraph()
notes_b9.add_run('Notes: ').bold = True
notes_b9_text = 'Dependent variable is 30-day cumulative abnormal returns (CAR). N = 898 observations across 374 unique firms. Model A uses heteroskedasticity-consistent (HC3) standard errors. Model B employs firm-level clustering to account for multiple breaches at the same firm. Coefficients are identical across both models; only standard errors differ. Firm-level clustering increases standard errors substantially (average increase of 73% for the FCC coefficient), making the specification more conservative. Despite larger standard errors, the primary findings remain robust: the FCC effect (fcc_reportable) stays significant at p<0.10, prior breach history remains highly significant, and health data breaches remain significant. The timing effect (immediate_disclosure) remains non-significant in both specifications, confirming that the null finding is not an artifact of standard error specification choice. The identical R-squared values confirm that clustering does not change model fit, only coefficient precision. Significance levels: * p<0.10, ** p<0.05, *** p<0.01.'
notes_b9.add_run(notes_b9_text)
notes_b9_format = notes_b9.paragraph_format
notes_b9_format.left_indent = Inches(0.25)
for run in notes_b9.runs:
    run.font.size = Pt(10)

# Save document
doc.save('Essay1_Appendix_Tables_B8_B9.docx')
print('Created: Essay1_Appendix_Tables_B8_B9.docx')
print('\nDocument contains:')
print('  + TABLE B8: Post-2007 Interaction Test')
print('  + TABLE B9: Standard Error Specification Comparison')
print('  + Full explanatory notes for both tables')
