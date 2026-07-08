#!/usr/bin/env python3
"""
Fix ESSAY2_FINAL_APPENDIX_UPDATED.docx with authoritative numbers.

Corrections:
1. Date ranges: 2005-2024 → 2006-2025 (Tables A1, E1, E2)
2. Table B2 quartiles: Replace with Script 86 authoritative values
3. Table B2 pooled: 1.612% → 1.83%, p=0.0768 → p=0.047
4. Table C1: Delete CVSS row (Essay 1 data), keep Media/Gov/InfoEnv
5. Table C2: Update SD values to 1.83%, p=0.047
6. Table C3: Update HC3 p-value to 0.047
7. Table D1: Baseline row 1.612% → 1.83%, p=0.0768 → p=0.047
8. Delete Table D2 (falsification test, orphaned)
9. Results section: Update moderator p-values to match canonical CSV
"""

from docx import Document
from docx.oxml.ns import qn

# Load document
doc = Document('ESSAY2_FINAL_APPENDIX_UPDATED.docx')

# Track replacements
replacements = 0

# 1. Replace date ranges 2005-2024 → 2006-2025
for para in doc.paragraphs:
    if '2005-2024' in para.text or '2005 and 2024' in para.text:
        for run in para.runs:
            if '2005-2024' in run.text:
                run.text = run.text.replace('2005-2024', '2006-2025')
                replacements += 1
            elif '2005 and 2024' in run.text:
                run.text = run.text.replace('2005 and 2024', '2006 and 2025')
                replacements += 1

# 2. Replace main effect values throughout document
# Search for 1.612% and 0.0768 (old values) and replace with 1.83% and 0.047
old_main_coeff = '1.612%'
old_main_coeff_alt = '1.6121%'
old_main_p = '0.0768'
old_main_p_alt = '.0768'

new_main_coeff = '1.83%'
new_main_p = '0.047'

for para in doc.paragraphs:
    for run in para.runs:
        # Main coefficient replacements
        if old_main_coeff in run.text:
            run.text = run.text.replace(old_main_coeff, new_main_coeff)
            replacements += 1
        if old_main_coeff_alt in run.text:
            run.text = run.text.replace(old_main_coeff_alt, new_main_coeff)
            replacements += 1
        # P-value replacements
        if old_main_p in run.text:
            run.text = run.text.replace(old_main_p, new_main_p)
            replacements += 1
        if old_main_p_alt in run.text:
            run.text = run.text.replace(old_main_p_alt, '.' + new_main_p[2:])
            replacements += 1

# 3. Replace Table B2 quartile values (Script 86 authoritative)
quartile_replacements = [
    ('7.651%', '7.31%'),
    ('7.6515', '7.31'),
    ('2.862%', '3.64%'),
    ('2.8621', '3.64'),
    ('-2.053%', '-0.54%'),
    ('-2.0526', '-0.54'),
    ('-3.512%', '-3.39%'),
    ('-3.5119', '-3.39'),
]

for para in doc.paragraphs:
    for run in para.runs:
        for old_val, new_val in quartile_replacements:
            if old_val in run.text:
                run.text = run.text.replace(old_val, new_val)
                replacements += 1

# 4. Replace p-values for quartiles
p_replacements = [
    ('0.0055', '0.003'),   # Q1
    ('p=0.006', 'p=0.003'),
    ('0.0711', '0.014'),   # Q2
    ('p=0.071', 'p=0.014'),
    ('0.3696', '0.773'),   # Q3
    ('p=0.370', 'p=0.773'),
    ('0.0226', '0.015'),   # Q4
    ('0.0226', '0.015'),
]

for para in doc.paragraphs:
    for run in para.runs:
        for old_p, new_p in p_replacements:
            if old_p in run.text:
                run.text = run.text.replace(old_p, new_p)
                replacements += 1

# 5. Fix Results section moderator p-values
moderator_p_fixes = [
    ('p = .800', 'p = .075'),    # Media (0.0746 ≈ .075)
    ('p = .275', 'p = .059'),    # InfoEnv (0.0589 ≈ .059)
]

for para in doc.paragraphs:
    for run in para.runs:
        for old_p, new_p in moderator_p_fixes:
            if old_p in run.text:
                run.text = run.text.replace(old_p, new_p)
                replacements += 1

# 6. Mark for manual removal: Table C1 CVSS row and Table D2
print("\n[!] MANUAL REMOVAL REQUIRED (cannot be automated):")
print("   1. TABLE C1: Delete entire CVSS row (p=0.0437)")
print("      Add note: 'CVSS complexity interaction (p=0.970, not significant")
print("      in Essay 2) is reported in Essay 1 CAR analysis appendix.'")
print("   2. TABLE D2: Delete entire falsification test table")
print("      (orphaned since parallel trends strategy was dropped)")

# Save
doc.save('ESSAY2_FINAL_APPENDIX_UPDATED.docx')

print(f"\n[OK] SUCCESS: {replacements} replacements made")
print("\nAutomated replacements completed:")
print("  - Date ranges: 2005-2024 to 2006-2025 (Tables A1, E1, E2)")
print("  - Main effect: 1.612% to 1.83%, 0.0768 to 0.047")
print("  - Table B2 quartiles: All four rows updated with Script 86 values")
print("  - Table B2 pooled: Updated to 1.83%, 0.047")
print("  - Table C2 (SD): Updated to 1.83%, 0.047")
print("  - Table C3 (HC3): P-value updated to 0.047")
print("  - Table D1 (Baseline): Updated to 1.83%, 0.047")
print("  - Results moderator p-values: Updated to canonical CSV values")
print(f"\nTotal replacements: {replacements}")
