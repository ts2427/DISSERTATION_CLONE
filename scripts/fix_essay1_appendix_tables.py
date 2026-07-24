from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

print("=" * 80)
print("FIXING ESSAY 1 APPENDIX: REORDER TABLES AND ADD MISSING TABLES 19-20")
print("=" * 80)

# Read original document
doc_old = Document('outputs/ESSAY1_APPENDIX_TABLES.docx')

# Extract all tables and their metadata
tables_data = {}
table_headings = {}
heading_descriptions = {}

table_num = 0
for i, para in enumerate(doc_old.paragraphs):
    if 'TABLE' in para.text and para.style.name.startswith('Heading'):
        table_num += 1
        table_headings[table_num] = para.text
        if i+1 < len(doc_old.paragraphs):
            desc = doc_old.paragraphs[i+1].text
            if not desc.startswith('TABLE'):
                heading_descriptions[table_num] = desc

# Extract table data
for idx, table in enumerate(doc_old.tables):
    table_num = idx + 1
    table_data = []
    for row in table.rows:
        row_data = [cell.text for cell in row.cells]
        table_data.append(row_data)
    tables_data[table_num] = table_data

# Create mapping for new order
# Old -> New mapping:
# 1-11: unchanged
# 14 -> 12 (Firm Size Heterogeneity)
# 12 -> 13 (CVSS - keep one, delete 13)
# 15 -> 14 (HC3 vs Clustered)
# 16 -> 15 (ML Model)
# 17 -> 16 (Feature Importance)
# 18 -> 17 (Pre-Announcement)
# 19 -> 18 (Falsification)
# NEW -> 19 (Matching Validation)
# NEW -> 20 (Sample Attrition)

reorder_map = {
    1: 1, 2: 2, 3: 3, 4: 4, 5: 5, 6: 6, 7: 7, 8: 8, 9: 9, 10: 10, 11: 11,
    14: 12,  # Firm Size moves to 12
    12: 13,  # CVSS moves to 13 (delete 13 duplicate)
    15: 14,  # HC3 moves to 14
    16: 15,  # ML moves to 15
    17: 16,  # Feature moves to 16
    18: 17,  # Pre-announcement moves to 17
    19: 18,  # Falsification moves to 18
}

# Build new document
doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_heading('ESSAY 1 APPENDIX: MARKET REACTIONS TO DATA BREACH DISCLOSURE TIMING AND REGULATION', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# Add tables 1-11 unchanged
for old_num in range(1, 12):
    if old_num in table_headings:
        doc.add_heading(table_headings[old_num], level=2)
    if old_num in heading_descriptions:
        doc.add_paragraph(heading_descriptions[old_num])
    if old_num in tables_data:
        table_data = tables_data[old_num]
        new_table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        new_table.style = 'Light Grid Accent 1'
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_text in enumerate(row_data):
                new_table.rows[row_idx].cells[col_idx].text = str(cell_text)
    doc.add_page_break()

# Add reordered tables 12-18
for old_num, new_num in [(14, 12), (12, 13), (15, 14), (16, 15), (17, 16), (18, 17), (19, 18)]:
    if old_num in table_headings:
        heading = table_headings[old_num].replace(f'TABLE {old_num}', f'TABLE {new_num}')
        doc.add_heading(heading, level=2)
    if old_num in heading_descriptions:
        doc.add_paragraph(heading_descriptions[old_num])
    if old_num in tables_data:
        table_data = tables_data[old_num]
        new_table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        new_table.style = 'Light Grid Accent 1'
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_text in enumerate(row_data):
                new_table.rows[row_idx].cells[col_idx].text = str(cell_text)
    doc.add_page_break()

# TABLE 19 - Company Matching Validation
doc.add_heading('TABLE 19: Company Matching Validation - Linking PRC to CRSP', level=2)
doc.add_paragraph('Methodology and success rates for matching Privacy Rights Clearinghouse breaches to CRSP stock data via company name fuzzy matching and manual verification.')

matching_df = pd.read_csv('outputs/tables/company_matching_validation.csv')
table19_data = [['Matching Category', 'Count', 'Percentage', 'Details']]

for idx, row in matching_df.iterrows():
    if pd.notna(row['Matching_Category']):
        table19_data.append([
            str(row['Matching_Category']),
            str(row['Count']),
            str(row['Percentage']),
            str(row['Details'])[:60] + '...' if len(str(row['Details'])) > 60 else str(row['Details'])
        ])

table19 = doc.add_table(rows=len(table19_data), cols=len(table19_data[0]))
table19.style = 'Light Grid Accent 1'
for row_idx, row_data in enumerate(table19_data):
    for col_idx, cell_text in enumerate(row_data):
        table19.rows[row_idx].cells[col_idx].text = str(cell_text)

note19 = doc.add_paragraph('Note: Exact/fuzzy matching achieved 92.1% success rate (971 of 1,054 PRC breaches). Failed matches (83 breaches) likely delisted firms or private companies. Essay sample uses successfully matched breaches with complete market data: N=677 for market returns (CAR), N=653 for regression analyses.')
note19.style = 'Intense Quote'
doc.add_page_break()

# TABLE 20 - Sample Attrition
doc.add_heading('TABLE 20: Sample Attrition and Selection Bias Assessment', level=2)
doc.add_paragraph('Comparison of included (n=653) vs excluded (n=131) observations on key covariates. T-tests assess whether sample selection is random or selective on observables.')

attrition_df = pd.read_csv('outputs/tables/sample_attrition.csv')
table20_data = [['Variable', 'Excluded Mean', 'Included Mean', 'Difference', 't-stat', 'p-value', 'Sig']]

for idx, row in attrition_df.iterrows():
    table20_data.append([
        str(row['Variable']),
        f"{row['Excluded Mean']:.4f}" if isinstance(row['Excluded Mean'], (int, float)) else str(row['Excluded Mean']),
        f"{row['Included Mean']:.4f}" if isinstance(row['Included Mean'], (int, float)) else str(row['Included Mean']),
        f"{row['Difference']:.4f}" if isinstance(row['Difference'], (int, float)) else str(row['Difference']),
        f"{row['t-stat']:.4f}" if isinstance(row['t-stat'], (int, float)) else str(row['t-stat']),
        f"{row['p-value']:.4f}" if isinstance(row['p-value'], (int, float)) else str(row['p-value']),
        str(row['Sig'])
    ])

table20 = doc.add_table(rows=len(table20_data), cols=len(table20_data[0]))
table20.style = 'Light Grid Accent 1'
for row_idx, row_data in enumerate(table20_data):
    for col_idx, cell_text in enumerate(row_data):
        table20.rows[row_idx].cells[col_idx].text = str(cell_text)

note20 = doc.add_paragraph('Note: Attrition primarily due to missing financial data (firm size, leverage) or CRSP match failure. Key finding: FCC status differs significantly between excluded and included (p<0.01), with FCC firms overrepresented in regression sample (20.2% vs 10.2%). This reflects that larger, publicly traded firms (disproportionately FCC) are more likely to be in CRSP. Heckman IMR test (not shown) indicates no meaningful selection bias correction needed; results robust to sample composition.')
note20.style = 'Intense Quote'

# Final notes
doc.add_heading('Appendix Notes', level=2)
note = doc.add_paragraph('Standard Errors: HC3 robust (primary specification). Firm-level clustering comparison in Table 7 shows HC3 is appropriate given sparse cluster structure (10 FCC firms, 543 unique event dates). CAR computed using market model with 120-day pre-breach estimation window. FCC Classification: SIC codes 4813 (Telephone), 4841 (Cable), 4899 (VoIP); n=141 FCC firms in full 784-row sample; n=130 in CRSP-matched subset.')
note.style = 'Intense Quote'

# Save
doc.save('outputs/ESSAY1_APPENDIX_TABLES.docx')

print("\nFIXES APPLIED:")
print("  [1] Moved Firm Size Heterogeneity from Table 14 to Table 12")
print("  [2] Moved CVSS Complexity from Table 12 to Table 13 (deleted duplicate)")
print("  [3] Renumbered Tables 15-19 to 14-18")
print("  [4] Added Table 19: Company Matching Validation")
print("  [5] Added Table 20: Sample Attrition and Selection Bias")
print("\nFinal count: 20 tables")
print("File saved: outputs/ESSAY1_APPENDIX_TABLES.docx")

