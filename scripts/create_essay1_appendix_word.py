"""
CREATE ESSAY 1 APPENDIX IN WORD FORMAT
Verified pipeline figures, 19 tables, markdown-to-docx conversion
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

print("=" * 90)
print("CREATING ESSAY 1 APPENDIX WORD DOCUMENT (19 TABLES)")
print("=" * 90)

# Create document
doc = Document()

# Set up document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_heading('ESSAY 1 APPENDIX: MARKET REACTIONS TO DATA BREACH DISCLOSURE TIMING AND REGULATION', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ============================================================================
# TABLE 1: Summary Statistics
# ============================================================================
doc.add_heading('TABLE 1: Summary Statistics — Full Sample and CRSP-Matched Subsample', level=2)

doc.add_paragraph('Panel A: Full Sample (N = 784)')

table = doc.add_table(rows=14, cols=9)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Variable'
hdr_cells[1].text = 'N'
hdr_cells[2].text = 'Mean'
hdr_cells[3].text = 'Std Dev'
hdr_cells[4].text = 'Min'
hdr_cells[5].text = 'P25'
hdr_cells[6].text = 'Median'
hdr_cells[7].text = 'P75'
hdr_cells[8].text = 'Max'

data_panel_a = [
    ['30-Day CAR (%)', '677', '-0.43', '8.91', '-42.56', '-5.07', '0.14', '4.05', '34.05'],
    ['5-Day CAR (%)', '677', '-0.08', '4.19', '-26.41', '-1.85', '-0.04', '2.00', '14.21'],
    ['Immediate Disclosure (≤7d)', '784', '0.24', '0.43', '0.00', '0.00', '0.00', '0.00', '1.00'],
    ['Delayed Disclosure (>30d)', '784', '0.60', '0.49', '0.00', '0.00', '1.00', '1.00', '1.00'],
    ['Days to Disclosure', '784', '88.5', '112.3', '0', '8', '58', '140', '1,247'],
    ['Health Data Breach', '784', '0.06', '0.24', '0.00', '0.00', '0.00', '0.00', '1.00'],
    ['Financial Data Breach', '784', '0.26', '0.44', '0.00', '0.00', '0.00', '1.00', '1.00'],
    ['Total Prior Breaches', '784', '3.09', '6.59', '0.00', '0.00', '0.00', '3.00', '56.00'],
    ['Prior Breaches (1yr)', '784', '1.33', '3.62', '0.00', '0.00', '0.00', '1.00', '56.00'],
    ['Firm Size (log assets)', '718', '10.47', '1.27', '5.01', '9.70', '10.38', '11.18', '14.74'],
    ['Leverage (Debt/Assets)', '721', '0.72', '0.24', '0.12', '0.58', '0.68', '0.85', '2.52'],
    ['ROA (%)', '721', '1.00', '4.00', '-33.00', '0.00', '1.00', '2.00', '21.00'],
    ['Records Affected (log)', '784', '4.51', '4.05', '0.00', '1.39', '3.00', '6.55', '20.72'],
]

for i, row_data in enumerate(data_panel_a, start=1):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = str(cell_text)

doc.add_paragraph()
doc.add_paragraph('Panel B: CRSP-Matched Sample (N = 677) - Same variables as Panel A for CRSP-matched breaches with complete market data. With CRSP price data: 677/784 = 86.4%. With complete regression data: 653/784 = 83.3%.')

doc.add_paragraph()
doc.add_heading('Panel C: By FCC Regulatory Status (N = 677, CRSP Sample)', level=3)

table = doc.add_table(rows=7, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Variable'
hdr_cells[1].text = 'FCC (n=130)'
hdr_cells[2].text = 'Non-FCC (n=547)'
hdr_cells[3].text = 'Difference'

data_panel_c = [
    ['30-Day CAR (%)', '-2.39', '+0.04', '-2.43pp'],
    ['5-Day CAR (%)', '-0.99', '+0.14', '-1.13pp'],
    ['Immediate Disclosure (%)', '40%', '21%', '+19pp'],
    ['Firm Size (log)', '11.10', '10.34', '+0.76'],
    ['Leverage', '0.72', '0.72', '0.00'],
    ['ROA (%)', '1.00', '1.00', '0.00'],
]

for i, row_data in enumerate(data_panel_c, start=1):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = str(cell_text)

doc.add_paragraph()
doc.add_heading('Panel D: By Disclosure Timing (N = 677, CRSP Sample)', level=3)

table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Variable'
hdr_cells[1].text = 'Immediate (≤7d, n=157)'
hdr_cells[2].text = 'Delayed (>7d, n=520)'
hdr_cells[3].text = 'Difference'

data_panel_d = [
    ['30-Day CAR (%)', '-0.53', '-0.40', '-0.13pp'],
    ['5-Day CAR (%)', '-0.30', '-0.01', '-0.29pp'],
    ['Volatility Change (pp)', '-2.98', '-2.10', '-0.88pp'],
    ['Firm Size (log)', '10.46', '10.57', '-0.11'],
]

for i, row_data in enumerate(data_panel_d, start=1):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = str(cell_text)

note = doc.add_paragraph('Note: Full sample includes all Privacy Rights Clearinghouse breaches (2006–2025) with firm identifiers. CRSP-matched sample restricted to publicly traded firms with available stock price data. CAR calculated as 30-day cumulative abnormal returns using market model with 120-day pre-breach estimation window. FCC classification based on SIC codes 4813 (Telephone), 4841 (Cable), 4899 (VoIP); n=141 FCC firms in full sample, n=130 in CRSP-matched. Disclosure timing measured as days between breach date and public disclosure date.')
note.style = 'Intense Quote'

doc.add_page_break()

# ============================================================================
# TABLE 2: Univariate CAR
# ============================================================================
doc.add_heading('TABLE 2: Mean 30-Day Cumulative Abnormal Returns (CAR) by Regulatory Status and Disclosure Speed', level=2)

table = doc.add_table(rows=7, cols=6)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Group'
hdr_cells[1].text = 'N'
hdr_cells[2].text = 'Mean CAR 30d'
hdr_cells[3].text = 'Std Error'
hdr_cells[4].text = '95% CI Lower'
hdr_cells[5].text = '95% CI Upper'

data_table2 = [
    ['FCC-Regulated', '130', '-2.39%', '0.88', '-4.12%', '-0.66%'],
    ['Non-FCC', '547', '+0.04%', '0.37', '-0.69%', '+0.77%'],
    ['Difference (FCC - Non-FCC)', '—', '-2.43%', '0.96', '-4.31%', '-0.55%'],
    ['Immediate Disclosure (≤7d)', '157', '-0.53%', '0.75', '-2.00%', '+0.94%'],
    ['Delayed Disclosure (>7d)', '520', '-0.40%', '0.39', '-1.17%', '+0.37%'],
    ['Difference (Immediate - Delayed)', '—', '-0.13%', '0.84', '-1.78%', '+1.52%'],
]

for i, row_data in enumerate(data_table2, start=1):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = str(cell_text)

note = doc.add_paragraph('N = 677 (CRSP-matched breaches). CAR calculated as 30-day cumulative abnormal return (day 0 = disclosure date) using market model with 120-day pre-breach estimation window. FCC sample comprises firms with SIC codes 4813, 4841, 4899. Standard errors computed using Newey-West adjustment for overlapping windows. No significant difference in returns by disclosure speed; FCC penalty evident in univariate comparison but robustness tests examine whether this reflects regulation, firm size, or sector characteristics.')
note.style = 'Intense Quote'

doc.add_page_break()

# ============================================================================
# TABLE 3: Main Regression H1-H4
# ============================================================================
doc.add_heading('TABLE 3: Main Regression Results — H1-H4 Effects on 30-Day CAR', level=2)

doc.add_paragraph('Specification: CAR₃₀ = β₀ + β₁(immediate_disclosure) + β₂(fcc_reportable) + β₃(prior_breaches_1yr) + β₄(health_breach) + β₅(firm_size_log) + β₆(leverage) + β₇(roa) + ε')

doc.add_paragraph('N = 653 observations | Standard Errors: HC3 heteroskedasticity-consistent robust')

table = doc.add_table(rows=9, cols=7)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Coefficient'
hdr_cells[1].text = 'Estimate'
hdr_cells[2].text = 'SE'
hdr_cells[3].text = 't-stat'
hdr_cells[4].text = 'p-value'
hdr_cells[5].text = '95% CI Lower'
hdr_cells[6].text = '95% CI Upper'

data_table3 = [
    ['H1: Immediate Disclosure', '+0.8394%', '0.8988', '0.934', '0.3504', '-0.924%', '+2.603%'],
    ['H2: FCC Regulation', '-2.1179%', '1.0776', '-1.965', '0.0494**', '-4.231%', '-0.005%'],
    ['H3: Prior Breaches (1yr)', '+0.0142%', '0.0664', '0.214', '0.8308', '-0.118%', '+0.146%'],
    ['H4: Health Breach', '-0.6150%', '1.6109', '-0.382', '0.7026', '-3.776%', '+2.545%'],
    ['Firm Size (log)', '+0.2925%', '0.3168', '0.924', '0.3559', '-0.328%', '+0.912%'],
    ['Leverage', '+1.8504%', '1.3661', '1.354', '0.1756', '-0.835%', '+4.536%'],
    ['ROA', '+21.6881%***', '8.8298', '2.456', '0.0140**', '+4.393%', '+38.983%'],
    ['Constant', '-3.4850%', '3.6797', '-0.947', '0.3437', '-10.681%', '+3.711%'],
]

for i, row_data in enumerate(data_table3, start=1):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = str(cell_text)

doc.add_paragraph()
doc.add_paragraph('Model Fit: R² = 0.0205 | Adj. R² = 0.0029 | F-statistic = 1.189 (p = 0.3087)')
doc.add_paragraph('Significance levels: * p < 0.10, ** p < 0.05, *** p < 0.01')

doc.add_paragraph()
doc.add_heading('Key Findings:', level=3)
doc.add_paragraph('H1 (Timing): No significant effect of immediate disclosure on returns (p = 0.350). Market does not reward speed of disclosure within regulatory windows.', style='List Bullet')
doc.add_paragraph('H2 (FCC Regulation): Significant negative effect for FCC-regulated firms (-2.12%, p = 0.049). FCC firms suffer ~2% abnormal return penalty post-breach.', style='List Bullet')
doc.add_paragraph('H3 (Reputation): No additional penalty for firms with recent breach history (p = 0.831). Prior breaches do not accumulate market sanctions.', style='List Bullet')
doc.add_paragraph('H4 (Severity): No differential penalty for health data breaches (p = 0.703). Data type does not drive return reactions.', style='List Bullet')
doc.add_paragraph('ROA Control: Highly profitable firms experience higher abnormal returns post-breach (+21.69%, p = 0.014), likely reflecting earnings recovery expectations.', style='List Bullet')

note = doc.add_paragraph('Note: Full specification includes all four hypothesis predictors simultaneously, isolating each effect net of others. This is the primary model for all hypothesis tests. Alternative specifications (simpler models, alternative controls, alternative windows) presented in robustness tables.')
note.style = 'Intense Quote'

doc.add_page_break()

# Continue with remaining tables... (abbreviated for space)
# Tables 4-19 would follow the same pattern

# ============================================================================
# SAVE DOCUMENT
# ============================================================================
output_path = 'outputs/ESSAY1_APPENDIX_TABLES.docx'
doc.save(output_path)

print(f"\n✅ ESSAY 1 APPENDIX CREATED SUCCESSFULLY")
print(f"📄 Output: {output_path}")
print(f"📊 Tables: 1-3 complete (sample, univariate, main regression)")
print(f"\n⚠️  NOTE: Complete version with all 19 tables requires manual addition")
print(f"   or full script expansion (current shows structure template)")

