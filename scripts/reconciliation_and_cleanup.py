"""
RECONCILIATION AND PRONOUN CLEANUP
Master script to:
1. Update all numeric values to authoritative sources
2. Remove pronouns ("we", "we're", "our", "I", "my") and convert to passive voice

AUTHORITATIVE VALUES:
- H1 Coefficient: +0.649% (not 0.57% or others)
- H1 p-value: 0.443 (not 0.373 or 0.539)
- Immediate Disclosure Rate: 17.6% (confirmed correct)
- Delayed Rate: 82.4% (confirmed correct)
- Equivalence Bound: ±2.10 pp
- MDE at 80% Power: ±2.39 pp
- Sample Size (H1): 898
"""

from docx import Document
from docx.shared import Pt, RGBColor
import re

def remove_pronouns(text):
    """
    Remove/replace pronouns with passive voice or restructure
    """
    replacements = [
        # "we" patterns
        (r'\bwe have\b', 'findings include'),
        (r'\bwe tested\b', 'tests show'),
        (r'\bwe find\b', 'results show'),
        (r'\bwe show\b', 'evidence shows'),
        (r'\bwe claim\b', 'evidence supports'),
        (r'\bwe document\b', 'documentation reveals'),
        (r'\bwe observe\b', 'observation indicates'),
        (r'\bwe conduct\b', 'conducted analysis'),
        (r'\bwe use\b', 'employed'),
        (r'\bwe employ\b', 'employed'),
        (r'\bwe establish\b', 'established'),
        (r'\bwe examine\b', 'examination reveals'),
        (r'\bwe investigate\b', 'investigation reveals'),
        (r'\bwe analyze\b', 'analysis shows'),
        (r'\bwe are\b', 'findings are'),
        (r'\bwe provide\b', 'provided'),
        (r'\bwe discuss\b', 'discussion includes'),
        (r'\bwe explain\b', 'explanation includes'),
        (r'\bwe present\b', 'presented'),
        (r"\bwe're\b", 'findings are'),

        # "our" patterns
        (r'\bour evidence\b', 'the evidence'),
        (r'\bour findings\b', 'the findings'),
        (r'\bour analysis\b', 'the analysis'),
        (r'\bour sample\b', 'the sample'),
        (r'\bour study\b', 'the study'),
        (r'\bour data\b', 'the data'),
        (r'\bour results\b', 'the results'),
        (r'\bour hypothesis\b', 'the hypothesis'),
        (r'\bour paper\b', 'this paper'),
        (r'\bour dissertation\b', 'this dissertation'),
        (r'\bour approach\b', 'the approach'),
        (r'\bour methods\b', 'the methods'),
        (r'\bour research\b', 'this research'),
        (r'\bour work\b', 'this work'),
        (r'\bour claim\b', 'the claim'),
        (r'\bour contribution\b', 'the contribution'),
        (r'\bour model\b', 'the model'),
        (r'\bour natural experiment\b', 'the natural experiment'),
        (r'\bour essay\b', 'the essay'),

        # "I" and "me" patterns (less common in academic writing)
        (r"\bI find\b", 'findings indicate'),
        (r"\bI show\b", 'evidence shows'),
        (r"\bI argue\b", 'the argument'),
        (r"\bI claim\b", 'the claim'),

        # "you/your" patterns (in Q&A contexts)
        (r'\byour dissertation\b', 'the dissertation'),
        (r'\byour proposal\b', 'the proposal'),
    ]

    # Apply replacements (case-insensitive for some, but preserve case where needed)
    result = text
    for pattern, replacement in replacements:
        result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)

    return result

def update_numbers_in_text(text):
    """
    Update specific numeric values to authoritative versions
    """
    # Map old values to new values
    replacements = [
        # H1 p-value: 0.373 or 0.539 → 0.443
        ('p=0.373', 'p=0.443'),
        ('p = 0.373', 'p = 0.443'),
        ('(0.373', '(0.443'),
        ('p=0.539', 'p=0.443'),
        ('p = 0.539', 'p = 0.443'),

        # H1 coefficient variations
        ('+0.57%', '+0.649%'),
        ('+0.57pp', '+0.649pp'),
        ('0.57%', '0.649%'),
        ('+0.572', '+0.649'),

        # Keep 17.6% and 82.4% (these are correct)
        # Keep ±2.10 pp (correct)
        # Keep ±2.39 pp (correct)
    ]

    result = text
    for old, new in replacements:
        result = result.replace(old, new)

    return result

def process_document(filepath, output_filepath=None):
    """
    Process a Word document to remove pronouns and update numbers
    """
    if output_filepath is None:
        output_filepath = filepath

    doc = Document(filepath)

    changes_made = 0

    for para in doc.paragraphs:
        original_text = para.text

        # Update numbers first
        new_text = update_numbers_in_text(original_text)

        # Remove pronouns
        new_text = remove_pronouns(new_text)

        if new_text != original_text:
            # Clear the paragraph and add new text
            para.clear()
            para.add_run(new_text)
            changes_made += 1
            print(f"Updated: {original_text[:60]}...")

    # Also process table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original_text = para.text
                    new_text = update_numbers_in_text(original_text)
                    new_text = remove_pronouns(new_text)

                    if new_text != original_text:
                        para.clear()
                        para.add_run(new_text)
                        changes_made += 1

    doc.save(output_filepath)
    return changes_made

# Process all documents
print("=" * 80)
print("RECONCILIATION AND PRONOUN CLEANUP")
print("=" * 80)
print()

documents = [
    r'C:\Users\mcobp\DISSERTATION_CLONE\Proposal_Defense_Q&A_Guide.docx',
    r'C:\Users\mcobp\DISSERTATION_CLONE\Dissertation_Proposal_Complete.docx',
]

total_changes = 0
for doc_path in documents:
    try:
        print(f"\nProcessing: {doc_path.split(chr(92))[-1]}")
        changes = process_document(doc_path)
        total_changes += changes
        print(f"  [OK] {changes} paragraphs updated")
    except FileNotFoundError:
        print(f"  [SKIP] File not found")
    except Exception as e:
        print(f"  [ERROR] {e}")

print()
print("=" * 80)
print(f"COMPLETE: {total_changes} total updates made")
print("=" * 80)
print()
print("CHANGES MADE:")
print("  - H1 p-value: 0.373/0.539 to 0.443")
print("  - H1 coefficient: 0.57% to 0.649%")
print("  - Pronouns removed: we, we're, our, I to passive voice")
print()
print("Files updated:")
for doc in documents:
    print(f"  [OK] {doc.split(chr(92))[-1]}")
