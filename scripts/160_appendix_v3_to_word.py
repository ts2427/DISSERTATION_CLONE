"""
REBUILD — APPENDIX V3 WORD RENDERING (16-table citation order)
Renders outputs/rebuild/appendix_v3/table_{1..16}.csv into a single Word
document for the dissertation appendix. Pure rendering: no numbers computed
here; the CSVs (asserted by Stage 8) are the source of truth. Each CSV's
trailing CAPTION row supplies the table note verbatim, so the document
cannot drift from the asserted captions; EXTRA holds interpretive addenda
only. The sample header is built from constants_v3.json.
Outputs: outputs/rebuild/APPENDIX_V3_TABLES.docx
         outputs/rebuild/INTEXT_TABLES_4_5.docx
"""
import sys
import json
import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Pt

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

TITLES = {
    1: 'Summary Statistics (CRSP Sample)',
    2: 'Disclosure-Date Verification (First 8-K Within 90 Days)',
    3: 'Mean and Median 30-Day CAR by Subgroup, with Welch Difference Tests',
    4: 'Main Regression: 30-Day CAR on Hypothesis Variables and Controls (HC3)',
    5: 'Equivalence Testing (TOST, Pre-Specified ±2.10pp Bound)',
    6: 'Timing Effect by Regulatory Regime (Subsamples and Interaction)',
    7: 'Sample Restrictions (Timing, FCC, and ROA Coefficients)',
    8: 'Standard-Error Methods',
    9: 'Timing Coefficient by Firm-Size Quartile',
    10: 'FCC Coefficient by Firm-Size Quartile (Treated Counts Shown)',
    11: 'Specification Robustness (Year Fixed Effects)',
    12: 'Alternative Explanations (Breach Severity)',
    13: 'Factor-Model Controls (FF3)',
    14: 'Abnormal Log Turnover (Event [-5,+25] vs Estimation [-240,-60])',
    15: 'Random-Forest Feature Importance',
    16: 'Pre-Announcement Abnormal Returns (Leakage Windows)',
}

# Interpretive addenda appended after the asserted caption; keep sparse.
EXTRA = {
    5: 'Verdicts: H3 bounded null; H1/H2/H4 null-inconclusive.',
    8: 'The firm-clustered row rests on 11 treated parent-CIK clusters; HC3 is primary.',
    10: 'Q1 holds 2 treated events — quartile estimates are illustrative, not inferential.',
    11: 'Treatment is near-collinear with the telecom sector, so within-industry '
        'specifications are not estimable with useful precision.',
    15: 'All four hypothesis variables rank below all three firm controls.',
    16: 'NOTIFICATION-ANCHOR CAVEAT: PRC breach dates are notification-anchored, so these '
        'pre-announcement windows are measured against the notification clock; with a '
        'median disclosure delay of 24 days, some events\' first market disclosure '
        'plausibly falls inside the [-20,-11] window. The significant positive drift there '
        'is therefore reported, not explained away; the supportable claim is the absence '
        'of anticipatory NEGATIVE leakage in all three windows.',
}

C = json.loads(Path('outputs/rebuild/constants_v3.json').read_text())
SAMPLE_HEADER = (
    'Sample definitions (name-your-sample rule): '
    f'FULL = {C["events_total"]} verified breach events ({C["treated_total"]} treated); '
    f'CRSP = {C["events_crsp"]} events with market data ({C["treated_crsp"]} treated); '
    f'REGRESSION = {C["N_regression"]} events with complete covariates '
    f'({C["treated_regression"]} treated events, {C["treated_parent_ciks_regression"]} '
    f'parent CIKs, {C["treated_orgs_regression"]} treated organizations). '
    f'Immediate-disclosure share: {100 * C["immediate_share_full"]:.1f}% full / '
    f'{100 * C["immediate_share_crsp"]:.1f}% CRSP / '
    f'{100 * C["immediate_share_regression"]:.1f}% regression.')


def load_table(i):
    """Return (dataframe without the CAPTION row, caption text)."""
    df = pd.read_csv(f'outputs/rebuild/appendix_v3/table_{i}.csv').fillna('')
    caption = ''
    last = df.iloc[-1]
    if str(last.iloc[0]) == 'CAPTION':
        caption = str(last.iloc[1])
        df = df.iloc[:-1]
    return df, caption


def render(doc, i, title, body_pt):
    df, caption = load_table(i)
    doc.add_heading(title, level=2)
    note = f'Note. {caption}'
    if i in EXTRA:
        note += ' ' + EXTRA[i]
    cap = doc.add_paragraph(note)
    cap.runs[0].font.size = Pt(8)
    cap.runs[0].font.italic = True
    t = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    t.style = 'Light Grid Accent 1'
    for j, c in enumerate(df.columns):
        t.rows[0].cells[j].text = str(c)
    for r, (_, row) in enumerate(df.iterrows(), start=1):
        for j, v in enumerate(row):
            t.rows[r].cells[j].text = str(v)
    for row in t.rows:
        for cell in row.cells:
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(body_pt)


doc = Document()
doc.add_heading('Essay 1 Appendix — Canonical V3 Tables', level=1)
p = doc.add_paragraph('All values regenerate from run_all.py (Canonical V3 chain, scripts '
                      '150–159) and assert against outputs/rebuild/constants_v3.json. '
                      'Table notes are the asserted captions emitted by Stage 8. '
                      'Generated by scripts/160.')
p.runs[0].font.size = Pt(9)
p2 = doc.add_paragraph(SAMPLE_HEADER)
p2.runs[0].font.size = Pt(9)
for i in range(1, 17):
    render(doc, i, f'Table {i}. {TITLES[i]}', body_pt=8)
out = Path('outputs/rebuild/APPENDIX_V3_TABLES.docx')
doc.save(out)
print(f'Saved: {out} (16 tables)')

# In-text renderings: main regression (Table 4) and TOST summary (Table 5)
doc2 = Document()
doc2.add_heading('In-Text Tables (Results section) — Canonical V3', level=1)
for i, intext_title in [(4, 'Main Regression Results'), (5, 'Equivalence Testing Summary')]:
    render(doc2, i, f'Table {i} (in-text). {intext_title}', body_pt=9)
out2 = Path('outputs/rebuild/INTEXT_TABLES_4_5.docx')
doc2.save(out2)
print(f'Saved: {out2}')
