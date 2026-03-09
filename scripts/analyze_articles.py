import os
import pdfplumber
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Configuration
ARTICLES_DIR = r"C:\Users\mcobp\DISSERTATION_CLONE\data\Articles"
OUTPUT_FILE = r"C:\Users\mcobp\DISSERTATION_CLONE\outputs\Article_Review_and_Dissertation_Relevance.docx"

# Dissertation context keywords for relevance assessment
CORE_KEYWORDS = [
    "FCC", "data breach", "cybersecurity breach", "security breach",
    "stock market reaction", "market reaction", "breach announcement",
    "disclosure", "notification law", "information asymmetry",
    "event study", "shareholder"
]

HIGH_KEYWORDS = [
    "cybersecurity", "cyber attack", "information disclosure",
    "regulatory", "governance", "crisis communication",
    "firm reputation", "risk management", "stock price"
]

MEDIUM_KEYWORDS = [
    "methodology", "statistical", "analysis", "model",
    "firm", "corporate", "management", "strategy"
]

def extract_pdf_content(pdf_path):
    """Extract full text and metadata from PDF"""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            metadata = pdf.metadata
            text = ""
            page_count = len(pdf.pages)

            for page in pdf.pages:
                text += page.extract_text() or ""

            return {
                'text': text,
                'metadata': metadata,
                'page_count': page_count,
                'title': metadata.get('Title', ''),
                'author': metadata.get('Author', ''),
                'success': True
            }
    except Exception as e:
        return {
            'text': '',
            'metadata': {},
            'page_count': 0,
            'title': '',
            'author': '',
            'success': False,
            'error': str(e)
        }

def get_file_title(filename):
    """Extract clean title from filename"""
    return filename.replace('.pdf', '').strip()

def assess_relevance(text, filename):
    """Assess article relevance based on keywords"""
    text_lower = (text + filename).lower()

    core_count = sum(1 for kw in CORE_KEYWORDS if kw.lower() in text_lower)
    high_count = sum(1 for kw in HIGH_KEYWORDS if kw.lower() in text_lower)
    medium_count = sum(1 for kw in MEDIUM_KEYWORDS if kw.lower() in text_lower)

    if core_count >= 3:
        return "CORE"
    elif core_count >= 1 or high_count >= 3:
        return "HIGH"
    elif high_count >= 1 or medium_count >= 3:
        return "MEDIUM"
    else:
        return "LOW"

def extract_abstract_and_intro(text):
    """Extract abstract and introduction sections"""
    lines = text.split('\n')
    abstract = ""
    intro = ""

    in_abstract = False
    in_intro = False
    line_count = 0

    for line in lines:
        line_lower = line.lower().strip()

        if 'abstract' in line_lower and not in_abstract:
            in_abstract = True
            continue

        if in_abstract:
            if any(x in line_lower for x in ['introduction', 'background', '1.', '1 ']):
                in_abstract = False
                in_intro = True
            else:
                abstract += line + "\n"
                if len(abstract) > 1000:
                    break

        if in_intro:
            intro += line + "\n"
            if len(intro) > 1500:
                break

    return abstract.strip()[:500], intro.strip()[:800]

def categorize_articles(articles_dir):
    """Categorize all articles"""
    articles = []

    # Get all PDF files
    pdf_files = sorted([f for f in os.listdir(articles_dir) if f.endswith('.pdf')])

    print(f"Found {len(pdf_files)} PDF files. Processing...")

    for idx, filename in enumerate(pdf_files, 1):
        print(f"  [{idx}/{len(pdf_files)}] Processing: {filename[:60]}...")

        filepath = os.path.join(articles_dir, filename)
        content = extract_pdf_content(filepath)

        if content['success']:
            abstract, intro = extract_abstract_and_intro(content['text'])
            relevance = assess_relevance(content['text'], filename)

            articles.append({
                'filename': filename,
                'title': get_file_title(filename),
                'metadata_title': content['title'],
                'author': content['author'],
                'pages': content['page_count'],
                'text': content['text'],
                'abstract': abstract,
                'relevance': relevance,
                'keywords': extract_keywords(content['text'])
            })
        else:
            articles.append({
                'filename': filename,
                'title': get_file_title(filename),
                'metadata_title': '',
                'author': '',
                'pages': 0,
                'text': '',
                'abstract': f"[Error reading PDF: {content['error']}]",
                'relevance': 'UNKNOWN',
                'keywords': []
            })

    return articles

def extract_keywords(text):
    """Extract key findings from text"""
    keywords = []
    text_lower = text.lower()

    # Look for key topics
    topics = {
        'market reaction': 'stock market response analyzed',
        'event study': 'event study methodology used',
        'disclosure': 'disclosure/transparency focus',
        'information asymmetry': 'information asymmetry examined',
        'FCC regulation': 'FCC regulatory aspect',
        'data breach': 'data breach analysis',
        'stock price': 'stock price impact studied',
        'cybersecurity': 'cybersecurity focus',
    }

    for keyword, label in topics.items():
        if keyword in text_lower:
            keywords.append(label)

    return keywords[:5]

def create_word_document(articles):
    """Create comprehensive Word document"""
    doc = Document()

    # Title
    title = doc.add_heading('Comprehensive Article Review', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('How 63 Journal Articles Relate to Your Dissertation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.italic = True
    subtitle_format.font.size = Pt(12)

    # Executive Summary
    doc.add_heading('Executive Summary', 1)

    relevance_counts = {}
    for article in articles:
        rel = article['relevance']
        relevance_counts[rel] = relevance_counts.get(rel, 0) + 1

    summary_text = f"""
This comprehensive review analyzes {len(articles)} journal articles collected for your dissertation research on FCC regulations, cybersecurity breaches, information asymmetry, and market reactions.

**Article Distribution by Relevance:**
"""
    doc.add_paragraph(summary_text)

    for rel in ['CORE', 'HIGH', 'MEDIUM', 'LOW', 'UNKNOWN']:
        count = relevance_counts.get(rel, 0)
        pct = (count / len(articles)) * 100 if len(articles) > 0 else 0
        doc.add_paragraph(f"• {rel}: {count} articles ({pct:.1f}%)", style='List Bullet')

    doc.add_paragraph("\nEach article below is annotated with its relevance category, page count, and how it supports your dissertation research.")

    # Table of Contents by Category
    doc.add_heading('Articles by Relevance Category', 1)

    for relevance in ['CORE', 'HIGH', 'MEDIUM', 'LOW']:
        articles_in_cat = [a for a in articles if a['relevance'] == relevance]

        if articles_in_cat:
            doc.add_heading(f'{relevance} Relevance ({len(articles_in_cat)} articles)', 2)

            for article in articles_in_cat:
                # Article heading with metadata
                article_heading = doc.add_heading(article['title'], 3)

                # Metadata table
                metadata_table = doc.add_table(rows=1, cols=2)
                metadata_table.autofit = False
                metadata_table.allow_autofit = False

                cells = metadata_table.rows[0].cells
                cells[0].text = "Author:"
                cells[1].text = article['author'] if article['author'] else "[Not specified in metadata]"

                row = metadata_table.add_row().cells
                row[0].text = "Pages:"
                row[1].text = str(article['pages'])

                row = metadata_table.add_row().cells
                row[0].text = "Relevance:"
                row[1].text = article['relevance']

                # Abstract
                if article['abstract'] and article['abstract'] != "[Error reading PDF]":
                    doc.add_paragraph("Abstract:", style='Heading 4')
                    abstract_para = doc.add_paragraph(article['abstract'])
                    abstract_para_format = abstract_para.runs[0] if abstract_para.runs else None
                    if abstract_para_format:
                        abstract_para_format.font.size = Pt(10)
                        abstract_para_format.italic = True

                # Keywords/Topics
                if article['keywords']:
                    doc.add_paragraph("Key Topics:", style='Heading 4')
                    for kw in article['keywords']:
                        doc.add_paragraph(kw, style='List Bullet')

                # Relevance explanation
                doc.add_paragraph("Relevance to Your Dissertation:", style='Heading 4')
                relevance_explanation = get_relevance_explanation(article, relevance)
                doc.add_paragraph(relevance_explanation)

                doc.add_paragraph()  # Spacing

    # Final Recommendations
    doc.add_heading('Recommendations', 1)

    core_articles = [a for a in articles if a['relevance'] == 'CORE']
    high_articles = [a for a in articles if a['relevance'] == 'HIGH']

    doc.add_heading('Must-Cite Articles (CORE)', 2)
    doc.add_paragraph(f"You should definitely cite the {len(core_articles)} CORE relevance articles. They directly support your dissertation hypotheses and frameworks.")

    for article in core_articles:
        doc.add_paragraph(f"• {article['title']}", style='List Bullet')

    doc.add_heading('Important Background (HIGH)', 2)
    doc.add_paragraph(f"The {len(high_articles)} HIGH relevance articles provide important context and methodology. Consider citing the most relevant ones.")

    # Save document
    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"\nDocument saved to: {OUTPUT_FILE}")

def get_relevance_explanation(article, relevance):
    """Get explanation of why article is relevant"""
    title = article['title'].lower()
    text = article['text'].lower()

    explanations = {
        'CORE': 'This article directly addresses key themes in your dissertation: data breaches, market reactions, disclosure regulations, or information asymmetry. Its findings or methodology are likely crucial for your arguments.',
        'HIGH': 'This article provides important supporting evidence or methodology for understanding the relationship between cybersecurity breaches, regulatory frameworks, and market dynamics relevant to your research.',
        'MEDIUM': 'This article contributes supporting context, methodology, or theoretical frameworks that help build the foundation for understanding your dissertation topic.',
        'LOW': 'This article provides supplementary background or tangential context that may be useful for contextualizing your research but is not central to your main arguments.',
        'UNKNOWN': 'Unable to determine relevance from PDF content. Manual review recommended.'
    }

    return explanations.get(relevance, '')

# Main execution
if __name__ == "__main__":
    print("Starting comprehensive article analysis...")
    print(f"Articles directory: {ARTICLES_DIR}\n")

    articles = categorize_articles(ARTICLES_DIR)

    print(f"\nAnalyzed {len(articles)} articles")

    print("Creating Word document...")
    create_word_document(articles)

    print("\n" + "="*60)
    print("ANALYSIS COMPLETE")
    print("="*60)
    print(f"Total articles analyzed: {len(articles)}")
    print(f"Output file: {OUTPUT_FILE}")
