"""
Create Word document with regression formulas and specifications for all essays.
Outputs: Dissertation_Regression_Formulas.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement

def add_formula_section(doc, essay_num, title, formulas_dict):
    """Add an essay section with formulas"""
    # Essay heading
    heading = doc.add_heading(f'Essay {essay_num}: {title}', level=1)

    for spec_name, spec_details in formulas_dict.items():
        # Specification subheading
        doc.add_heading(spec_name, level=2)

        # Formula
        doc.add_paragraph('Specification:', style='Normal')
        formula_para = doc.add_paragraph()
        formula_run = formula_para.add_run(spec_details['formula'])
        formula_run.italic = True
        formula_run.font.name = 'Courier New'
        formula_run.font.size = Pt(10)

        # Sample size
        doc.add_paragraph(f"Sample Size: {spec_details['n']}")

        # Estimation method
        doc.add_paragraph(f"Estimation Method: {spec_details['method']}")

        # Variable descriptions
        doc.add_paragraph('Variable Definitions:', style='Normal')
        vars_table = doc.add_table(rows=len(spec_details['variables']) + 1, cols=3)
        vars_table.style = 'Light Grid Accent 1'

        # Header
        header_cells = vars_table.rows[0].cells
        header_cells[0].text = 'Variable Name'
        header_cells[1].text = 'Type'
        header_cells[2].text = 'Description'
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Variable rows
        for row_idx, (var_name, var_info) in enumerate(spec_details['variables'].items(), 1):
            row = vars_table.rows[row_idx]
            row.cells[0].text = var_name
            row.cells[1].text = var_info['type']
            row.cells[2].text = var_info['desc']

        doc.add_paragraph()

# Create document
doc = Document()

# Title page
title = doc.add_paragraph()
title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("REGRESSION SPECIFICATIONS\n& FORMULAS")
title_run.bold = True
title_run.font.size = Pt(18)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("Complete Dissertation Analysis\nEssays 1-3 with Variable Definitions")
subtitle_run.font.size = Pt(12)

doc.add_paragraph()

note = doc.add_paragraph()
note.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
note_run = note.add_run("All models use heteroscedasticity-robust standard errors (HC3)")
note_run.italic = True
note_run.font.size = Pt(10)

doc.add_page_break()

# ============================================================================
# ESSAY 1: MARKET RETURNS
# ============================================================================
essay1_formulas = {
    'Model 1: FCC Effect Only': {
        'formula': 'CAR_30d = β₀ + β₁·FCC_Reportable + ε',
        'n': '898',
        'method': 'OLS with HC3 robust standard errors',
        'variables': {
            'CAR_30d': {
                'type': 'Continuous',
                'desc': 'Cumulative abnormal return over 30-day event window [-5, +25 days], estimated using market model with Fama-French 3-factor adjustment'
            },
            'FCC_Reportable': {
                'type': 'Binary (0/1)',
                'desc': 'Indicator for firms subject to FCC Rule 37.3 (47 CFR § 64.2011). 1=telecommunications carriers (SIC 4813, 4841, 4899), 0=all other industries'
            }
        }
    },
    'Model 2: FCC + Breach Characteristics': {
        'formula': 'CAR_30d = β₀ + β₁·FCC_Reportable + β₂·Health_Breach + β₃·Financial_Breach + β₄·Prior_Breaches + ε',
        'n': '898',
        'method': 'OLS with HC3 robust standard errors',
        'variables': {
            'CAR_30d': {
                'type': 'Continuous',
                'desc': 'Cumulative abnormal return, 30-day window'
            },
            'FCC_Reportable': {
                'type': 'Binary (0/1)',
                'desc': 'FCC regulation indicator'
            },
            'Health_Breach': {
                'type': 'Binary (0/1)',
                'desc': 'Indicator = 1 if breach involves protected health information (PHI), 0 otherwise'
            },
            'Financial_Breach': {
                'type': 'Binary (0/1)',
                'desc': 'Indicator = 1 if breach involves financial/payment card data, 0 otherwise'
            },
            'Prior_Breaches': {
                'type': 'Count',
                'desc': 'Total number of prior breaches for firm in dataset (reputation history)'
            }
        }
    },
    'Model 3: Full Specification with Firm Controls': {
        'formula': 'CAR_30d = β₀ + β₁·FCC_Reportable + β₂·Health_Breach + β₃·Financial_Breach + β₄·Prior_Breaches + β₅·Firm_Size_log + β₆·Leverage + β₇·ROA + ε',
        'n': '898',
        'method': 'OLS with HC3 robust standard errors',
        'variables': {
            'CAR_30d': {
                'type': 'Continuous',
                'desc': 'Cumulative abnormal return, 30-day window'
            },
            'FCC_Reportable': {
                'type': 'Binary (0/1)',
                'desc': 'FCC regulation indicator'
            },
            'Health_Breach': {
                'type': 'Binary (0/1)',
                'desc': 'Breach involves protected health information'
            },
            'Financial_Breach': {
                'type': 'Binary (0/1)',
                'desc': 'Breach involves financial/payment card data'
            },
            'Prior_Breaches': {
                'type': 'Count',
                'desc': 'Total prior breaches (reputation)'
            },
            'Firm_Size_log': {
                'type': 'Continuous',
                'desc': 'Natural logarithm of total assets (Compustat), measured in millions USD'
            },
            'Leverage': {
                'type': 'Continuous',
                'desc': 'Total debt / Total assets ratio (Compustat). Range: [0, 1]'
            },
            'ROA': {
                'type': 'Continuous',
                'desc': 'Return on Assets = Net Income / Total Assets (Compustat). Range: typically [-1, 1]'
            }
        }
    },
    'Model 4: Adding Disclosure Timing': {
        'formula': 'CAR_30d = β₀ + β₁·FCC_Reportable + β₂·Immediate_Disclosure + β₃·Health_Breach + β₄·Financial_Breach + β₅·Prior_Breaches + β₆·Firm_Size_log + β₇·Leverage + β₈·ROA + ε',
        'n': '898',
        'method': 'OLS with HC3 robust standard errors',
        'variables': {
            'CAR_30d': {
                'type': 'Continuous',
                'desc': 'Cumulative abnormal return, 30-day window'
            },
            'FCC_Reportable': {
                'type': 'Binary (0/1)',
                'desc': 'FCC regulation indicator'
            },
            'Immediate_Disclosure': {
                'type': 'Binary (0/1)',
                'desc': 'Indicator = 1 if breach disclosed within 7 days of detection, 0 otherwise'
            },
            'Health_Breach': {
                'type': 'Binary (0/1)',
                'desc': 'Breach involves protected health information'
            },
            'Financial_Breach': {
                'type': 'Binary (0/1)',
                'desc': 'Breach involves financial/payment card data'
            },
            'Prior_Breaches': {
                'type': 'Count',
                'desc': 'Total prior breaches (reputation)'
            },
            'Firm_Size_log': {
                'type': 'Continuous',
                'desc': 'Natural log of total assets'
            },
            'Leverage': {
                'type': 'Continuous',
                'desc': 'Total debt / Total assets'
            },
            'ROA': {
                'type': 'Continuous',
                'desc': 'Return on assets'
            }
        }
    }
}

add_formula_section(doc, 1, 'Market Reactions (CAR-30d)', essay1_formulas)
doc.add_page_break()

# ============================================================================
# ESSAY 2: VOLATILITY (MAIN MODELS)
# ============================================================================
essay2_formulas = {
    'Main Model: FCC Effect on Volatility': {
        'formula': 'ΔVolatility = β₀ + β₁·FCC_Reportable + β₂·Disclosure_Delay + β₃·Pre_Volatility + β₄·Firm_Size_log + β₅·Health_Breach + β₆·Leverage + β₇·ROA + ε',
        'n': '891',
        'method': 'OLS with HC3 robust standard errors',
        'variables': {
            'ΔVolatility': {
                'type': 'Continuous',
                'desc': 'Change in return volatility = σ_post - σ_pre, where σ calculated over 20-trading-day windows (pre: -25 to -5 days; post: +5 to +25 days)'
            },
            'FCC_Reportable': {
                'type': 'Binary (0/1)',
                'desc': 'FCC regulation indicator (telecommunications carriers)'
            },
            'Disclosure_Delay': {
                'type': 'Continuous',
                'desc': 'Days from breach detection to public disclosure (disclosure_delay_days)'
            },
            'Pre_Volatility': {
                'type': 'Continuous',
                'desc': 'Pre-breach return volatility (σ_pre), baseline information environment'
            },
            'Firm_Size_log': {
                'type': 'Continuous',
                'desc': 'Natural log of total assets'
            },
            'Health_Breach': {
                'type': 'Binary (0/1)',
                'desc': 'Breach involves protected health information'
            },
            'Leverage': {
                'type': 'Continuous',
                'desc': 'Total debt / Total assets'
            },
            'ROA': {
                'type': 'Continuous',
                'desc': 'Return on assets'
            }
        }
    }
}

add_formula_section(doc, 2, 'Information Asymmetry (Volatility)', essay2_formulas)

# Essay 2 Mechanisms
doc.add_heading('Essay 2 (Continued): Heterogeneous Mechanisms', level=1)

mech_formulas = {
    'Mechanism 1: Firm Size Heterogeneity': {
        'formula': 'ΔVolatility = β₀ + β₁·FCC + β₂·Firm_Size_Q1 + β₃·Firm_Size_Q2 + β₄·Firm_Size_Q3 + β₅·FCC×Firm_Size_Q1 + β₆·FCC×Firm_Size_Q2 + β₇·FCC×Firm_Size_Q3 + [controls] + ε',
        'n': '891',
        'method': 'OLS with HC3 robust standard errors',
        'variables': {
            'ΔVolatility': {
                'type': 'Continuous',
                'desc': 'Change in return volatility'
            },
            'FCC': {
                'type': 'Binary (0/1)',
                'desc': 'FCC indicator'
            },
            'Firm_Size_Q1, Q2, Q3': {
                'type': 'Binary (0/1)',
                'desc': 'Firm size quartile indicators (Q4 largest is reference). Q1=smallest'
            },
            'FCC × Firm_Size_Qi': {
                'type': 'Binary (0/1)',
                'desc': 'Interaction: FCC effect conditional on firm size quartile'
            }
        }
    },
    'Mechanism 2: CVSS Complexity': {
        'formula': 'ΔVolatility = β₀ + β₁·FCC + β₂·High_Complexity + β₃·FCC×High_Complexity + [controls] + ε',
        'n': '891',
        'method': 'OLS with HC3 robust standard errors',
        'variables': {
            'ΔVolatility': {
                'type': 'Continuous',
                'desc': 'Change in return volatility'
            },
            'FCC': {
                'type': 'Binary (0/1)',
                'desc': 'FCC indicator'
            },
            'High_Complexity': {
                'type': 'Binary (0/1)',
                'desc': 'Indicator = 1 if CVSS score in top quartile, 0 otherwise'
            },
            'FCC × High_Complexity': {
                'type': 'Binary (0/1)',
                'desc': 'Interaction: whether complexity amplifies FCC effect'
            }
        }
    },
    'Mechanism 3: Governance Quality': {
        'formula': 'ΔVolatility = β₀ + β₁·FCC + β₂·Low_Governance + β₃·FCC×Low_Governance + [controls] + ε',
        'n': '891',
        'method': 'OLS with HC3 robust standard errors',
        'variables': {
            'ΔVolatility': {
                'type': 'Continuous',
                'desc': 'Change in return volatility'
            },
            'FCC': {
                'type': 'Binary (0/1)',
                'desc': 'FCC indicator'
            },
            'Low_Governance': {
                'type': 'Binary (0/1)',
                'desc': 'Indicator = 1 if governance quality in bottom quartile, 0 otherwise'
            },
            'FCC × Low_Governance': {
                'type': 'Binary (0/1)',
                'desc': 'Interaction: governance quality moderation'
            }
        }
    },
    'Mechanism 4: Information Environment (Media + Reputation)': {
        'formula': 'ΔVolatility = β₀ + β₁·FCC + β₂·High_Info_Risk + β₃·FCC×High_Info_Risk + [controls] + ε',
        'n': '891',
        'method': 'OLS with HC3 robust standard errors',
        'variables': {
            'ΔVolatility': {
                'type': 'Continuous',
                'desc': 'Change in return volatility'
            },
            'FCC': {
                'type': 'Binary (0/1)',
                'desc': 'FCC indicator'
            },
            'High_Info_Risk': {
                'type': 'Binary (0/1)',
                'desc': 'Composite: 1 if high media attention OR repeat offender, 0 otherwise'
            },
            'FCC × High_Info_Risk': {
                'type': 'Binary (0/1)',
                'desc': 'Interaction: information environment moderation'
            }
        }
    }
}

for mech_name, mech_details in mech_formulas.items():
    doc.add_heading(mech_name, level=2)

    doc.add_paragraph('Specification:', style='Normal')
    formula_para = doc.add_paragraph()
    formula_run = formula_para.add_run(mech_details['formula'])
    formula_run.italic = True
    formula_run.font.name = 'Courier New'
    formula_run.font.size = Pt(10)

    doc.add_paragraph(f"Sample Size: {mech_details['n']}")
    doc.add_paragraph(f"Estimation Method: {mech_details['method']}")

    doc.add_paragraph('Variable Definitions:', style='Normal')
    vars_table = doc.add_table(rows=len(mech_details['variables']) + 1, cols=3)
    vars_table.style = 'Light Grid Accent 1'

    header_cells = vars_table.rows[0].cells
    header_cells[0].text = 'Variable Name'
    header_cells[1].text = 'Type'
    header_cells[2].text = 'Description'
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row_idx, (var_name, var_info) in enumerate(mech_details['variables'].items(), 1):
        row = vars_table.rows[row_idx]
        row.cells[0].text = var_name
        row.cells[1].text = var_info['type']
        row.cells[2].text = var_info['desc']

    doc.add_paragraph()

doc.add_page_break()

# ============================================================================
# ESSAY 3: EXECUTIVE TURNOVER
# ============================================================================
essay3_formulas = {
    'Model: Executive Turnover (Logit)': {
        'formula': 'P(Executive_Change_30d = 1) = Λ(β₀ + β₁·FCC_Reportable + β₂·Immediate_Disclosure + β₃·Health_Breach + β₄·Prior_Breaches + β₅·Firm_Size_log + β₆·Leverage + β₇·ROA)',
        'n': '896',
        'method': 'Logistic regression (binary logit); marginal effects reported (dP/dx evaluated at means)',
        'variables': {
            'Executive_Change_30d': {
                'type': 'Binary (0/1)',
                'desc': 'Indicator = 1 if any executive departure within 30 days of breach announcement, 0 otherwise (executive turnover measured via SEC 8-K filings)'
            },
            'FCC_Reportable': {
                'type': 'Binary (0/1)',
                'desc': 'FCC regulation indicator (telecommunications carriers)'
            },
            'Immediate_Disclosure': {
                'type': 'Binary (0/1)',
                'desc': 'Indicator = 1 if breach disclosed within 7 days, 0 otherwise'
            },
            'Health_Breach': {
                'type': 'Binary (0/1)',
                'desc': 'Breach involves protected health information'
            },
            'Prior_Breaches': {
                'type': 'Count',
                'desc': 'Total prior breaches for firm (reputation history)'
            },
            'Firm_Size_log': {
                'type': 'Continuous',
                'desc': 'Natural log of total assets'
            },
            'Leverage': {
                'type': 'Continuous',
                'desc': 'Total debt / Total assets'
            },
            'ROA': {
                'type': 'Continuous',
                'desc': 'Return on assets'
            }
        }
    },
    'Extended Model: 90-day Window': {
        'formula': 'P(Executive_Change_90d = 1) = Λ(β₀ + β₁·FCC_Reportable + β₂·Immediate_Disclosure + β₃·Health_Breach + β₄·Prior_Breaches + β₅·Firm_Size_log + β₆·Leverage + β₇·ROA)',
        'n': '896',
        'method': 'Logistic regression; marginal effects reported',
        'variables': {
            'Executive_Change_90d': {
                'type': 'Binary (0/1)',
                'desc': 'Indicator = 1 if any executive departure within 90 days of breach announcement, 0 otherwise'
            },
            'FCC_Reportable': {
                'type': 'Binary (0/1)',
                'desc': 'FCC regulation indicator'
            },
            'Immediate_Disclosure': {
                'type': 'Binary (0/1)',
                'desc': 'Disclosure within 7 days'
            },
            'Health_Breach': {
                'type': 'Binary (0/1)',
                'desc': 'Breach involves protected health information'
            },
            'Prior_Breaches': {
                'type': 'Count',
                'desc': 'Total prior breaches'
            },
            'Firm_Size_log': {
                'type': 'Continuous',
                'desc': 'Natural log of total assets'
            },
            'Leverage': {
                'type': 'Continuous',
                'desc': 'Total debt / Total assets'
            },
            'ROA': {
                'type': 'Continuous',
                'desc': 'Return on assets'
            }
        }
    },
    'Extended Model: 180-day Window': {
        'formula': 'P(Executive_Change_180d = 1) = Λ(β₀ + β₁·FCC_Reportable + β₂·Immediate_Disclosure + β₃·Health_Breach + β₄·Prior_Breaches + β₅·Firm_Size_log + β₆·Leverage + β₇·ROA)',
        'n': '896',
        'method': 'Logistic regression; marginal effects reported',
        'variables': {
            'Executive_Change_180d': {
                'type': 'Binary (0/1)',
                'desc': 'Indicator = 1 if any executive departure within 180 days of breach announcement, 0 otherwise'
            },
            'FCC_Reportable': {
                'type': 'Binary (0/1)',
                'desc': 'FCC regulation indicator'
            },
            'Immediate_Disclosure': {
                'type': 'Binary (0/1)',
                'desc': 'Disclosure within 7 days'
            },
            'Health_Breach': {
                'type': 'Binary (0/1)',
                'desc': 'Breach involves protected health information'
            },
            'Prior_Breaches': {
                'type': 'Count',
                'desc': 'Total prior breaches'
            },
            'Firm_Size_log': {
                'type': 'Continuous',
                'desc': 'Natural log of total assets'
            },
            'Leverage': {
                'type': 'Continuous',
                'desc': 'Total debt / Total assets'
            },
            'ROA': {
                'type': 'Continuous',
                'desc': 'Return on assets'
            }
        }
    }
}

add_formula_section(doc, 3, 'Executive Governance Response (Turnover)', essay3_formulas)

doc.add_page_break()

# ============================================================================
# APPENDIX: VARIABLE SUMMARY TABLE
# ============================================================================
doc.add_heading('APPENDIX: Complete Variable Summary', level=1)

doc.add_paragraph('Summary of all variables used across the three essays, with ranges and descriptive statistics.')

summary_table_data = [
    ['Variable', 'Type', 'Range/Distribution', 'Data Source', 'Used In Essays'],
    ['CAR_30d', 'Continuous', '[-0.30, +0.15]', 'CRSP (Fama-French 3-factor)', '1'],
    ['Volatility_Change', 'Continuous', '[-0.10, +0.20]', 'CRSP daily returns', '2'],
    ['Executive_Change_*d', 'Binary (0/1)', '{0, 1}', 'SEC 8-K filings', '3'],
    ['FCC_Reportable', 'Binary (0/1)', '{0, 1} [~24% = 1]', 'FCC Rule 37.3, SIC codes', '1, 2, 3'],
    ['Immediate_Disclosure', 'Binary (0/1)', '{0, 1} [~31% = 1]', 'PRC/DataBreaches.gov dates', '1, 3'],
    ['Disclosure_Delay_Days', 'Count', '[0, 730] days', 'PRC/DataBreaches.gov dates', '2'],
    ['Health_Breach', 'Binary (0/1)', '{0, 1} [~18% = 1]', 'PRC breach classification', '1, 2, 3'],
    ['Financial_Breach', 'Binary (0/1)', '{0, 1} [~22% = 1]', 'PRC breach classification', '1'],
    ['Prior_Breaches_Total', 'Count', '[0, 21]', 'PRC historical data', '1, 3'],
    ['Firm_Size_log', 'Continuous', '[2.5, 12.8]', 'Compustat total assets', '1, 2, 3'],
    ['Leverage', 'Continuous', '[0.0, 0.95]', 'Compustat debt/assets', '1, 2, 3'],
    ['ROA', 'Continuous', '[-0.45, 0.35]', 'Compustat net income/assets', '1, 2, 3'],
    ['Pre_Volatility', 'Continuous', '[0.005, 0.15]', 'CRSP [-25, -5] window', '2'],
    ['High_Complexity', 'Binary (0/1)', '{0, 1}', 'CVSS score quartile', '2 (mechanism)'],
    ['Low_Governance', 'Binary (0/1)', '{0, 1}', 'Governance quality metric', '2 (mechanism)'],
    ['High_Media_Coverage', 'Binary (0/1)', '{0, 1}', 'Media database', '2 (mechanism)'],
    ['Is_Repeat_Offender', 'Binary (0/1)', '{0, 1}', 'PRC prior breaches', '2 (mechanism)'],
]

summary_vars_table = doc.add_table(rows=len(summary_table_data), cols=5)
summary_vars_table.style = 'Light Grid Accent 1'

for row_idx, row_data in enumerate(summary_table_data):
    cells = summary_vars_table.rows[row_idx].cells
    for col_idx, cell_text in enumerate(row_data):
        cells[col_idx].text = cell_text
        if row_idx == 0:
            for paragraph in cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()

# ============================================================================
# METHODOLOGICAL NOTES
# ============================================================================
doc.add_page_break()
doc.add_heading('METHODOLOGICAL NOTES', level=1)

notes_sections = {
    'Estimation Method': 'All models use heteroscedasticity-robust standard errors (HC3) to address potential violations of homoskedasticity assumptions. For Essay 3 (logit), marginal effects are reported and evaluated at the sample mean of each covariate, providing interpretation as probability changes.',

    'Causal Identification': 'The FCC Rule 37.3 (47 CFR § 64.2011) natural experiment is used to identify causal effects of mandatory disclosure. Firms in specific SIC codes (4813, 4841, 4899) are subject to a 7-day mandatory disclosure rule starting January 1, 2007. Control group comprises all other industries with no such mandate. Parallel trends assumption tested via pre-2007 vs. post-2007 interaction.',

    'Event Study Window': 'For abnormal return calculations (Essay 1), the event window is [-5, +25] trading days relative to announcement date. Pre-event estimation window is [-240, -60] days. Volatility windows (Essay 2) are 20-trading-day blocks pre-breach ([-25, -5]) and post-breach ([+5, +25]).',

    'Data Sources': '(1) Privacy Rights Clearinghouse (PRC) DataBreaches.gov: ~1,054 publicly-traded firm breaches, 2006-2025. (2) CRSP: daily stock returns and trading volume. (3) Compustat: firm-level accounting data (assets, debt, net income). (4) SEC EDGAR: 8-K filings for executive departures. (5) FCC/FTC: regulatory status. CRSP match rate: 92%.',

    'Sample Construction': 'Observations are firm-breach level. Sample restricted to breaches affecting publicly-traded firms with complete market data (CRSP) and financial data (Compustat). Essays 1-2 exclude Essay 3 observations with missing executive data (different windows). Listwise deletion applied; no imputation.',
}

for section_title, section_text in notes_sections.items():
    doc.add_heading(section_title, level=2)
    doc.add_paragraph(section_text)

# Save
doc.save('Dissertation_Regression_Formulas.docx')
print("SUCCESS: Created Dissertation_Regression_Formulas.docx")
