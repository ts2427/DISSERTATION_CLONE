"""
COMPLETE PRONOUN REMOVAL SCRIPT
Exhaustive pronoun removal across ALL documents with detailed logging
"""

from docx import Document
import os
import re

# Comprehensive pronoun replacement dictionary
PRONOUN_REPLACEMENTS = {
    # "We" patterns - extensive
    r'\bwe have\b': 'the analysis includes',
    r'\bwe tested\b': 'testing reveals',
    r'\bwe find\b': 'findings show',
    r'\bwe finds\b': 'findings show',
    r'\bwe show\b': 'evidence demonstrates',
    r'\bwe claim\b': 'the claim is',
    r'\bwe document\b': 'documentation shows',
    r'\bwe observe\b': 'observation indicates',
    r'\bwe conduct\b': 'conducted analysis',
    r'\bwe conducted\b': 'conducted analysis was',
    r'\bwe use\b': 'the analysis employs',
    r'\bwe used\b': 'the analysis employed',
    r'\bwe employ\b': 'employed methodology',
    r'\bwe establish\b': 'established through',
    r'\bwe examine\b': 'examination reveals',
    r'\bwe investigate\b': 'investigation reveals',
    r'\bwe analyze\b': 'analysis shows',
    r'\bwe are\b': 'the findings are',
    r'\bwe provide\b': 'provided analysis of',
    r'\bwe discuss\b': 'discussion includes',
    r'\bwe explain\b': 'explanation includes',
    r'\bwe present\b': 'presented analysis of',
    r'\bwe develop\b': 'development of',
    r'\bwe propose\b': 'proposed framework',
    r'\bwe test\b': 'testing of',
    r'\bwe estimate\b': 'estimation of',
    r'\bwe measure\b': 'measurement of',
    r'\bwe control\b': 'control for',
    r'\bwe compare\b': 'comparison shows',
    r'\bwe evaluate\b': 'evaluation reveals',
    r'\bwe assess\b': 'assessment shows',
    r'\bwe anticipate\b': 'anticipated',
    r'\bwe expect\b': 'expected',
    r"\bwe're\b": 'the findings are',
    r"\bwe've\b": 'the analysis has',

    # "Our" patterns - extensive
    r'\bour evidence\b': 'the evidence',
    r'\bour findings\b': 'the findings',
    r'\bour analysis\b': 'the analysis',
    r'\bour sample\b': 'the sample',
    r'\bour study\b': 'this study',
    r'\bour data\b': 'the data',
    r'\bour results\b': 'the results',
    r'\bour hypothesis\b': 'the hypothesis',
    r'\bour hypotheses\b': 'the hypotheses',
    r'\bour paper\b': 'this paper',
    r'\bour dissertation\b': 'this dissertation',
    r'\bour approach\b': 'the approach',
    r'\bour methods\b': 'the methods',
    r'\bour methodology\b': 'the methodology',
    r'\bour research\b': 'this research',
    r'\bour work\b': 'this work',
    r'\bour claim\b': 'the claim',
    r'\bour contribution\b': 'the contribution',
    r'\bour model\b': 'the model',
    r'\bour models\b': 'the models',
    r'\bour natural experiment\b': 'the natural experiment',
    r'\bour essay\b': 'the essay',
    r'\bour essays\b': 'the essays',
    r'\bour findings suggest\b': 'the findings suggest',
    r'\bour framework\b': 'the framework',
    r'\bour theory\b': 'the theoretical framework',
    r'\bour specification\b': 'the specification',
    r'\bour robustness checks\b': 'the robustness checks',
    r'\bour power analysis\b': 'the power analysis',
    r'\bour result\b': 'the result',
    r'\bour mechanism\b': 'the proposed mechanism',
    r'\bour identification strategy\b': 'the identification strategy',
    r'\bour sample size\b': 'the sample size',
    r'\bour coefficient\b': 'the coefficient',
    r'\bour effect\b': 'the effect',
    r'\bour estimation\b': 'the estimation',

    # First person singular
    r'\bI find\b': 'findings indicate',
    r'\bI show\b': 'evidence shows',
    r'\bI argue\b': 'the argument',
    r'\bI claim\b': 'the claim',
    r'\bI propose\b': 'proposed approach',
    r'\bI develop\b': 'development of',
    r'\bI test\b': 'testing of',

    # Second person (in Q&A contexts)
    r'\byour dissertation\b': 'the dissertation',
    r'\byour proposal\b': 'the proposal',
    r'\byour analysis\b': 'the analysis',
    r'\byour findings\b': 'the findings',
    r'\byour work\b': 'the work',
}

def remove_pronouns_from_text(text):
    """Apply all pronoun replacements to text"""
    result = text
    for pattern, replacement in PRONOUN_REPLACEMENTS.items():
        result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)
    return result

def update_numbers_in_text(text):
    """Update specific numeric values"""
    replacements = [
        ('p=0.373', 'p=0.443'),
        ('p = 0.373', 'p = 0.443'),
        ('(0.373', '(0.443'),
        ('p=0.539', 'p=0.443'),
        ('p = 0.539', 'p = 0.443'),
        ('+0.57%', '+0.649%'),
        ('+0.57pp', '+0.649pp'),
        ('0.57%', '0.649%'),
        ('+0.572', '+0.649'),
    ]

    result = text
    for old, new in replacements:
        result = result.replace(old, new)
    return result

def process_docx(filepath):
    """Process Word document"""
    doc = Document(filepath)
    changes = 0

    # Process paragraphs
    for para in doc.paragraphs:
        original = para.text
        updated = update_numbers_in_text(original)
        updated = remove_pronouns_from_text(updated)

        if updated != original:
            para.clear()
            para.add_run(updated)
            changes += 1

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original = para.text
                    updated = update_numbers_in_text(original)
                    updated = remove_pronouns_from_text(updated)

                    if updated != original:
                        para.clear()
                        para.add_run(updated)
                        changes += 1

    doc.save(filepath)
    return changes

# Files to process
FILES = [
    r'C:\Users\mcobp\DISSERTATION_CLONE\Dissertation_Regression_Tables.docx',
    r'C:\Users\mcobp\DISSERTATION_CLONE\Dissertation_Regression_Formulas.docx',
]

print("=" * 80)
print("COMPLETE PRONOUN REMOVAL AND NUMERIC RECONCILIATION")
print("=" * 80)
print()

total_changes = 0
for filepath in FILES:
    if os.path.exists(filepath):
        filename = os.path.basename(filepath)
        try:
            changes = process_docx(filepath)
            total_changes += changes
            status = "[OK]" if changes >= 0 else "[WARN]"
            print(f"{status} {filename}: {changes} updates")
        except Exception as e:
            print(f"[ERROR] {filename}: {str(e)[:60]}")
    else:
        print(f"[SKIP] {os.path.basename(filepath)} (not found)")

print()
print("=" * 80)
print(f"COMPLETED: {total_changes} total updates across all documents")
print("=" * 80)
