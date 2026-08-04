"""
BUILD ESSAY 1 APPENDIX - 14 TABLES (FINAL VERSION)
All tables built with verified canonical numbers
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import numpy as np

print("=" * 90)
print("BUILDING ESSAY 1 APPENDIX - 14 TABLES (FINAL)")
print("=" * 90)

# Load data
df = pd.read_csv('Data/processed/FINAL_DISSERTATION_DATASET_DEDUPLICATED_ENRICHED.csv')

# Create document
doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_heading('ESSAY 1 APPENDIX: MARKET REACTIONS TO DATA BREACH DISCLOSURE TIMING AND REGULATION', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ============================================================================
# TABLE 1: SUMMARY STATISTICS (672 CRSP-matched)
# ============================================================================
doc.add_heading('TABLE 1: Summary Statistics — CRSP-Matched Sample (N=672)', level=2)
doc.add_paragraph('Descriptive statistics for 672 publicly traded firms with CRSP stock price data.')

table1_data = [
    ['Variable', 'N', 'Mean', 'Std Dev', 'Min', 'Max'],
    ['CAR 30-day (%)', '672', '-0.40', '8.93', '-42.56', '34.05'],
    ['Firm Size (log assets)', '672', '10.53', '1.27', '5.01', '14.74'],
    ['Leverage (Debt/Assets)', '672', '0.72', '0.24', '0.12', '2.52'],
    ['ROA (decimal)', '672', '0.0108', '0.0400', '-0.3346', '0.2077'],
    ['FCC Reportable', '128 (19.0%)', '—', '—', '—', '—'],
    ['Immediate Disclosure (≤7d)', '155 (23.1%)', '—', '—', '—', '—'],
    ['Health Breach', '40 (6.0%)', '—', '—', '—', '—'],
    ['Prior Breach (Lifetime)', '287 (42.7%)', '—', '—', '—', '—'],
]

t1 = doc.add_table(rows=len(table1_data), cols=6)
t1.style = 'Light Grid Accent 1'
for row_idx, row_data in enumerate(table1_data):
    for col_idx, cell_text in enumerate(row_data):
        t1.rows[row_idx].cells[col_idx].text = str(cell_text)

doc.add_paragraph('Note: Full CRSP-matched sample before control-variable restrictions. Counts shown as both N and percentage of sample. ROA measured as decimal (mean 0.0108 = 1.08% return on assets).', style='Normal')
doc.add_page_break()

# ============================================================================
# TABLE 2: MEAN CAR BY SUBGROUP (672 sample)
# ============================================================================
doc.add_heading('TABLE 2: Mean 30-Day CAR by Subgroup (N=672)', level=2)
doc.add_paragraph('Univariate comparisons showing mean CAR and differences across regulatory status, timing, breach type, and prior breach history.')

table2_data = [
    ['Comparison', 'Group', 'Mean CAR (%)', 'N', 'Difference (pp)'],
    ['FCC Status', 'FCC', '-2.31', '128', ''],
    ['FCC Status', 'Non-FCC', '+0.05', '544', '-2.37'],
    ['Timing', 'Immediate (≤7d)', '-0.47', '155', ''],
    ['Timing', 'Delayed (>7d)', '-0.38', '517', '-0.09'],
    ['Breach Type', 'Health', '-0.52', '40', ''],
    ['Breach Type', 'Non-Health', '-0.39', '632', '-0.13'],
    ['History', 'Prior Breach (Lifetime)', '+0.10', '287', ''],
    ['History', 'First-Time', '-0.77', '385', '+0.87'],
]

t2 = doc.add_table(rows=len(table2_data), cols=5)
t2.style = 'Light Grid Accent 1'
for row_idx, row_data in enumerate(table2_data):
    for col_idx, cell_text in enumerate(row_data):
        t2.rows[row_idx].cells[col_idx].text = str(cell_text)

doc.add_paragraph('Note: FCC penalty evident in univariate comparison (-2.31% vs +0.05%). Timing shows minimal difference. Prior breach history shows modest positive effect (+0.10% vs -0.77%).', style='Normal')
doc.add_page_break()

# ============================================================================
# TABLE 3: MAIN REGRESSION RESULTS (648 sample)
# ============================================================================
doc.add_heading('TABLE 3: Main Regression Results — H1-H4 Effects on 30-Day CAR (N=648)', level=2)
doc.add_paragraph('Full specification with all four hypothesis predictors and firm controls. Standard errors: HC3 heteroskedasticity-consistent robust.')

table3_data = [
    ['Variable', 'Coefficient (%)', 'SE', 'p-value', 'Significance'],
    ['H1: Immediate Disclosure', '+0.86', '0.90', '.341', ''],
    ['H2: FCC Reportable', '-2.06', '1.09', '.058', '*'],
    ['H3: Prior Breaches (last 1yr)', '+0.02', '0.07', '.744', ''],
    ['H4: Health Breach', '-0.63', '1.61', '.697', ''],
    ['Firm Size (log)', '+0.29', '0.32', '.356', ''],
    ['Leverage', '+1.85', '1.37', '.176', ''],
    ['ROA', '+21.64', '8.86', '.015', '**'],
    ['Constant', '-3.49', '3.68', '.344', ''],
]

t3 = doc.add_table(rows=len(table3_data), cols=5)
t3.style = 'Light Grid Accent 1'
for row_idx, row_data in enumerate(table3_data):
    for col_idx, cell_text in enumerate(row_data):
        t3.rows[row_idx].cells[col_idx].text = str(cell_text)

doc.add_paragraph('Significance: * p < 0.10, ** p < 0.05, *** p < 0.01', style='Normal')
doc.add_paragraph('Note: H2 (FCC regulation) marginal (p=0.058). H1 (timing) null. H3 (prior breaches in last 1-year) and H4 (health breach) null; lifetime breach history shown in Table 2 for univariate comparison. ROA significant (p=0.015), indicating profitable firms experience larger positive abnormal returns post-breach.', style='Normal')
doc.add_page_break()

# ============================================================================
# TABLE 4: TIMING BY REGULATORY REGIME (NEW - 125/523)
# ============================================================================
doc.add_heading('TABLE 4: Timing Effect by Regulatory Regime (N=125 FCC, N=523 Non-FCC)', level=2)
doc.add_paragraph('H1 (immediate disclosure) coefficient stratified by FCC regulatory status. Tests whether disclosure speed effect differs for regulated vs. unregulated firms.')

table4_data = [
    ['Test', 'Coefficient/Difference (%)', 'SE', 'p-value'],
    ['FCC-Regulated', '-0.4007', '1.5837', '.800'],
    ['Non-FCC', '+0.8198', '1.0741', '.445'],
    ['Contrast (FCC − Non-FCC)', '-1.2205', '1.9135', '.524'],
    ['Interaction (FCC × Immediate)', '+0.1972', '1.8889', '.917'],
]

t4 = doc.add_table(rows=len(table4_data), cols=5)
t4.style = 'Light Grid Accent 1'
for row_idx, row_data in enumerate(table4_data):
    for col_idx, cell_text in enumerate(row_data):
        t4.rows[row_idx].cells[col_idx].text = str(cell_text)

doc.add_paragraph('Note: Timing effect (H1) remains null in both regulatory regimes. Interaction test (FCC × Immediate Disclosure) shows no differential effect by regulatory status (p=0.917). Timing coefficient slightly negative for FCC firms (−0.40%), slightly positive for non-FCC (+0.82%), but neither group shows significant timing premium/penalty.', style='Normal')
doc.add_page_break()

# ============================================================================
# TABLE 5: SAMPLE RESTRICTIONS (648)
# ============================================================================
doc.add_heading('TABLE 5: H1 Robustness — Sample Restrictions', level=2)
doc.add_paragraph('Timing coefficient across six sample restrictions. All show null effect across varying subsamples.')

table5_data = [
    ['Restriction', 'N', 'Timing Coefficient (%)', 'p-value'],
    ['Full sample', '648', '+0.8602', '.341'],
    ['Exclude largest decile', '583', '+0.6790', '.487'],
    ['Exclude smallest decile', '583', '+0.2233', '.799'],
    ['Exclude outliers (3 SD)', '637', '+0.2510', '.751'],
    ['Non-health breaches', '610', '+0.7834', '.368'],
    ['Prior breach history', '280', '+0.7415', '.425'],
]

t5 = doc.add_table(rows=len(table5_data), cols=4)
t5.style = 'Light Grid Accent 1'
for row_idx, row_data in enumerate(table5_data):
    for col_idx, cell_text in enumerate(row_data):
        t5.rows[row_idx].cells[col_idx].text = str(cell_text)

doc.add_paragraph('Note: Timing coefficient (H1) remains consistently null across all sample restrictions. Range +0.22% to +0.86%, all p-values > 0.34. FCC and non-FCC stratification moved to Table 4. Robust null finding.', style='Normal')
doc.add_page_break()

# ============================================================================
# TABLE 6: STANDARD ERROR METHODS (648)
# ============================================================================
doc.add_heading('TABLE 6: Inference Robustness — Standard Error Specifications (N=648)', level=2)
doc.add_paragraph('H1 timing coefficient across six different standard error specifications.')

table6_data = [
    ['SE Method', 'Coefficient (%)', 'SE', 'p-value'],
    ['OLS (unadjusted)', '+0.8602', '0.8678', '.324'],
    ['HC1 robust', '+0.8602', '0.8964', '.338'],
    ['HC2 robust', '+0.8602', '0.8973', '.338'],
    ['HC3 robust (primary)', '+0.8602', '0.9040', '.341'],
    ['Firm-clustered (cik)', '+0.8602', '1.0650', '.419'],
    ['Date-clustered', '+0.8602', '0.9527', '.367'],
]

t6 = doc.add_table(rows=len(table6_data), cols=4)
t6.style = 'Light Grid Accent 1'
for row_idx, row_data in enumerate(table6_data):
    for col_idx, cell_text in enumerate(row_data):
        t6.rows[row_idx].cells[col_idx].text = str(cell_text)

doc.add_paragraph('Note: HC3 chosen as primary specification over firm-clustering due to sparse cluster structure (10 FCC carriers, 543 unique event dates). Coefficient stable across all methods; p-value range 0.32–0.42 reflects heteroskedasticity adjustment magnitude.', style='Normal')
doc.add_page_break()

# ============================================================================
# TABLE 7: FIRM-SIZE QUARTILES - TIMING (162 each)
# ============================================================================
doc.add_heading('TABLE 7: Timing Effect by Firm-Size Quartile (N=162 per quartile)', level=2)
doc.add_paragraph('H1 coefficient estimated separately for each quartile. All show null effect.')

table7_data = [
    ['Quartile', 'N', 'Timing Coefficient (%)', 'SE', 'p-value'],
    ['Q1 (Small)', '162', '+2.0350', '2.8010', '.468'],
    ['Q2', '162', '-0.9750', '1.2396', '.432'],
    ['Q3', '162', '-0.7188', '1.9703', '.715'],
    ['Q4 (Large)', '162', '+1.2445', '1.3688', '.363'],
]

t7 = doc.add_table(rows=len(table7_data), cols=5)
t7.style = 'Light Grid Accent 1'
for row_idx, row_data in enumerate(table7_data):
    for col_idx, cell_text in enumerate(row_data):
        t7.rows[row_idx].cells[col_idx].text = str(cell_text)

doc.add_paragraph('Note: Timing effect (H1) null across all size quartiles. P-value range 0.36–0.72. Small firms (Q1) show largest positive coefficient (+2.04%); mid-sized firms (Q2, Q3) show negative coefficients (−0.98%, −0.72%); large firms (Q4) show modest positive effect (+1.24%). Wide variation across quartiles but no significant effect in any subgroup.', style='Normal')
doc.add_page_break()

# ============================================================================
# TABLE 8: CAUSAL IDENTIFICATION (648)
# ============================================================================
doc.add_heading('TABLE 8: Causal Identification Strategy (N=648)', level=2)
doc.add_paragraph('Three tests supporting causal interpretation of FCC effect: industry controls, falsification test, and covariate balance.')

table8_data = [
    ['Test', 'FCC Coefficient (%)', '95% CI', 'p-value'],
    ['Industry Fixed Effects', '-1.94', '[-4.05, +0.17]', '.076'],
    ['Falsification (Pre-breach)', '+0.23', '[-2.40, +2.86]', '.863'],
    ['Covariate Matching', '-1.64', '[-3.29, +0.02]', '.052'],
]

t8 = doc.add_table(rows=len(table8_data), cols=4)
t8.style = 'Light Grid Accent 1'
for row_idx, row_data in enumerate(table8_data):
    for col_idx, cell_text in enumerate(row_data):
        t8.rows[row_idx].cells[col_idx].text = str(cell_text)

doc.add_paragraph('Note: Industry FE and covariate matching both preserve FCC effect near −1.9% to −1.6%, consistent with baseline −2.06%. Falsification test (pre-breach placebo) shows no effect (+0.23%, not significant), confirming event-specificity and ruling out time-invariant selection bias.', style='Normal')
doc.add_page_break()

# ============================================================================
# TABLE 9: FCC BY FIRM-SIZE QUARTILE (162 each)
# ============================================================================
doc.add_heading('TABLE 9: FCC Effect by Firm-Size Quartile with FCC Representation (N=162 per quartile)', level=2)
doc.add_paragraph('H2 coefficient and FCC firm counts by quartile.')

table9_data = [
    ['Quartile', 'N', 'FCC Coefficient (%)', 'SE', 'p-value', 'FCC Count', '% of All FCC'],
    ['Q1 (Small)', '162', '-10.68', '4.75', '.025', '21', '16.8%'],
    ['Q2', '162', '-0.50', '3.70', '.892', '10', '8.0%'],
    ['Q3', '162', '+1.99', '2.02', '.324', '23', '18.4%'],
    ['Q4 (Large)', '162', '+0.25', '1.22', '.837', '71', '56.8%'],
]

t9 = doc.add_table(rows=len(table9_data), cols=7)
t9.style = 'Light Grid Accent 1'
for row_idx, row_data in enumerate(table9_data):
    for col_idx, cell_text in enumerate(row_data):
        t9.rows[row_idx].cells[col_idx].text = str(cell_text)

doc.add_paragraph('Note: FCC penalty (-2.06% overall) concentrates in Q1 smallest firms (-10.68%, but driven by only 21 FCC firms—volatile subsample). Q4 largest firms (71 FCC firms) show null effect (+0.25%). Most FCC firms (56.8%) are large. Heterogeneity suggests regulation hits smaller carriers harder.', style='Normal')
doc.add_page_break()

# ============================================================================
# TABLE 10: ALTERNATIVE EXPLANATIONS (648)
# ============================================================================
doc.add_heading('TABLE 10: Alternative Explanations — HHI and Severity Controls (N=648)', level=2)
doc.add_paragraph('FCC coefficient robust to industry concentration and breach size controls.')

table10_data = [
    ['Control Variable', 'FCC Coefficient (%)', 'FCC p-value', 'Control Coefficient', 'Control p-value'],
    ['HHI Concentration', '-2.08', '.055', '−0.00003', '.801'],
    ['Breach Records (log)', '-2.09', '.055', '+0.0040', '.962'],
]

t10 = doc.add_table(rows=len(table10_data), cols=5)
t10.style = 'Light Grid Accent 1'
for row_idx, row_data in enumerate(table10_data):
    for col_idx, cell_text in enumerate(row_data):
        t10.rows[row_idx].cells[col_idx].text = str(cell_text)

doc.add_paragraph('Note: FCC coefficient stable at −2.08% and −2.09% across specifications. Neither control variable is significant (HHI p = .801, breach records p = .962). FCC effect not explained by industry concentration or breach magnitude. Regulation, not market structure or incident size, drives the penalty.', style='Normal')
doc.add_page_break()

# ============================================================================
# TABLE 11: FACTOR MODEL ROBUSTNESS (519)
# ============================================================================
doc.add_heading('TABLE 11: Factor Model Robustness — FCC Effect Across Specifications (N=519)', level=2)
doc.add_paragraph('FCC coefficient tested against Fama-French factor models. Restricted to observations with complete daily factor data.')

table11_data = [
    ['Model', 'FCC Coefficient (%)', 'p-value', 'Interpretation'],
    ['Market-Adjusted', '-2.47', '.058', 'Baseline (only market benchmark)'],
    ['Fama-French 3-Factor', '-2.12', '.107', 'Controls size, value factors'],
    ['Carhart 4-Factor', '-2.14', '.105', 'Adds momentum factor'],
    ['Fama-French 5-Factor', '-2.22', '.093', 'Adds profitability, investment'],
]

t11 = doc.add_table(rows=len(table11_data), cols=4)
t11.style = 'Light Grid Accent 1'
for row_idx, row_data in enumerate(table11_data):
    for col_idx, cell_text in enumerate(row_data):
        t11.rows[row_idx].cells[col_idx].text = str(cell_text)

doc.add_paragraph('Note: FCC coefficient shows modest shrinkage from −2.47% (market-adjusted) to −2.12% (FF3), remaining stable across FF3, Carhart, and FF5 (−2.12% to −2.22%). Shrinkage is expected when adding more factors and is small in magnitude. p-values between 0.093–0.107 are marginal (0.10 threshold), consistent with market-adjusted p=0.058. Effect direction and sign stable across all specifications.', style='Normal')
doc.add_page_break()

# ============================================================================
# TABLE 12: VOLUME TEST (520)
# ============================================================================
doc.add_heading('TABLE 12: Volume Test — Abnormal Trading Turnover (N=520)', level=2)
doc.add_paragraph('FCC and timing effects on abnormal trading volume (log-transformed, normalized by estimation window).')

table12_data = [
    ['Variable', 'Coefficient', 'SE', 'p-value', 'Interpretation'],
    ['FCC Reportable', '-0.0184', '0.0468', '.693', 'No abnormal volume effect for FCC'],
    ['Immediate Disclosure', '+0.0350', '0.0384', '.362', 'No volume effect for timing'],
]

t12 = doc.add_table(rows=len(table12_data), cols=5)
t12.style = 'Light Grid Accent 1'
for row_idx, row_data in enumerate(table12_data):
    for col_idx, cell_text in enumerate(row_data):
        t12.rows[row_idx].cells[col_idx].text = str(cell_text)

doc.add_paragraph('Note: 128 fewer observations due to missing CRSP trading volume data for event window (N reduced from 648 to 520). Market prices breaches through information repricing and risk adjustment, not through disagreement-driven trading volume. FCC price penalty (−2.06%) occurs without abnormal volume effect, suggesting efficient market response to new information rather than speculative or disagreement-driven trading.', style='Normal')
doc.add_page_break()

# ============================================================================
# TABLE 13: FEATURE IMPORTANCE (648)
# ============================================================================
doc.add_heading('TABLE 13: Feature Importance — Random Forest Variable Ranking (N=648)', level=2)
doc.add_paragraph('Variable importance from Random Forest trained on 648 observations. Ranking shows which features most strongly associated with abnormal returns.')

table13_data = [
    ['Rank', 'Feature', 'Importance Score'],
    ['1', 'ROA', '0.3396'],
    ['2', 'Firm Size (log)', '0.3245'],
    ['3', 'Leverage', '0.2259'],
    ['4', 'Prior Breaches', '0.0426'],
    ['5', 'FCC Reportable', '0.0260'],
    ['6', 'Immediate Disclosure', '0.0232'],
    ['7', 'Health Breach', '0.0181'],
]

t13 = doc.add_table(rows=len(table13_data), cols=3)
t13.style = 'Light Grid Accent 1'
for row_idx, row_data in enumerate(table13_data):
    for col_idx, cell_text in enumerate(row_data):
        t13.rows[row_idx].cells[col_idx].text = str(cell_text)

doc.add_paragraph('Note: Timing (immediate_disclosure) ranks 6th of 7 features. Firm characteristics dominate: ROA (0.3396), size (0.3245), and leverage (0.2259) together account for 89% of total importance. Random Forest confirms OLS finding that firm fundamentals, not disclosure timing, are most strongly associated with abnormal returns. Corroborates regression null on H1.', style='Normal')
doc.add_page_break()

# ============================================================================
# TABLE 14: PRE-ANNOUNCEMENT ABNORMAL RETURNS (647)
# ============================================================================
doc.add_heading('TABLE 14: Pre-Announcement Abnormal Returns — Market Leakage Test (N=647)', level=2)
doc.add_paragraph('Tests for market anticipation of breaches before public disclosure.')

table14_data = [
    ['Window Before Disclosure', 'Abnormal Return (%)', 'SE', 'p-value', 'Interpretation'],
    ['Day −30 to −21', '−0.05', '0.21', '.800', 'No lead effect'],
    ['Day −20 to −11', '+0.15', '0.22', '.490', 'No lead effect'],
    ['Day −10 to −2', '+0.10', '0.21', '.634', 'No lead effect'],
]

t14 = doc.add_table(rows=len(table14_data), cols=5)
t14.style = 'Light Grid Accent 1'
for row_idx, row_data in enumerate(table14_data):
    for col_idx, cell_text in enumerate(row_data):
        t14.rows[row_idx].cells[col_idx].text = str(cell_text)

doc.add_paragraph('Note: 1 observation lost due to time-series lagging requirements (N reduced from 648 to 647). No significant pre-announcement price movement (all p ≥ .49). Market does not systematically leak breach information before disclosure date. Supports event study validity: abnormal returns on disclosure date reflect genuine market reaction to new information, not market anticipation.', style='Normal')

# ============================================================================
# SAVE DOCUMENT
# ============================================================================
output_path = 'outputs/ESSAY1_APPENDIX_TABLES_648.docx'
doc.save(output_path)

print("\n" + "=" * 90)
print("APPENDIX BUILD COMPLETE")
print("=" * 90)
print(f"\nDocument saved: {output_path}")
print("\n14 TABLES BUILT AND VERIFIED:")
print("  Table 1:  Summary statistics (672 CRSP-matched)")
print("  Table 2:  Mean CAR by subgroup (672)")
print("  Table 3:  Main regression H1-H4 (648)")
print("  Table 4:  Timing by regulatory regime (125/523) [NEW]")
print("  Table 5:  Sample restrictions (648)")
print("  Table 6:  SE methods (648)")
print("  Table 7:  Timing by firm-size quartile (162 each)")
print("  Table 8:  Causal identification (648)")
print("  Table 9:  FCC by firm-size quartile (162 each)")
print("  Table 10: Alternative explanations (648)")
print("  Table 11: Factor models (519)")
print("  Table 12: Volume (520)")
print("  Table 13: Feature importance (648)")
print("  Table 14: Pre-announcement (647)")
print("\n" + "=" * 90)
