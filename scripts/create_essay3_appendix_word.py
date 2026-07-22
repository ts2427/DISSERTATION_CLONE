"""
CREATE ESSAY 3 APPENDIX IN WORD FORMAT
Convert markdown tables to python-docx with proper formatting
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd

print("=" * 90)
print("CREATING ESSAY 3 APPENDIX WORD DOCUMENT")
print("=" * 90)

# Create document
doc = Document()

# Set up document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_heading('ESSAY 3 APPENDIX: EXECUTIVE TURNOVER AND GOVERNANCE RESPONSE', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ============================================================================
# TABLE A1
# ============================================================================
doc.add_heading('TABLE A1: Sample Characteristics and CEO Turnover Rates by Regulatory Status and Disclosure Timing', level=2)

doc.add_paragraph('N = 651 observations')

doc.add_heading('Overall Sample Composition', level=3)
doc.add_paragraph('Total observations: 651', style='List Bullet')
doc.add_paragraph('FCC-regulated firms: n = 140 (21.5%)', style='List Bullet')
doc.add_paragraph('Non-FCC firms: n = 511 (78.5%)', style='List Bullet')

doc.add_heading('Turnover Events by Temporal Window', level=3)
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Window'
hdr_cells[1].text = 'Events'
hdr_cells[2].text = '% of Sample'
table.rows[1].cells[0].text = '30-day'
table.rows[1].cells[1].text = '247'
table.rows[1].cells[2].text = '37.9%'
table.rows[2].cells[0].text = '90-day'
table.rows[2].cells[1].text = '379'
table.rows[2].cells[2].text = '58.2%'
table.rows[3].cells[0].text = '180-day'
table.rows[3].cells[1].text = '385'
table.rows[3].cells[2].text = '59.1%'

doc.add_heading('Turnover Rate by FCC Regulatory Status', level=3)
table = doc.add_table(rows=4, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Window'
hdr_cells[1].text = 'FCC Rate'
hdr_cells[2].text = 'Non-FCC Rate'
hdr_cells[3].text = 'Difference (pp)'
table.rows[1].cells[0].text = '30-day'
table.rows[1].cells[1].text = '37.1%'
table.rows[1].cells[2].text = '38.2%'
table.rows[1].cells[3].text = '-1.1'
table.rows[2].cells[0].text = '90-day'
table.rows[2].cells[1].text = '60.0%'
table.rows[2].cells[2].text = '57.7%'
table.rows[2].cells[3].text = '+2.3'
table.rows[3].cells[0].text = '180-day'
table.rows[3].cells[1].text = '60.7%'
table.rows[3].cells[2].text = '58.7%'
table.rows[3].cells[3].text = '+2.0'

doc.add_heading('Turnover Rate by Disclosure Speed', level=3)
table = doc.add_table(rows=4, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Window'
hdr_cells[1].text = 'Immediate (≤7d)'
hdr_cells[2].text = 'Delayed (>7d)'
hdr_cells[3].text = 'Difference (pp)'
table.rows[1].cells[0].text = '30-day'
table.rows[1].cells[1].text = '33.8%'
table.rows[1].cells[2].text = '39.3%'
table.rows[1].cells[3].text = '-5.5'
table.rows[2].cells[0].text = '90-day'
table.rows[2].cells[1].text = '60.6%'
table.rows[2].cells[2].text = '57.4%'
table.rows[2].cells[3].text = '+3.2'
table.rows[3].cells[0].text = '180-day'
table.rows[3].cells[1].text = '62.5%'
table.rows[3].cells[2].text = '58.0%'
table.rows[3].cells[3].text = '+4.5'

doc.add_heading('Mean Firm Characteristics by FCC Status', level=3)
table = doc.add_table(rows=4, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Characteristic'
hdr_cells[1].text = 'FCC Mean'
hdr_cells[2].text = 'Non-FCC Mean'
hdr_cells[3].text = 'Difference'
table.rows[1].cells[0].text = 'Log Market Cap'
table.rows[1].cells[1].text = '11.29'
table.rows[1].cells[2].text = '10.49'
table.rows[1].cells[3].text = '+0.80'
table.rows[2].cells[0].text = 'Leverage (D/A)'
table.rows[2].cells[1].text = '0.568'
table.rows[2].cells[2].text = '0.561'
table.rows[2].cells[3].text = '+0.007'
table.rows[3].cells[0].text = 'ROA'
table.rows[3].cells[1].text = '0.0108'
table.rows[3].cells[2].text = '0.0145'
table.rows[3].cells[3].text = '-0.0037'

note = doc.add_paragraph('Note: Sample drawn from Privacy Rights Clearinghouse database (2005–2025), deduplicated to one observation per firm-breach-date pair. Turnover rates remarkably consistent across regulatory conditions, suggesting executive changes driven by breach events themselves rather than regulatory mandates.')
note.style = 'Intense Quote'

doc.add_page_break()

# ============================================================================
# TABLE A2
# ============================================================================
doc.add_heading('TABLE A2: Reduced-Form Logistic Regression — FCC Effect on CEO Turnover', level=2)

doc.add_paragraph('Specification: logit(P(Turnover = 1)) = β₀ + β₁(fcc_reportable) + β₂(firm_size_log) + β₃(leverage) + β₄(roa) + ε')

table = doc.add_table(rows=8, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Coefficient'
hdr_cells[1].text = '30-day'
hdr_cells[2].text = '90-day'
hdr_cells[3].text = '180-day'

data = [
    ['FCC logit coefficient', '0.1101', '0.0045', '-0.0605'],
    ['Standard error', '0.2072', '0.2050', '0.2049'],
    ['p-value', '.595', '.982', '.768'],
    ['Average Marginal Effect (pp)', '+2.55', '+0.11', '-1.42'],
    ['AME Std Error (pp)', '4.769', '4.843', '4.804'],
    ['MDE at 80% power (pp)', '13.36', '13.57', '13.46'],
    ['N', '651', '651', '651'],
]

for i, row_data in enumerate(data, start=1):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text

note = doc.add_paragraph('Controls: firm_size_log, leverage, roa, industry fixed effects | Standard Errors: Industry-clustered robust (HC3)')
note.style = 'Intense Quote'

note2 = doc.add_paragraph('Dependent variable = 1 if CEO departed within specified window post-disclosure, 0 otherwise. Average marginal effects converted to percentage-point probability scale at sample mean turnover rate (~38%). All effects statistically non-significant and well below minimum detectable effect.')
note2.style = 'Intense Quote'

doc.add_page_break()

# ============================================================================
# TABLE A3
# ============================================================================
doc.add_heading('TABLE A3: Control Variable Coefficients from Base Logistic Regression', level=2)

table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Control Variable'
hdr_cells[1].text = '30-day'
hdr_cells[2].text = '90-day'
hdr_cells[3].text = '180-day'

data = [
    ['Leverage', '0.9524** (0.0065)', '-0.1821 (0.5942)', '-0.2658 (0.4378)'],
    ['Prior Breach History', '-0.0158 (0.2245)', '-0.0283* (0.0189)', '-0.0272* (0.0234)'],
    ['Health Breach', '0.5013 (0.1417)', '0.7559† (0.0563)', '0.8686** (0.0348)'],
    ['Return on Assets', '-2.2607 (0.3565)', '-4.7959† (0.0610)', '-5.1581* (0.0460)'],
    ['Firm Size (log)', '0.0202 (0.7701)', '-0.0850 (0.2194)', '-0.0610 (0.3784)'],
]

for i, row_data in enumerate(data, start=1):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text

note = doc.add_paragraph('Format: coefficient (p-value) | † p < .10; * p < .05; ** p < .01')
note.style = 'Intense Quote'

note2 = doc.add_paragraph('Reported in log-odds scale. Industry fixed effects included but not reported. FCC regulatory status reported separately in Table A2. Significance patterns reveal breach characteristics and firm financial condition drive turnover; regulatory status does not.')
note2.style = 'Intense Quote'

doc.add_page_break()

# ============================================================================
# TABLE A4
# ============================================================================
doc.add_heading('TABLE A4: Two One-Sided Tests (TOST) of Equivalence', level=2)

doc.add_paragraph('Equivalence Bound: ±10 percentage points (governance-meaningful threshold)')

table = doc.add_table(rows=7, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = '30-day'
hdr_cells[2].text = '90-day'
hdr_cells[3].text = '180-day'

data = [
    ['Point Estimate (AME, pp)', '2.55', '0.11', '-1.42'],
    ['Standard Error (pp)', '4.769', '4.843', '4.804'],
    ['90% CI Lower', '-5.30', '-7.86', '-9.32'],
    ['90% CI Upper', '+10.40', '+8.08', '+6.48'],
    ['Equivalence Verdict', 'Inconclusive', 'EQUIVALENT', 'EQUIVALENT'],
]

for i, row_data in enumerate(data, start=1):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text

note = doc.add_paragraph('90% confidence intervals per Lakens et al. (2018) TOST convention. Equivalence bound of ±10pp represents smallest governance-relevant shift in board accountability given baseline turnover ~38%. "Equivalent" indicates both CI bounds fall within equivalence threshold. "Inconclusive" indicates 30-day window lacks precision to confirm negligibility at upper tail (+10.40pp vs. +10.00pp bound).')
note.style = 'Intense Quote'

doc.add_page_break()

# ============================================================================
# TABLE A5
# ============================================================================
doc.add_heading('TABLE A5: First-Stage Logistic and Mediation Model Coefficients', level=2)

doc.add_heading('Panel A: First Stage — FCC → Immediate Disclosure', level=3)

table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Coefficient'
hdr_cells[1].text = 'Estimate'

data_a = [
    ['FCC logit coefficient', '0.8313'],
    ['Standard error', '0.2352'],
    ['p-value', '< .001 ***'],
    ['Average marginal effect', '+14.52pp'],
]

for i, row_data in enumerate(data_a, start=1):
    row = table.rows[i]
    row.cells[0].text = row_data[0]
    row.cells[1].text = row_data[1]

doc.add_paragraph('N = 651 | Dependent Variable: immediate_disclosure = 1 if firm disclosed within 7 days')

doc.add_heading('Panel B: Mediation Model — Both Variables in One Specification', level=3)

table = doc.add_table(rows=7, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Variable'
hdr_cells[1].text = '30-day'
hdr_cells[2].text = '90-day'
hdr_cells[3].text = '180-day'

data_b = [
    ['FCC coefficient', '0.1704', '0.1164', '0.0352'],
    ['FCC SE', '0.2093', '0.2103', '0.2094'],
    ['FCC p-value', '.416', '.580', '.867'],
    ['Immediate Disclosure coef', '-0.6502', '-1.0085', '-0.9440'],
    ['Immediate Disclosure SE', '0.2117', '0.1976', '0.1965'],
    ['Immediate Disclosure p-value', '.002 **', '<.001 ***', '<.001 ***'],
]

for i, row_data in enumerate(data_b, start=1):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text

note = doc.add_paragraph('N = 651 | Standard Errors: Industry-clustered robust')
note.style = 'Intense Quote'

note2 = doc.add_paragraph('Immediate disclosure treated as mediator (not treatment). FCC coefficient in Panel B represents direct effect after accounting for disclosure timing. Negative immediate disclosure coefficient reflects compositional differences between rapid and delayed disclosers and should NOT be interpreted as causal.')
note2.style = 'Intense Quote'

doc.add_page_break()

# ============================================================================
# TABLE A6
# ============================================================================
doc.add_heading('TABLE A6: Bootstrap Estimates of Indirect Effects', level=2)

doc.add_paragraph('Method: Bias-corrected bootstrap, 1,000 iterations')

table = doc.add_table(rows=4, cols=5)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Temporal Window'
hdr_cells[1].text = 'Indirect Effect (pp)'
hdr_cells[2].text = '95% CI Lower'
hdr_cells[3].text = '95% CI Upper'
hdr_cells[4].text = 'Significant?'

data = [
    ['30-day', '2.69', '-6.87', '+12.93', 'No'],
    ['90-day', '0.25', '-8.99', '+10.96', 'No'],
    ['180-day', '-1.45', '-11.54', '+8.16', 'No'],
]

for i, row_data in enumerate(data, start=1):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text

note = doc.add_paragraph('Bootstrap confidence intervals spanning zero indicate no significant mediation. The indirect effect (FCC → immediate_disclosure → turnover) does not significantly explain the total FCC effect at any window.')
note.style = 'Intense Quote'

doc.add_page_break()

# ============================================================================
# TABLE A7
# ============================================================================
doc.add_heading('TABLE A7: Covariate Balance Test — FCC vs. Non-FCC Firms', level=2)

doc.add_paragraph('Sample: FCC n = 140, Non-FCC n = 511')

table = doc.add_table(rows=4, cols=6)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Covariate'
hdr_cells[1].text = 'FCC Mean'
hdr_cells[2].text = 'Non-FCC Mean'
hdr_cells[3].text = 'Difference'
hdr_cells[4].text = "Cohen's d"
hdr_cells[5].text = 'Balanced?'

data = [
    ['Log Market Cap', '11.29', '10.49', '+0.80', '0.65', 'NO (expected)'],
    ['Leverage (D/A)', '0.568', '0.561', '+0.007', '0.03', 'YES'],
    ['Return on Assets', '0.0108', '0.0145', '-0.0037', '-0.10', 'MARGINAL'],
]

for i, row_data in enumerate(data, start=1):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text

note = doc.add_paragraph('FCC firms significantly larger than non-FCC (d = 0.65), as expected. Leverage is well-balanced (p = .675). Marginal ROA imbalance is economically small and addressed through regression controls.')
note.style = 'Intense Quote'

doc.add_page_break()

# ============================================================================
# TABLE A8
# ============================================================================
doc.add_heading('TABLE A8: Placebo Tests — FCC Effect on Alternative Governance Outcomes', level=2)

doc.add_paragraph('Specification: Identical to Table A2, with alternative dependent variables substituted')

table = doc.add_table(rows=4, cols=5)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Placebo Outcome'
hdr_cells[1].text = 'FCC Coefficient'
hdr_cells[2].text = 'p-value'
hdr_cells[3].text = 'Significant?'
hdr_cells[4].text = 'Interpretation'

data = [
    ['Multiple executive changes (180d)', '-0.025', '.900', 'No', 'FCC independent of org churn'],
    ['Weak governance indicator†', '181.02', '.706', 'No', 'Singular matrix—not interpretable'],
    ['Governance weakness score', '-0.0000', '.052', 'No', 'Marginal; NS at α = .05'],
]

for i, row_data in enumerate(data, start=1):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text

note = doc.add_paragraph('All placebo tests are null (p > .05), confirming FCC effects are specific to CEO turnover, not general governance deterioration or organizational disruption.')
note.style = 'Intense Quote'

doc.add_page_break()

# ============================================================================
# TABLE A9
# ============================================================================
doc.add_heading('TABLE A9: Dose-Response Analysis — FCC × Breach Severity Interactions', level=2)

doc.add_paragraph('Specification: logit(P) = β₀ + β₁(fcc) + β₂(severity) + β₃(fcc × severity) + β₄(X) + ε')
doc.add_paragraph('N = 651 | All interactions estimated at 30-day window (primary specification)')

table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Interaction Term'
hdr_cells[1].text = 'Coefficient'
hdr_cells[2].text = 'p-value'
hdr_cells[3].text = 'Significant?'

data = [
    ['FCC × Financial Data Breach', '0.5314', '.234', 'No'],
    ['FCC × Health Data Breach', '22.8102†', '.998', 'No'],
    ['FCC × High Record Count (>100K)', '0.2469', '.870', 'No'],
    ['FCC × Ransomware Attack', '-0.6250', '.617', 'No'],
]

for i, row_data in enumerate(data, start=1):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text

note = doc.add_paragraph('All interactions null (p > .10), indicating FCC effect does not scale with breach severity. Uniformity rules out reputational harm amplification mechanism.')
note.style = 'Intense Quote'

note2 = doc.add_paragraph('† Near-perfect separation. Insufficient observations at the intersection of FCC-regulated status and health breach classification prevented stable estimation. Coefficient is not interpretable.')
note2.style = 'Intense Quote'

doc.add_page_break()

# ============================================================================
# TABLE A10
# ============================================================================
doc.add_heading('TABLE A10: Robustness to Alternative Model Specifications', level=2)

doc.add_heading('Panel A: OLS Linear Probability Model', level=3)

doc.add_paragraph('Outcome: Binary CEO turnover indicator; interpretation: percentage-point change')

table = doc.add_table(rows=4, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Coefficient'
hdr_cells[1].text = '30-day'
hdr_cells[2].text = '90-day'
hdr_cells[3].text = '180-day'

data = [
    ['FCC coefficient (pp)', '0.0241', '0.0015', '-0.0138'],
    ['Standard error', '0.0505', '0.0511', '0.0509'],
    ['p-value', '.632', '.977', '.786'],
]

for i, row_data in enumerate(data, start=1):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text

doc.add_paragraph('N = 651 | Standard Errors: Industry-clustered robust (HC3)')

doc.add_heading('Panel B: Cox Proportional Hazards Model', level=3)

doc.add_paragraph('Outcome: Time-to-CEO-turnover in days; right-censored at 180 days')

table = doc.add_table(rows=7, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Estimate'

data_cox = [
    ['N observations', '651'],
    ['N events', '385'],
    ['FCC hazard ratio', '1.127'],
    ['95% CI for HR', '[0.611, 2.080]'],
    ['p-value', '.702'],
    ['Proportional hazards assumption', 'Violated for FCC variable (p = .020); see note'],
    ['Schoenfeld residuals test p-value (FCC)', '0.020'],
]

table = doc.add_table(rows=8, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Estimate'

for i, row_data in enumerate(data_cox, start=1):
    row = table.rows[i]
    row.cells[0].text = row_data[0]
    row.cells[1].text = row_data[1]

doc.add_heading('Panel C: Negative Binomial Regression', level=3)

doc.add_paragraph('Outcome: Count of executive departures within 180 days')

table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Estimate'

data_nb = [
    ['N observations', '651'],
    ['FCC incidence rate ratio (IRR)', '1.220'],
    ['95% CI for IRR', '[1.004, 1.480]'],
    ['p-value', '.069'],
]

for i, row_data in enumerate(data_nb, start=1):
    row = table.rows[i]
    row.cells[0].text = row_data[0]
    row.cells[1].text = row_data[1]

note = doc.add_paragraph('The Schoenfeld residuals test indicates the proportional hazards assumption is violated for the FCC variable (p = .020). The primary logistic regression at fixed windows does not impose this assumption and remains the authoritative specification. The Cox model is reported for directional consistency only.')
note.style = 'Intense Quote'

note2 = doc.add_paragraph('All three specifications confirm null FCC effect. OLS and Cox are fully null. Negative binomial IRR = 1.220 is marginally non-significant (p = .069) and consistent with null governance response. 95% CI computed using Wald method.')
note2.style = 'Intense Quote'

doc.add_page_break()

# ============================================================================
# TABLE A11
# ============================================================================
doc.add_heading('TABLE A11: Firm-Size Heterogeneity Analysis', level=2)

doc.add_paragraph('Outcome: CEO turnover by firm size quartile (quarterly break points by log market cap)')
doc.add_paragraph('Specification: logit(P) = β₀ + β₁(fcc) + β₂(Q2) + β₃(Q3) + β₄(Q4) + β₅(controls) + ε')

table = doc.add_table(rows=5, cols=6)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Firm Size Quartile'
hdr_cells[1].text = 'Definition'
hdr_cells[2].text = '30-day'
hdr_cells[3].text = '90-day'
hdr_cells[4].text = '180-day'
hdr_cells[5].text = 'Status'

data = [
    ['Q1', 'Smallest (n=164)', '†', '†', '†', 'Singular matrix'],
    ['Q2', 'n=162', '†', '†', '†', 'Singular matrix'],
    ['Q3', 'Mid-large (n=163)', '-17.54pp (p=.08)', '-11.49pp (p=.22)', '-17.61pp (p=.06)', 'Exploratory'],
    ['Q4', 'Largest (n=162)', 'Reference', 'Reference', 'Reference', 'Reference'],
]

for i, row_data in enumerate(data, start=1):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text

note = doc.add_paragraph('† Singular matrix. Logit model did not converge due to absence of within-cell outcome variation in FCC-regulated observations within this quartile. Estimates are not interpretable.')
note.style = 'Intense Quote'

note2 = doc.add_paragraph('Q3 pattern is exploratory and should NOT be interpreted as a causal finding. Thin within-cell variation limits statistical inference. The aggregate null in Table A2 is the primary result.')
note2.style = 'Intense Quote'

# Save document
doc.save('ESSAY3_H6_APPENDIX.docx')

print("\n[OK] Complete appendix created: ESSAY3_H6_APPENDIX.docx")
print("Tables A1-A11 with all canonical data")
print("=" * 90)
