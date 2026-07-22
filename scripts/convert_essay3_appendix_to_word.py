"""Convert Essay 3 H6 appendix markdown to formatted Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_table_from_markdown(doc, markdown_table):
    """Parse markdown table and add to document."""
    lines = markdown_table.strip().split('\n')
    if len(lines) < 2:
        return

    # Parse header
    header_line = lines[0]
    separator_line = lines[1]

    # Extract columns
    headers = [h.strip() for h in header_line.split('|')[1:-1]]

    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Format header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)

    # Add data rows
    for line in lines[2:]:
        if line.strip() and '|' in line:
            cells_data = [c.strip() for c in line.split('|')[1:-1]]
            if len(cells_data) == len(headers):
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(cells_data):
                    row_cells[i].text = cell_data
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_paragraph()
title_run = title.add_run('ESSAY 3 APPENDIX: EXECUTIVE TURNOVER AND GOVERNANCE RESPONSE')
title_run.font.size = Pt(14)
title_run.font.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('Mandatory Disclosure Requirements and CEO Accountability')
subtitle_run.font.size = Pt(12)
subtitle_run.font.italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Blank line

# Main content sections
sections_data = [
    {
        'title': 'APPENDIX A: MAIN RESULTS & SPECIFICATION TESTS',
        'subsections': [
            ('TABLE A1: Sample Characteristics and Turnover Rates (N=651 Mediation Analysis Sample)',
             '''| Characteristic | N | Turnover 30d | Turnover 90d | Turnover 180d | % of Total |
|---|---|---|---|---|---|
| **Full Sample** | 651 | 247 (37.9%) | 379 (58.2%) | 385 (59.1%) | 100.0% |
| **By FCC Status:** | | | | | |
| FCC-Regulated | 140 | 52 (37.1%) | 84 (60.0%) | 85 (60.7%) | 21.5% |
| Non-FCC Firms | 511 | 195 (38.2%) | 295 (57.7%) | 300 (58.7%) | 78.5% |
| Difference (FCC - Non-FCC) | — | -1.1pp | +2.3pp | +2.0pp | — |''',
             'Turnover rates are consistent across regulatory conditions (FCC vs. non-FCC difference: 1-2 percentage points across all windows), a pattern that foreshadows the primary null finding. Baseline turnover is driven by breach events themselves, not regulatory status.'),
        ]
    }
]

# Add content
for section_data in sections_data:
    # Section header
    section_heading = doc.add_heading(section_data['title'], level=1)
    section_heading.runs[0].font.size = Pt(13)

    for table_title, table_md, notes in section_data['subsections']:
        # Table heading
        table_heading = doc.add_heading(table_title, level=2)
        table_heading.runs[0].font.size = Pt(11)

        # Add table
        add_table_from_markdown(doc, table_md)

        # Notes
        notes_para = doc.add_paragraph()
        notes_para.paragraph_format.left_indent = Inches(0)
        notes_run = notes_para.add_run('Notes: ')
        notes_run.font.bold = True
        notes_para.add_run(notes)
        for run in notes_para.runs:
            run.font.size = Pt(10)

        doc.add_paragraph()  # Blank line

# Save
output_path = 'ESSAY3_H6_APPENDIX.docx'
doc.save(output_path)
print(f'[OK] Converted Essay 3 appendix to Word format: {output_path}')
print('     Note: This is a template. Open in Word and format tables as needed.')
print('     The markdown file (ESSAY3_H6_APPENDIX.md) has complete data for copying into Word.')
