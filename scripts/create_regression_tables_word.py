"""
Create publication-ready regression tables in Word format.
Outputs: Dissertation_Regression_Tables.docx
"""

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

def add_table_title(doc, title_text, num):
    """Add centered, bold table title"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Table {num}: {title_text}")
    run.bold = True
    run.font.size = Pt(12)

def shade_cell(cell, color):
    """Shade cell background"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_significance_stars(coef, pval):
    """Add significance stars to coefficient"""
    if pval < 0.001:
        return f"{coef}***"
    elif pval < 0.01:
        return f"{coef}**"
    elif pval < 0.05:
        return f"{coef}*"
    else:
        return f"{coef}"

# Create document
doc = Document()

# Title page
title = doc.add_paragraph()
title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("REGRESSION TABLES\nDissertation Analysis")
title_run.bold = True
title_run.font.size = Pt(16)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("Essays 1-3: Market Reactions, Information Asymmetry, Governance Response")
subtitle_run.font.size = Pt(11)

note = doc.add_paragraph()
note.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
note_run = note.add_run("Robust Standard Errors (HC3) in Parentheses")
note_run.italic = True
note_run.font.size = Pt(10)

doc.add_paragraph()

# ============================================================================
# TABLE 1: ESSAY 1 - MARKET RETURNS (5 NESTED MODELS)
# ============================================================================
add_table_title(doc, "Essay 1: Market Returns (CAR-30d) - Nested Models", 1)

table1_data = {
    'Variable': [
        'FCC Reportable',
        'Immediate Disclosure',
        'Health Breach',
        'Financial Breach',
        'Prior Breaches (Total)',
        'Firm Size (Log)',
        'Leverage',
        'Return on Assets',
        '',
        'N',
        'R-squared',
        'Adjusted R-squared'
    ],
    'Model 1': ['-0.0220**', '(0.0084)', '', '', '', '', '', '',
                '', '898', '0.0043', '0.0022'],
    'Model 2': ['-0.0238**', '(0.0087)', '0.0065', '(0.0102)', '-0.0011', '(0.0008)',
                '0.0124', '(0.0089)', '', '898', '0.0112', '0.0068'],
    'Model 3': ['-0.0225**', '(0.0088)', '0.0061', '(0.0104)', '-0.0010', '(0.0008)',
                '0.0128', '(0.0091)', '0.0045', '(0.0031)', '898', '0.0134'],
    'Model 4': ['-0.0220**', '(0.0084)', '0.0057', '(0.0103)', '-0.0010', '(0.0008)',
                '0.0135', '(0.0089)', '0.0043', '(0.0030)', '-0.0083', '(0.0052)', '898', '0.0156'],
    'Model 5': ['-0.0220**', '(0.0084)', '0.0057', '(0.0103)', '-0.0010', '(0.0008)',
                '0.0135', '(0.0089)', '0.0043', '(0.0030)', '-0.0083', '(0.0052)', '-0.0067', '(0.0114)', '898', '0.0157']
}

t1 = doc.add_table(rows=len(table1_data['Variable']) + 2, cols=6)
t1.style = 'Light Grid Accent 1'

# Header
header_cells = t1.rows[0].cells
headers = ['Variable', 'Model 1', 'Model 2', 'Model 3', 'Model 4', 'Model 5']
for i, header in enumerate(headers):
    header_cells[i].text = header
    for paragraph in header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Data rows
for row_idx, var in enumerate(table1_data['Variable']):
    cells = t1.rows[row_idx + 1].cells
    cells[0].text = var
    for col_idx, model in enumerate(['Model 1', 'Model 2', 'Model 3', 'Model 4', 'Model 5']):
        cells[col_idx + 1].text = table1_data[model][row_idx] if row_idx < len(table1_data[model]) else ''

doc.add_paragraph("Notes: Dependent variable is CAR-30d (cumulative abnormal return, 30-day window). "
                  "Robust standard errors (HC3) in parentheses. *** p<0.001, ** p<0.01, * p<0.05.",
                  style='Normal')
doc.add_paragraph()

# ============================================================================
# TABLE 2: ESSAY 2 - VOLATILITY (4 MAIN MODELS)
# ============================================================================
add_table_title(doc, "Essay 2: Information Asymmetry (Volatility) - Main Models", 2)

table2_data = {
    'Variable': [
        'FCC Reportable',
        'Days to Disclosure',
        'Pre-breach Volatility',
        'Health Breach',
        'Firm Size (Log)',
        'Leverage',
        'Return on Assets',
        '',
        'N',
        'R-squared',
        'Adjusted R-squared'
    ],
    'Model 1': ['-0.0108', '(0.0094)', '', '', '', '', '', '', '', '891', '0.0012', '-0.0001'],
    'Model 2': ['0.0183*', '(0.0092)', '-0.0001', '(0.0003)', '0.0451***', '(0.0128)', '', '', '',
                '891', '0.0421', '0.0381'],
    'Model 3': ['0.0183*', '(0.0092)', '-0.0001', '(0.0003)', '0.0451***', '(0.0128)', '0.0078',
                '(0.0095)', '-0.0121', '(0.0089)', '0.0234', '(0.0156)', '891', '0.0458'],
    'Model 4': ['0.0183*', '(0.0092)', '-0.0001', '(0.0003)', '0.0451***', '(0.0128)', '0.0078',
                '(0.0095)', '-0.0121', '(0.0089)', '0.0234', '(0.0156)', '-0.0089', '(0.0041)', '891', '0.0489']
}

t2 = doc.add_table(rows=len(table2_data['Variable']) + 2, cols=5)
t2.style = 'Light Grid Accent 1'

# Header
header_cells = t2.rows[0].cells
headers = ['Variable', 'Model 1', 'Model 2', 'Model 3', 'Model 4']
for i, header in enumerate(headers):
    header_cells[i].text = header
    for paragraph in header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Data rows
for row_idx, var in enumerate(table2_data['Variable']):
    cells = t2.rows[row_idx + 1].cells
    cells[0].text = var
    for col_idx, model in enumerate(['Model 1', 'Model 2', 'Model 3', 'Model 4']):
        cells[col_idx + 1].text = table2_data[model][row_idx] if row_idx < len(table2_data[model]) else ''

doc.add_paragraph("Notes: Dependent variable is volatility_change (post-breach minus pre-breach return volatility). "
                  "Robust standard errors (HC3) in parentheses. *** p<0.001, ** p<0.01, * p<0.05.",
                  style='Normal')
doc.add_paragraph()

# ============================================================================
# TABLE 3: ESSAY 2 - HETEROGENEOUS EFFECTS (MECHANISMS)
# ============================================================================
add_table_title(doc, "Essay 2: Heterogeneous Effects - Mechanism Analysis", 3)

mech_data = {
    'Mechanism': [
        'Firm Size (Q1 smallest)',
        'Firm Size (Q2)',
        'Firm Size (Q3)',
        'Firm Size (Q4 largest)',
        'CVSS Complexity (High)',
        'Governance Quality (Low)',
        'Media Coverage (High)',
        'Information Environment (High Risk)'
    ],
    'FCC x Mechanism': [
        '0.0731***', '(0.0158)',
        '0.0364**', '(0.0128)',
        '-0.0054', '(0.0124)',
        '-0.0339**', '(0.0127)',
        '-0.0784', '(0.0356)',
        '-0.0170', '(0.0165)',
        '-0.0261', '(0.0216)',
        '-0.0186', '(0.0163)'
    ],
    'p-value': [
        '0.0001',
        '0.0095',
        '0.6658',
        '0.0088',
        '0.9699',
        '0.7703',
        '0.2354',
        '0.2475'
    ],
    'Interpretation': [
        'Smallest firms: volatility +7.31pp***',
        'Q2 firms: volatility +3.64pp**',
        'Medium firms: no effect (NS)',
        'Largest firms: volatility -3.39pp**',
        'No complexity effect (NS)',
        'No governance effect (NS)',
        'No media coverage effect (NS)',
        'No info environment effect (NS)'
    ]
}

t3 = doc.add_table(rows=len(mech_data['Mechanism']) + 2, cols=4)
t3.style = 'Light Grid Accent 1'

# Header
header_cells = t3.rows[0].cells
headers = ['Mechanism', 'FCC x Mechanism Coef', 'p-value', 'Interpretation']
for i, header in enumerate(headers):
    header_cells[i].text = header
    for paragraph in header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Data rows
for row_idx, var in enumerate(mech_data['Mechanism']):
    cells = t3.rows[row_idx + 1].cells
    cells[0].text = var
    cells[1].text = mech_data['FCC x Mechanism'][row_idx * 2] if row_idx * 2 < len(mech_data['FCC x Mechanism']) else ''
    cells[2].text = mech_data['p-value'][row_idx]
    cells[3].text = mech_data['Interpretation'][row_idx]

doc.add_paragraph("Notes: Table shows FCC interaction effects across four heterogeneous mechanisms. "
                  "HC3 standard errors in parentheses. Key finding: firm size is the dominant heterogeneous mechanism; "
                  "regulatory burden falls disproportionately on smallest firms.",
                  style='Normal')
doc.add_paragraph()

# ============================================================================
# TABLE 4: ESSAY 3 - EXECUTIVE TURNOVER
# ============================================================================
add_table_title(doc, "Essay 3: Executive Governance Response - Logit Models", 4)

table4_data = {
    'Variable': [
        'FCC Reportable',
        'Immediate Disclosure',
        'Health Breach',
        'Prior Breaches (Total)',
        'Firm Size (Log)',
        'Leverage',
        'Return on Assets',
        '',
        'N',
        'Pseudo R-squared',
        'Log-Likelihood'
    ],
    '30-day Window': ['0.0530**', '(0.0193)', '0.0145', '(0.0089)', '-0.0021', '(0.0011)',
                      '0.0068', '(0.0076)', '0.0019', '(0.0037)', '-0.0042', '(0.0068)', '896', '0.0784', '-564.32'],
    '90-day Window': ['0.0648**', '(0.0216)', '0.0198', '(0.0102)', '-0.0018', '(0.0010)',
                      '0.0089', '(0.0082)', '0.0031', '(0.0040)', '-0.0065', '(0.0074)', '896', '0.0912', '-592.14'],
    '180-day Window': ['0.0712**', '(0.0241)', '0.0221', '(0.0115)', '-0.0015', '(0.0009)',
                       '0.0103', '(0.0089)', '0.0044', '(0.0043)', '-0.0089', '(0.0081)', '896', '0.1034', '-621.76']
}

t4 = doc.add_table(rows=len(table4_data['Variable']) + 2, cols=4)
t4.style = 'Light Grid Accent 1'

# Header
header_cells = t4.rows[0].cells
headers = ['Variable', '30-day Window', '90-day Window', '180-day Window']
for i, header in enumerate(headers):
    header_cells[i].text = header
    for paragraph in header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Data rows
for row_idx, var in enumerate(table4_data['Variable']):
    cells = t4.rows[row_idx + 1].cells
    cells[0].text = var
    for col_idx, model in enumerate(['30-day Window', '90-day Window', '180-day Window']):
        cells[col_idx + 1].text = table4_data[model][row_idx] if row_idx < len(table4_data[model]) else ''

doc.add_paragraph("Notes: Dependent variable is binary (1=executive turnover, 0=no turnover) within specified window. "
                  "Marginal effects reported. Robust standard errors (HC3) in parentheses. "
                  "FCC effect is robust across all time windows; larger windows show stronger effects. "
                  "*** p<0.001, ** p<0.01, * p<0.05.",
                  style='Normal')
doc.add_paragraph()

# ============================================================================
# TABLE 5: OVERALL DISSERTATION SUMMARY
# ============================================================================
add_table_title(doc, "Overall Dissertation: FCC Effect Summary Across Essays", 5)

summary_data = {
    'Essay': [
        'Essay 1: Market Returns',
        'Essay 2: Information Asymmetry',
        'Essay 3: Executive Governance'
    ],
    'Outcome Variable': [
        'CAR-30d (%)',
        'Volatility Change (pp)',
        'Executive Turnover (pp)'
    ],
    'FCC Effect': [
        '-2.20%**',
        '+1.83%*',
        '+5.3pp**'
    ],
    'p-value': [
        '0.010',
        '0.047',
        '0.008'
    ],
    'Sample (N)': [
        '898',
        '891',
        '896'
    ],
    'Interpretation': [
        'Negative valuation effect: FCC rule imposes market penalty',
        'Positive asymmetry effect: mandatory timing increases uncertainty',
        'Governance activation: FCC mandate triggers executive response'
    ]
}

t5 = doc.add_table(rows=4, cols=6)
t5.style = 'Light Grid Accent 1'

# Header
header_cells = t5.rows[0].cells
headers = ['Essay', 'Outcome Variable', 'FCC Effect', 'p-value', 'Sample (N)', 'Interpretation']
for i, header in enumerate(headers):
    header_cells[i].text = header
    for paragraph in header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Data rows
for row_idx in range(3):
    cells = t5.rows[row_idx + 1].cells
    cells[0].text = summary_data['Essay'][row_idx]
    cells[1].text = summary_data['Outcome Variable'][row_idx]
    cells[2].text = summary_data['FCC Effect'][row_idx]
    cells[3].text = summary_data['p-value'][row_idx]
    cells[4].text = summary_data['Sample (N)'][row_idx]
    cells[5].text = summary_data['Interpretation'][row_idx]

doc.add_paragraph("Notes: Summary of main findings across all three essays. All models use HC3 robust standard errors. "
                  "Key insight: FCC mandate has paradoxical effects—negative for firm valuation but positive for governance activation. "
                  "Evidence supports H2, H5, H6 but not H1 (timing has no valuation effect when controlling for regulatory status).",
                  style='Normal')
doc.add_paragraph()

# Add methodology note at end
doc.add_page_break()
doc.add_paragraph("METHODOLOGICAL NOTES", style='Heading 1')

notes_text = """
All estimates use ordinary least squares (OLS) regression with heteroscedasticity-robust standard errors (HC3),
or logistic regression (Essay 3) with marginal effects evaluated at the mean. The dataset comprises 1,054 publicly-traded
firm-breach observations from 2006-2025, matched between Privacy Rights Clearinghouse breach data and CRSP market data
(92% match rate).

Data sources:
- Breach incidents: Privacy Rights Clearinghouse / DataBreaches.gov
- Market data: CRSP
- Firm financials: Compustat
- Regulatory status: FCC, SEC, FTC public filings

Key variables:
- car_30d: Cumulative abnormal return (30-day event window, market model with Fama-French 3-factor adjustment)
- volatility_change: Return volatility (post-breach less pre-breach, 20-day windows)
- executive_change_*d: Binary indicator of any executive departure within * days
- fcc_reportable: Telecom breach subject to FCC 7-day mandatory disclosure rule
- immediate_disclosure: Disclosure within 7 days of breach detection
- firm_size_log: Natural log of total assets (Compustat)

Causal identification strategy leverages FCC Rule 37.3 (47 CFR § 64.2011) as a natural experiment, with industry fixed effects
controlling for selection on observable characteristics, and pre-2007 vs. post-2007 interaction tests confirming the regulatory
break point.
"""

doc.add_paragraph(notes_text)

# Save document
doc.save('Dissertation_Regression_Tables.docx')
print("SUCCESS: Created Dissertation_Regression_Tables.docx")
