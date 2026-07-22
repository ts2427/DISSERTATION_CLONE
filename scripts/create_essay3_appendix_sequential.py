"""Create Essay 3 H6 Appendix with SEQUENTIAL A1-A9 numbering (matching Results section order)."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_appendix():
    """Create Essay 3 H6 appendix Word document with sequential A1-A9 numbering."""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('ESSAY 3 APPENDIX: EXECUTIVE TURNOVER AND GOVERNANCE RESPONSE')
    title_run.font.size = Pt(14)
    title_run.font.bold = True

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Mandatory Disclosure Requirements and CEO Accountability')
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True

    doc.add_paragraph()  # Blank line

    # TABLE A1: Sample characteristics
    doc.add_heading('TABLE A1: Sample Characteristics and Turnover Rates (N=651)', level=2)

    table = doc.add_table(rows=15, cols=5)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    headers = ['Characteristic', 'N', 'Turnover 30d', 'Turnover 90d', 'Turnover 180d']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    data = [
        ['Full Sample', '651', '247 (37.9%)', '379 (58.2%)', '385 (59.1%)'],
        ['By FCC Status:', '', '', '', ''],
        ['FCC-Regulated', '140', '52 (37.1%)', '84 (60.0%)', '85 (60.7%)'],
        ['Non-FCC Firms', '511', '195 (38.2%)', '295 (57.7%)', '300 (58.7%)'],
        ['Difference (FCC - Non-FCC)', '—', '-1.1pp', '+2.3pp', '+2.0pp'],
        ['By Disclosure Timing:', '', '', '', ''],
        ['Immediate Disclosure (≤7d)', '160', '54 (33.8%)', '97 (60.6%)', '100 (62.5%)'],
        ['Delayed Disclosure (>7d)', '491', '193 (39.3%)', '282 (57.4%)', '285 (58.0%)'],
        ['Difference (Immediate - Delayed)', '—', '-5.5pp', '+3.2pp', '+4.5pp'],
        ['By Firm Size Quartile:', '', '', '', ''],
        ['Q1 (Smallest)', '164', '69 (42.1%)', '125 (76.2%)', '125 (76.2%)'],
        ['Q2', '162', '33 (20.4%)', '94 (58.0%)', '96 (59.3%)'],
        ['Q3', '163', '87 (53.4%)', '109 (66.9%)', '114 (69.9%)'],
        ['Q4 (Largest)', '162', '58 (35.8%)', '89 (54.9%)', '90 (55.6%)'],
    ]

    for i, row_data in enumerate(data):
        row_cells = table.rows[i+1].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if row_data[0] in ['By FCC Status:', 'By Disclosure Timing:', 'By Firm Size Quartile:']:
                for run in row_cells[0].paragraphs[0].runs:
                    run.font.bold = True

    notes = doc.add_paragraph()
    notes_run = notes.add_run('Notes: ')
    notes_run.font.bold = True
    notes.add_run('Turnover rates are consistent across regulatory conditions (FCC vs. non-FCC difference: 1-2 percentage points across all windows), foreshadowing the primary null finding.')
    for run in notes.runs:
        run.font.size = Pt(10)

    doc.add_paragraph()

    # TABLE A2: Main H6 regression
    doc.add_heading('TABLE A2: Main H6 Regression - FCC Effect on Executive Turnover (N=651)', level=2)

    table = doc.add_table(rows=4, cols=8)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    headers = ['Window', 'FCC Coef', 'Std. Error', 'P-Value', 'AME (pp)', '95% CI', 'Pseudo R²', 'Sig']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    data = [
        ['30 days', '0.110', '0.207', '0.595', '2.55', '[-4.39, 9.49]', '0.0164', 'NS'],
        ['90 days', '0.005', '0.205', '0.982', '0.11', '[-4.04, 4.26]', '0.0219', 'NS'],
        ['180 days', '-0.061', '0.205', '0.768', '-1.42', '[-5.46, 2.63]', '0.0226', 'NS'],
    ]

    for row_data in data:
        row_cells = table.add_row().cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    notes = doc.add_paragraph()
    notes_run = notes.add_run('Notes: ')
    notes_run.font.bold = True
    notes.add_run('All FCC effects non-significant (p > 0.59) across all windows. MDE at 80% power = 13.36-13.57pp. Observed effects well below detection threshold. Logistic regression with controls (health_breach, prior_breaches_total, firm_size_log, leverage, roa). Standard errors HC3.')
    for run in notes.runs:
        run.font.size = Pt(10)

    doc.add_paragraph()

    # TABLE A3: TOST equivalence
    doc.add_heading('TABLE A3: TOST Equivalence Testing (N=651, ±10pp Bounds)', level=2)

    table = doc.add_table(rows=4, cols=7)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    headers = ['Window', 'Point Est. (pp)', 'Std. Error', '95% CI', 'Lower Test', 'Upper Test', 'Status']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    data = [
        ['30d', '2.55', '4.769', '[-6.80, 11.90]', 'PASS', 'FAIL', 'Inconclusive'],
        ['90d', '0.11', '4.843', '[-9.38, 9.60]', 'PASS', 'PASS', 'Equivalent'],
        ['180d', '-1.42', '4.804', '[-10.84, 8.00]', 'FAIL', 'PASS', 'Inconclusive'],
    ]

    for row_data in data:
        row_cells = table.add_row().cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    notes = doc.add_paragraph()
    notes_run = notes.add_run('Notes: ')
    notes_run.font.bold = True
    notes.add_run('At 90d, CI falls entirely within equivalence bounds, confirming effect is economically negligible. At 30d/180d, inconclusive due to precision limits. Equivalence bound ±10pp is more conservative than 80% power MDE.')
    for run in notes.runs:
        run.font.size = Pt(10)

    doc.add_paragraph()

    # TABLE A4: Mediation first stage & reduced form
    doc.add_heading('TABLE A4: Mediation Analysis - First Stage and Reduced Form (N=651)', level=2)

    doc.add_paragraph('A4A: FIRST STAGE (a-path: FCC → Immediate Disclosure)').runs[0].font.bold = True

    table = doc.add_table(rows=2, cols=6)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    headers = ['Coefficient', 'Value', 'Std. Error', 'P-Value', 'AME (pp)', 'Pseudo R²']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    row_cells = table.add_row().cells
    data = ['FCC (logit)', '0.8313', '0.2352', '< 0.001***', '14.52', '0.0887']
    for j, cell_data in enumerate(data):
        row_cells[j].text = cell_data
        for paragraph in row_cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

    interp = doc.add_paragraph()
    interp_run = interp.add_run('Interpretation: ')
    interp_run.font.bold = True
    interp.add_run('Strong first stage. FCC regulation increases probability of immediate disclosure by 14.52 percentage points.')
    for run in interp.runs:
        run.font.size = Pt(10)

    doc.add_paragraph()

    doc.add_paragraph('A4B: REDUCED FORM (c-path: FCC → Turnover Total Effect)').runs[0].font.bold = True

    table = doc.add_table(rows=4, cols=6)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    headers = ['Window', 'FCC Coef', 'Std. Error', 'P-Value', 'AME (pp)', '95% CI']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    data = [
        ['30d', '0.1101', '0.2072', '0.5952', '2.55', '[-4.39, 9.49]'],
        ['90d', '0.0045', '0.2050', '0.9823', '0.11', '[-4.04, 4.26]'],
        ['180d', '-0.0605', '0.2049', '0.7677', '-1.42', '[-5.46, 2.63]'],
    ]

    for row_data in data:
        row_cells = table.add_row().cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    interp = doc.add_paragraph()
    interp_run = interp.add_run('Interpretation: ')
    interp_run.font.bold = True
    interp.add_run('Total FCC effect on turnover is null across all windows.')
    for run in interp.runs:
        run.font.size = Pt(10)

    doc.add_paragraph()

    # TABLE A5: Bootstrap indirect effects (NEWLY ADDED)
    doc.add_heading('TABLE A5: Bootstrap Indirect Effects - Mediation Pathway (1,000 iterations, N=651)', level=2)

    table = doc.add_table(rows=4, cols=5)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    headers = ['Window', 'Indirect Effect (pp)', '95% CI Lower', '95% CI Upper', 'Significant?']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    data = [
        ['30d', '2.27', '-7.31', '12.74', 'NO (CI crosses zero)'],
        ['90d', '0.34', '-9.78', '9.76', 'NO (CI crosses zero)'],
        ['180d', '-1.24', '-11.18', '8.48', 'NO (CI crosses zero)'],
    ]

    for row_data in data:
        row_cells = table.add_row().cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    notes = doc.add_paragraph()
    notes_run = notes.add_run('Notes: ')
    notes_run.font.bold = True
    notes.add_run('Mediation pathway is null across all windows. All confidence intervals include zero. Disclosure timing does not transmit FCC pressure into governance response. The negative b-path (immediate disclosure → lower turnover) reflects selection bias, not causation.')
    for run in notes.runs:
        run.font.size = Pt(10)

    doc.add_paragraph()

    # TABLE A6: Covariate balance (relabeled from C1)
    doc.add_heading('TABLE A6: Covariate Balance Test (FCC vs. Non-FCC Firms, N=651)', level=2)

    table = doc.add_table(rows=4, cols=6)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    headers = ['Variable', 'FCC Mean', 'Non-FCC Mean', 'Difference', "Cohen's d", 'P-Value']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    data = [
        ['firm_size_log', '11.29', '10.49', '+0.80', '0.65', '< 0.001***'],
        ['leverage', '0.568', '0.561', '+0.007', '0.03', '0.675'],
        ['roa', '0.0108', '0.0145', '-0.0037', '-0.10', '0.045*'],
    ]

    for row_data in data:
        row_cells = table.add_row().cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    notes = doc.add_paragraph()
    notes_run = notes.add_run('Notes: ')
    notes_run.font.bold = True
    notes.add_run('FCC firms larger (expected), controlled via regression. Leverage well-balanced. ROA marginal imbalance addressed via regression. Overall: adequate balance for quasi-experimental design.')
    for run in notes.runs:
        run.font.size = Pt(10)

    doc.add_paragraph()

    # TABLE A7: Placebo tests (relabeled from C2)
    doc.add_heading('TABLE A7: Placebo Tests - Alternative Governance Outcomes (N=651)', level=2)

    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    headers = ['Outcome Variable', 'FCC Coefficient', 'P-Value', 'Interpretation']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    data = [
        ['Multiple executive changes (180d)', '-0.0252', '0.9000', 'Not significant'],
        ['Weak governance indicator†', '181.02†', '0.7061†', 'Singular matrix'],
        ['Governance weakness score', '-0.0004', '0.0520', 'Marginal only'],
    ]

    for row_data in data:
        row_cells = table.add_row().cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    notes = doc.add_paragraph()
    notes_run = notes.add_run('Notes: ')
    notes_run.font.bold = True
    notes.add_run('All placebo outcomes null or marginal (p ≥ 0.052). Confirms analysis targets CEO turnover specifically, not general governance disruption. ')
    notes_run2 = notes.add_run('†Singular matrix. Logit model failed to converge due to absence of within-cell outcome variation in weak-governance indicator for FCC-regulated observations. Coefficient 181.02 is not a meaningful estimate; included for full transparency on model estimation.')
    for run in notes.runs:
        run.font.size = Pt(10)

    doc.add_paragraph()

    # TABLE A8: Dose-response (relabeled from C3)
    doc.add_heading('TABLE A8: Dose-Response Analysis - FCC × Breach Severity Interactions (N=651)', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    headers = ['Interaction Term', 'Coefficient', 'P-Value']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    data = [
        ['FCC × Financial Breach', '0.2148', '0.2336'],
        ['FCC × Health Breach', '0.0082', '0.9975'],
        ['FCC × High Record Count', '-0.1208', '0.8704'],
        ['FCC × Ransomware', '-0.0891', '0.6166'],
    ]

    for row_data in data:
        row_cells = table.add_row().cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    notes = doc.add_paragraph()
    notes_run = notes.add_run('Notes: ')
    notes_run.font.bold = True
    notes.add_run('All interactions null (p > 0.10). FCC effect does NOT scale with breach severity. Inconsistent with reputational-harm mechanism. Consistent with null primary finding.')
    for run in notes.runs:
        run.font.size = Pt(10)

    doc.add_paragraph()

    # TABLE A9: Firm-size heterogeneity (relabeled from D1)
    doc.add_heading('TABLE A9: Exploratory - Firm-Size Heterogeneity (30-Day Window, N=651)', level=2)

    table = doc.add_table(rows=5, cols=7)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    headers = ['Quartile', 'N', 'Baseline', 'FCC Coef', 'Std. Error', 'P-Value', 'AME (pp)']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    data = [
        ['Q1 (Smallest)', '164', '42.1%', '-0.1696', '0.5173', '0.7431', '-3.90'],
        ['Q2†', '162', '20.4%', '-42.12†', '655040085†', '1.0000†', '-22.07†'],
        ['Q3 (Mid-High)', '163', '53.4%', '-0.8242', '0.4712', '0.0802*', '-17.54'],
        ['Q4 (Largest)', '162', '35.8%', '0.2644', '0.3781', '0.4844', '+5.63'],
    ]

    for row_data in data:
        row_cells = table.add_row().cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    notes = doc.add_paragraph()
    notes_run = notes.add_run('Notes: ')
    notes_run.font.bold = True
    notes.add_run('†Singular matrix. Logit model failed to converge due to absence of within-cell outcome variation in FCC-regulated observations across all windows. Estimate not interpretable. Q3 shows consistent negative pattern across all windows (-17.5 to -11.5pp), marginal at 30d (p=0.080). This pattern is exploratory; larger sample needed for definitive inference. Aggregate null result is robust across identifiable quartiles.')
    for run in notes.runs:
        run.font.size = Pt(10)

    # Save
    output_path = 'ESSAY3_H6_APPENDIX.docx'
    doc.save(output_path)

    print(f'\n[OK] Created Word document: {output_path}')
    print(f'     Formatted with SEQUENTIAL A1-A9 numbering')
    print(f'     Matches Results section table references in order')
    print(f'     Includes all 9 tables: A1-A9')
    print(f'     Added missing Table A5 (Bootstrap indirect effects)')

if __name__ == '__main__':
    create_appendix()
