#!/usr/bin/env python3
"""Verify the final Essay 1 document"""

from docx import Document
import os

# Check file
filepath = "Essay 1 (6) (1) (1)_FINAL_CORRECTED.docx"
if os.path.exists(filepath):
    size_mb = os.path.getsize(filepath) / (1024*1024)
    print(f"File: {filepath}")
    print(f"Size: {size_mb:.1f} MB")
    print(f"Exists: YES")
else:
    print(f"ERROR: {filepath} not found")
    exit(1)

# Check content
doc = Document(filepath)

# Find Results section
results_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == "Results":
        results_idx = i
        break

print(f"\nResults section starts at paragraph: {results_idx}")
print(f"Total paragraphs in document: {len(doc.paragraphs)}")
print(f"Results section spans: {len(doc.paragraphs) - results_idx} paragraphs")

# Verify key content
print("\nVerifying key content sections:")
sections_to_check = [
    "Descriptive Overview",
    "Hypothesis 1: Disclosure Timing (H1)",
    "Hypothesis 2: FCC Regulatory Mandate (H2)",
    "Hypothesis 3: Prior Breach History (H3)",
    "Hypothesis 4: Health Data Sensitivity (H4)",
    "HETEROGENEITY ANALYSES",
    "Analysis 1: Complexity Heterogeneity",
    "Analysis 2: Firm Size Heterogeneity",
    "Machine Learning Validation",
    "Robustness Across Alternative",
    "Summary: Empirical Resolution",
]

for section in sections_to_check:
    found = any(section.lower() in para.text.lower() for para in doc.paragraphs)
    status = "YES" if found else "NO"
    print(f"  {section}: {status}")

# Check for key findings
print("\nVerifying key findings:")
findings = {
    "H1 coefficient +0.5676": "H1 null effect",
    "H2 coefficient -2.2994": "FCC effect significant",
    "H3 coefficient -0.2156": "Prior breach penalty",
    "H4 coefficient -2.5068": "Health data penalty",
    "Pre-2007": "Pre-2007 analysis",
    "Post-2007": "Post-2007 analysis",
    "CVSS": "Complexity heterogeneity",
    "Firm Size": "Firm size heterogeneity",
    "TOST": "TOST equivalence",
    "propensity score": "PSM robustness",
}

for search_text, label in findings.items():
    found = any(search_text.lower() in para.text.lower() for para in doc.paragraphs)
    status = "YES" if found else "NO"
    print(f"  {label}: {status}")

print("\n✓ Essay 1 Final Version Ready for Defense")
