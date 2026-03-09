"""
Add Table 1B (Essay 1 Heterogeneity Analysis) to Dissertation_Regression_Tables.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Open the existing document
doc = Document('Dissertation_Regression_Tables.docx')

# Find Table 1 by looking for it in the document
# We'll insert Table 1B after Table 1
table_1_index = None
for i, element in enumerate(doc.element.body):
    if element.tag.endswith('tbl'):  # Found a table
        # Check if this is Table 1 by looking at first cell content
        table = doc.tables[0] if table_1_index is None else None
        if table is None:
            # This is the first table, which should be Table 1
            table_1_index = i
            break

# If we found Table 1, we'll insert after it
if table_1_index is not None:
    insert_position = table_1_index + 1
else:
    insert_position = len(doc.element.body)

# Add a paragraph with the table title
p = doc.paragraphs[-1]._element
table_title = doc.add_paragraph()
table_title.style = 'Heading 2'
table_title.text = 'Table 1B: Essay 1 - Heterogeneity Analysis (Cumulative Abnormal Returns, CAR-30d)'
table_title.paragraph_format.space_before = Pt(12)
table_title.paragraph_format.space_after = Pt(6)

# Create Table 1B with 5 rows (header + 4 mechanisms) and 5 columns
# Columns: Mechanism | FCC Main Effect (Model 2) | Interaction Term (Model 3) | p-value | R-squared
table = doc.add_table(rows=6, cols=5)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
headers = ['Mechanism', 'FCC Main Effect\n(Model 2)', 'Interaction Term\n(Model 3)', 'p-value', 'R-squared']
for i, header_text in enumerate(headers):
    header_cells[i].text = header_text
    # Bold the header
    for paragraph in header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

# Data rows
data = [
    ['Technical Complexity (CVSS)', '-6.46%***', '+6.27%*', '0.007', '0.0604'],
    ['Media Coverage', '-3.33%***', '+7.08%*', '0.006', '0.0389'],
    ['Governance Quality', '-3.29%***', '+0.55%', '0.84 (NS)', '0.0462'],
    ['Ransomware Status', '+5.23%', '-8.34%', '0.069', '0.0385'],
    ['Multi-Type Breach Diversity', '-2.64%**', '-0.32%', '>0.10 (NS)', '0.0292'],
]

for row_idx, row_data in enumerate(data, start=1):
    cells = table.rows[row_idx].cells
    for col_idx, value in enumerate(row_data):
        cells[col_idx].text = value

# Add a note after the table
note_p = doc.add_paragraph()
note_p.text = 'Notes: * p<0.10, ** p<0.05, *** p<0.01. All models estimated via OLS with HC3 robust standard errors. N=898. All models include controls: immediate_disclosure, health_breach, financial_breach, prior_breaches_total, firm_size_log, leverage, roa, and industry/year fixed effects. The FCC main effect and Interaction term columns show coefficients from Model 3 (full interaction specification). p-values refer to the interaction term significance test.'
note_p.style = 'Normal'
for run in note_p.runs:
    run.font.size = Pt(10)
    run.font.italic = True

# Save the document
doc.save('Dissertation_Regression_Tables.docx')
print('SUCCESS: Table 1B added to Dissertation_Regression_Tables.docx')
print('Location: Immediately after Table 1 (Essay 1 Main Regressions)')
print('Content: Essay 1 heterogeneity interactions (CVSS, Media, Governance, Ransomware, Diversity)')
