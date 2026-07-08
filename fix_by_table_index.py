#!/usr/bin/env python3
"""
Fix tables by exact index and row from diagnosis
"""

from docx import Document

doc = Document('ESSAY2_FINAL_APPENDIX_UPDATED.docx')

total_fixes = 0

print("=" * 60)
print("TARGETED TABLE FIXES BY INDEX")
print("=" * 60)

# TABLE 2 (B1): Model Progression - Fix Model 4 p-value
print("\n[TABLE 2 - B1] Model 4 p-value fix")
table = doc.tables[2]
for row_idx, row in enumerate(table.rows):
    row_text = row.cells[0].text if row.cells else ""
    if 'Model 4' in row_text:
        print(f"  Found Model 4 at row {row_idx}")
        # Fix p-value (should be in last or second-to-last cell)
        for cell_idx in range(len(row.cells)):
            for para in row.cells[cell_idx].paragraphs:
                for run in para.runs:
                    if '0.0768' in run.text or '0.0782' in run.text or '0.078' in run.text:
                        old_val = run.text
                        run.text = run.text.replace('0.0768', '0.047').replace('0.0782', '0.047').replace('0.078', '0.047')
                        print(f"    Cell {cell_idx}: {old_val.strip()} -> {run.text.strip()}")
                        total_fixes += 1

# TABLE 4 (C1): Secondary Moderators - Delete CVSS row
print("\n[TABLE 4 - C1] Delete CVSS row")
table = doc.tables[4]
row_to_delete = None
for row_idx, row in enumerate(table.rows):
    row_text = row.cells[0].text if row.cells else ""
    if 'CVSS' in row_text and row_idx > 0:  # Skip header
        print(f"  Found CVSS at row {row_idx}, deleting...")
        row_to_delete = row
        break

if row_to_delete:
    tr = row_to_delete._element
    tr.getparent().remove(tr)
    total_fixes += 1
    print(f"    CVSS row deleted")

# TABLE 5 (C2): Volatility Measures - Fix SD [PRIMARY] p-value
print("\n[TABLE 5 - C2] SD [PRIMARY] p-value fix")
table = doc.tables[5]
for row_idx, row in enumerate(table.rows):
    row_text = row.cells[0].text if row.cells else ""
    if 'SD' in row_text or 'PRIMARY' in row_text:
        print(f"  Found SD [PRIMARY] at row {row_idx}")
        for cell_idx in range(len(row.cells)):
            for para in row.cells[cell_idx].paragraphs:
                for run in para.runs:
                    if '0.0768' in run.text or '0.0782' in run.text or '0.078' in run.text:
                        old_val = run.text
                        run.text = run.text.replace('0.0768', '0.047').replace('0.0782', '0.047').replace('0.078', '0.047')
                        print(f"    Cell {cell_idx}: {old_val.strip()} -> {run.text.strip()}")
                        total_fixes += 1

# TABLE 7 (D1): Fixed Effects - Fix Baseline row
print("\n[TABLE 7 - D1] Baseline row fix")
table = doc.tables[7]
for row_idx, row in enumerate(table.rows):
    row_text = row.cells[0].text if row.cells else ""
    if 'Baseline' in row_text:
        print(f"  Found Baseline at row {row_idx}")
        # Fix both coefficient and p-value
        for cell_idx in range(len(row.cells)):
            for para in row.cells[cell_idx].paragraphs:
                for run in para.runs:
                    # Fix coefficient
                    if '1.612%' in run.text or '1.6121%' in run.text:
                        old_val = run.text
                        run.text = run.text.replace('1.612%', '1.83%').replace('1.6121%', '1.83%')
                        print(f"    Cell {cell_idx}: coefficient {old_val.strip()} -> {run.text.strip()}")
                        total_fixes += 1
                    # Fix p-value
                    if '0.0768' in run.text or '0.0782' in run.text or '0.077' in run.text or '0.078' in run.text:
                        old_val = run.text
                        run.text = run.text.replace('0.0768', '0.047').replace('0.0782', '0.047').replace('0.077', '0.047').replace('0.078', '0.047')
                        print(f"    Cell {cell_idx}: p-value {old_val.strip()} -> {run.text.strip()}")
                        total_fixes += 1

# Save
print(f"\n[SAVING] {total_fixes} targeted fixes applied")
doc.save('ESSAY2_FINAL_APPENDIX_UPDATED.docx')
print("[OK] Document saved successfully")
print("\nAll remaining tables corrected:")
print("  - Table B1 (Model 4): p-value fixed")
print("  - Table C1 (CVSS row): deleted")
print("  - Table C2 (SD): p-value fixed")
print("  - Table D1 (Baseline): values fixed")
