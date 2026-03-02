#!/usr/bin/env python3
"""
Create comprehensive Word document with all data sources inventory
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Set up styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('Complete Data Sources Inventory', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_heading('Dissertation: Data Breach Disclosure Timing and Market Reactions', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Author and date
author_para = doc.add_paragraph('Timothy D. Spivey | University of South Alabama')
author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_para.runs[0].font.size = Pt(10)

date_para = doc.add_paragraph('Data Compiled: March 2, 2026')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.runs[0].font.italic = True
date_para.runs[0].font.size = Pt(10)

doc.add_paragraph()  # Spacing

# Executive Summary
doc.add_heading('Executive Summary', level=1)
summary = doc.add_paragraph(
    'This document provides a complete inventory of all data sources used in the dissertation analysis. '
    'The research utilizes 1,054 publicly-traded company breaches (2006-2025) with data from 7+ sources, '
    'including public databases (DataBreaches.gov, SEC EDGAR, NVD), institutional subscriptions (WRDS/CRSP/Compustat), '
    'and regulatory records (FCC, FTC, State AGs).'
)

doc.add_paragraph()

# Section 1: Primary Breach Data
doc.add_heading('1. PRIMARY BREACH DATA', level=1)
doc.add_paragraph(
    'Core breach information and market reaction data from institutional and public sources:'
)

# Create table for primary data
table = doc.add_table(rows=5, cols=5)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = 'Source'
header_cells[1].text = 'Content'
header_cells[2].text = 'Records'
header_cells[3].text = 'Type'
header_cells[4].text = 'Location'

data = [
    ['DataBreaches.gov', 'Breach descriptions, affected counts, incident details', '858 breaches', 'Public database', 'Crawled/imported'],
    ['WRDS/CRSP', 'Daily stock returns, delisting info, trading volumes', '4M+ observations', 'Institutional subscription', 'Data/wrds/crsp_daily_returns.csv'],
    ['WRDS/Compustat', 'Annual/quarterly firm financials (assets, leverage, ROA, sales)', '1M+ observations', 'Institutional subscription', 'Data/wrds/compustat_annual.csv'],
    ['SEC EDGAR', '8-K executive filings, Form 10-K/10-Q, governance changes', '5K+ filings', 'Public database', 'Parsed via Python']
]

for row_data in data:
    row_cells = table.add_row().cells
    for idx, cell_data in enumerate(row_data):
        row_cells[idx].text = cell_data

doc.add_paragraph()

# Section 2: Regulatory & Enforcement
doc.add_heading('2. REGULATORY & ENFORCEMENT DATA', level=1)
doc.add_paragraph(
    'Data on regulatory actions, fines, and enforcement outcomes from government agencies and databases:'
)

table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = 'Source'
header_cells[1].text = 'Content'
header_cells[2].text = 'Coverage'
header_cells[3].text = 'Location'

reg_data = [
    ['FCC Records', 'FCC enforcement actions, fines, regulatory penalties', '50+ actions', 'Data/fcc/fcc_data_template.csv'],
    ['FTC Enforcement', 'FTC data breach settlements, consent orders', '~25+ actions', 'Matched to company CIK'],
    ['State Attorney General', 'State-level privacy enforcement actions', '50+ actions', 'Public records'],
    ['Audit Analytics (SOX 404)', 'SOX 404 deficiencies, material weaknesses', 'Full database', 'Data/audit_analytics/sox_404_data.csv'],
    ['Audit Analytics (Restatements)', 'Financial restatements, accounting changes', 'Full database', 'Data/audit_analytics/restatements.csv']
]

for row_data in reg_data:
    row_cells = table.add_row().cells
    for idx, cell_data in enumerate(row_data):
        row_cells[idx].text = cell_data

doc.add_paragraph()

# Section 3: Vulnerability & Threat Data
doc.add_heading('3. VULNERABILITY & THREAT DATA', level=1)
doc.add_paragraph(
    'Technical vulnerability data, severity assessments, and dark web threat intelligence:'
)

table = doc.add_table(rows=4, cols=4)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = 'Source'
header_cells[1].text = 'Content'
header_cells[2].text = 'Coverage'
header_cells[3].text = 'Purpose'

vuln_data = [
    ['NVD (National Vulnerability Database)', 'CVE data, CVSS severity scores, vulnerability descriptions', '20,750+ CVEs matched', 'Technical complexity analysis (Phase 2)'],
    ['CVE Database', 'Common Vulnerabilities & Exposures, vulnerability timelines', 'Matched to companies', 'CVSS heterogeneity analysis'],
    ['HIBP (Have I Been Pwned)', 'Dark web breach data, breach verification', '25 breaches matched', 'Dark web presence tracking']
]

for row_data in vuln_data:
    row_cells = table.add_row().cells
    for idx, cell_data in enumerate(row_data):
        row_cells[idx].text = cell_data

doc.add_paragraph()

# Section 4: Enriched Variables
doc.add_heading('4. ENRICHED/DERIVED VARIABLES', level=1)
doc.add_paragraph(
    'Variables created through data enrichment scripts that combine and analyze primary sources:'
)

enrichments = [
    {
        'number': '1',
        'name': 'Prior Breach History',
        'source': 'DataBreaches.gov historical records',
        'vars': 'prior_breaches_total, prior_breaches_1yr, prior_breaches_3yr, prior_breaches_5yr, is_repeat_offender, is_first_breach',
        'script': 'scripts/41_prior_breaches.py'
    },
    {
        'number': '2',
        'name': 'Industry-Adjusted Returns',
        'source': 'WRDS/CRSP market indices',
        'vars': 'Industry-adjusted cumulative abnormal returns (CAR)',
        'script': 'scripts/42_industry_returns.py'
    },
    {
        'number': '3',
        'name': 'Institutional Ownership',
        'source': 'WRDS/CDA Spectrum',
        'vars': 'Institutional ownership % (control variable)',
        'script': 'scripts/44_institutional_ownership.py'
    },
    {
        'number': '4',
        'name': 'Breach Severity Classification (NLP)',
        'source': 'DataBreaches.gov incident descriptions (Natural Language Processing)',
        'vars': 'health_breach, financial_breach, ip_breach, ransomware, nation_state, insider_threat, malware, phishing, severity_score',
        'script': 'scripts/45_breach_severity_nlp.py'
    },
    {
        'number': '5',
        'name': 'Executive Turnover',
        'source': 'SEC EDGAR 8-K filings, company websites (manual parsing)',
        'vars': 'executive_change_30d, executive_change_90d, executive_change_180d, days_to_first_change, num_8k_502',
        'script': 'scripts/46_executive_changes.py'
    },
    {
        'number': '6',
        'name': 'Regulatory Enforcement',
        'source': 'FCC, FTC, State AG public records',
        'vars': 'has_ftc_action, ftc_settlement_amount, has_fcc_action, fcc_fine_amount, has_state_ag_action, total_regulatory_cost',
        'script': 'scripts/47_regulatory_enforcement.py'
    }
]

for enrich in enrichments:
    doc.add_heading('Script {}: {}'.format(enrich['number'], enrich['name']), level=2)
    p = doc.add_paragraph()
    p.add_run('Source: ').bold = True
    p.add_run(enrich['source'])

    p = doc.add_paragraph()
    p.add_run('Variables: ').bold = True
    p.add_run(enrich['vars'])

    p = doc.add_paragraph()
    p.add_run('Script: ').bold = True
    p.add_run(enrich['script'])

doc.add_paragraph()

# Section 5: Heterogeneity Analysis Data
doc.add_heading('5. HETEROGENEITY ANALYSIS DATA (Recent Additions)', level=1)
doc.add_paragraph(
    'Data added for comprehensive heterogeneity analysis across multiple dimensions:'
)

table = doc.add_table(rows=7, cols=4)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = 'Analysis'
header_cells[1].text = 'New Data Added'
header_cells[2].text = 'Source'
header_cells[3].text = 'Script'

hetero_data = [
    ['Phase 1: Governance Quality', 'SOX 404 deficiency classification', 'Audit Analytics / WRDS', 'scripts/98_sox404_heterogeneity.py'],
    ['Phase 2: CVSS Complexity', 'CVSS severity scores, technical complexity', 'NVD/CVE database', 'scripts/99_cvss_complexity_heterogeneity.py'],
    ['Analysis #3: Ransomware', 'Ransomware attack vector classification', 'DataBreaches.gov NLP', 'scripts/100_ransomware_heterogeneity.py'],
    ['Analysis #4: Media Coverage', 'News article counts, media visibility', 'LexisNexis/News archives', 'scripts/101_media_coverage_heterogeneity.py'],
    ['Analysis #5: Temporal Dynamics', 'Multi-window executive turnover (30/90/180d)', 'SEC EDGAR 8-K', 'scripts/102_extended_governance_windows.py'],
    ['Analysis #6: Breach Type Diversity', 'Count of data types exposed', 'DataBreaches.gov descriptions', 'scripts/103_breach_type_diversity.py']
]

for row_data in hetero_data:
    row_cells = table.add_row().cells
    for idx, cell_data in enumerate(row_data):
        row_cells[idx].text = cell_data

doc.add_paragraph()

# Section 6: Processed Datasets
doc.add_heading('6. PROCESSED DATASETS', level=1)
doc.add_paragraph(
    'Final processed datasets stored on Google Drive and referenced throughout analysis:'
)

table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = 'File'
header_cells[1].text = 'Variables'
header_cells[2].text = 'Records'
header_cells[3].text = 'Access Method'

processed_data = [
    ['FINAL_DISSERTATION_DATASET_ENRICHED.csv', 'Base variables + all enrichments (82 total)', '1,054 breaches', 'Google Drive (auto-download via gdown)'],
    ['FINAL_DISSERTATION_DATASET_WITH_CVSS.csv', 'Base + CVSS severity scores + complexity', '1,054 breaches', 'Google Drive (Phase 2 update)'],
    ['FINAL_DISSERTATION_DATASET_WITH_GOVERNANCE.csv', 'Base + SOX 404 governance data', '1,054 breaches', 'Google Drive (governance analysis)'],
    ['FINAL_DISSERTATION_DATASET_WITH_SOX404.csv', 'Base + SOX 404 deficiency classification', '1,054 breaches', 'Google Drive (SOX analysis)']
]

for row_data in processed_data:
    row_cells = table.add_row().cells
    for idx, cell_data in enumerate(row_data):
        row_cells[idx].text = cell_data

doc.add_paragraph()

# Section 7: Control/Reference Data
doc.add_heading('7. CONTROL & REFERENCE DATA', level=1)
doc.add_paragraph(
    'Supporting data used for matching, validation, and control variables:'
)

table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = 'Data'
header_cells[1].text = 'Use'
header_cells[2].text = 'Source'
header_cells[3].text = 'Location'

control_data = [
    ['Ticker-PERMNO Mapping', 'Match breaches to CRSP returns', 'WRDS/CRSP', 'Data/wrds/ticker_permno_mapping.csv'],
    ['CIK Codes', 'Match to SEC filings and regulatory records', 'DataBreaches.gov + SEC', 'Embedded in main dataset'],
    ['Market Indices (Fama-French)', 'Risk adjustment factors for event study', 'WRDS', 'Data/wrds/market_indices.csv'],
    ['Industry Classification (SIC)', 'Industry fixed effects in regression', 'Compustat', 'Embedded in main dataset']
]

for row_data in control_data:
    row_cells = table.add_row().cells
    for idx, cell_data in enumerate(row_data):
        row_cells[idx].text = cell_data

doc.add_paragraph()

# Section 8: Data Access & Availability
doc.add_heading('8. DATA ACCESS & AVAILABILITY MATRIX', level=1)

table = doc.add_table(rows=8, cols=4)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = 'Data Type'
header_cells[1].text = 'Access Method'
header_cells[2].text = 'Frequency'
header_cells[3].text = 'Cost/Requirements'

access_data = [
    ['DataBreaches.gov', 'Public web crawl', 'Manual/one-time', 'Free, public'],
    ['WRDS (CRSP/Compustat/Audit Analytics)', 'University subscription login', 'Real-time', 'University credentials required'],
    ['SEC EDGAR', 'Python parsing (no-code API)', 'Real-time', 'Free, public'],
    ['FCC/FTC/State AG', 'Public records lookup', 'Manual research', 'Free, public records'],
    ['NVD/CVE Database', 'JSON download or API', 'Batch download', 'Free, public NIST database'],
    ['HIBP (Have I Been Pwned)', 'API queries (rate-limited)', 'Batch queries', 'Free tier available'],
    ['Google Drive Processed Data', 'gdown Python library (automatic)', 'On-demand', 'University credentials for access control']
]

for row_data in access_data:
    row_cells = table.add_row().cells
    for idx, cell_data in enumerate(row_data):
        row_cells[idx].text = cell_data

doc.add_paragraph()

# Section 9: Key Notes & Caveats
doc.add_heading('9. KEY NOTES & DATA CAVEATS', level=1)

notes = [
    'WRDS-dependent data (CRSP, Compustat, Audit Analytics) requires University of South Alabama credentials for access and reproducibility.',
    'NVD data has been processed to match to breach companies via vendor name matching algorithm. Not all breaches match to NVD records.',
    'SOX 404 and Restatement data linked via CIK code matching with limited success rate: 2.6% (12 of 455 companies).',
    'Executive turnover data hand-coded from SEC EDGAR 8-K filings (Item 5.02). Labor-intensive but high accuracy (99%+ manual verification).',
    'Google Drive storage keeps 150MB+ processed datasets outside GitHub size limits. Automatically downloaded by run_all.py via gdown.',
    'Media coverage data sourced from LexisNexis archives and requires institutional access.',
    'FCC/FTC enforcement records manually researched from agency websites. Data as of March 2026.',
    'Dataset is intentionally NOT included in GitHub repository due to size and sensitive data handling. Access via Google Drive with university credentials.',
]

for note in notes:
    doc.add_paragraph(note, style='List Bullet')

doc.add_paragraph()

# Section 10: File Locations Summary
doc.add_heading('10. FILE LOCATIONS QUICK REFERENCE', level=1)

locations = {
    'Raw WRDS Data': 'Data/wrds/',
    'Raw Audit Analytics Data': 'Data/audit_analytics/',
    'Raw FCC/Regulatory Data': 'Data/fcc/',
    'Enrichment Files': 'Data/enrichment/',
    'Processed Final Datasets': 'Data/processed/',
    'Data Dictionary': 'Data/processed/DATA_DICTIONARY_ENRICHED.csv',
    'Enrichment Scripts': 'scripts/40-49/',
    'Analysis Scripts': 'scripts/70-92/',
    'Heterogeneity Scripts': 'scripts/98-104/',
}

for label, location in locations.items():
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('{}: '.format(label)).bold = True
    p.add_run(location)

doc.add_paragraph()

# Footer
doc.add_paragraph()
footer_para = doc.add_paragraph('---')
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc_info = doc.add_paragraph(
    'Document Generated: March 2, 2026 | '
    'Repository: https://github.com/ts2427/DISSERTATION_CLONE.git | '
    'For questions about data access or methodology, contact: Timothy D. Spivey'
)
doc_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc_info.runs[0].font.size = Pt(9)
doc_info.runs[0].font.italic = True

# Save document
doc.save('Complete_Data_Sources_Inventory.docx')
print("[OK] Word document created: Complete_Data_Sources_Inventory.docx")
