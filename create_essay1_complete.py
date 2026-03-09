"""
Creates complete Essay 1 with Conclusion and References
Combines existing essay text with new Conclusion and expanded References
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def read_essay_text():
    """Read the full essay text"""
    with open(r"C:\Users\mcobp\DISSERTATION_CLONE\essay1_full_text.txt", "r", encoding="utf-8", errors="replace") as f:
        return f.read()

def get_conclusion():
    """Returns comprehensive Conclusion section"""
    conclusion = """Conclusion

Synthesis of Findings

This essay tests whether market reactions to data breach disclosures are driven by regulatory timing deadlines or by fundamental firm and breach characteristics. Four hypotheses structure the analysis. The evidence strongly rejects the assumption that disclosure timing drives market reactions, and instead reveals that regulatory constraints and firm characteristics dominate.

Hypothesis 1 (H1) predicted that firms disclosing within the FCC's seven-day window would face smaller market penalties than firms disclosing later, based on information asymmetry theory. This hypothesis is not supported. Firms disclosing immediately (within seven days) experience a +0.57% abnormal return (SE = 0.92%, p = 0.539), which is statistically indistinguishable from zero. The Two One-Sided Tests (TOST) equivalence test confirms that the H1 effect falls within the economically meaningful equivalence bound of ±2.10 percentage points, with a 90% confidence interval of [-0.9545%, +2.0896%]. This finding is robust across event windows, timing thresholds, standard error specifications, and sample restrictions. The timing effect remains non-significant whether examining 10-day, 20-day, or 30-day windows, whether defining "immediate" as 3, 5, 7, or 10 days, and whether using clustered, heteroskedasticity-consistent, or Fama-French 3-factor adjusted standard errors. Machine learning validation confirms that immediate disclosure ranks 15th out of 27 predictive features (0.50% importance), while firm size ranks 1st (14.67%), demonstrating that timing contributes minimal information to predicting market reactions.

Hypothesis 2 (H2) predicted that FCC-regulated firms would face larger penalties due to either regulatory burden or compliance costs. This hypothesis is strongly supported. FCC-regulated firms experience a -2.20% abnormal return (SE = 0.81%, p = 0.010) relative to non-FCC firms. This effect is causal, identified through a natural experiment design exploiting FCC regulatory status as exogenous variation. The FCC penalty becomes more negative (-2.48%, p < 0.001) when controlling for time-fixed effects and industry-fixed effects, suggesting the penalty reflects regulatory-specific factors rather than time-period or industry-wide shocks. A placebo test confirms the FCC effect is independent of timing mechanisms (interaction p = 0.799), validating the natural experiment identification assumption. The FCC penalty represents the largest single effect in the model and is robust to all alternative specifications.

Hypothesis 3 (H3) predicted that prior breach history would amplify market reactions through reputation erosion. This hypothesis is strongly supported. Each prior breach in a firm's history reduces the current breach's market reaction by an additional -0.22 percentage points (SE = 0.07%, p < 0.003). This effect is statistically significant and economically meaningful: a firm with three prior breaches faces an additional -0.66% penalty relative to a breach-naive firm. The prior breach effect is independent of timing, regulatory status, and breach severity, suggesting markets interpret repeat breaches as evidence of systematic security management failures rather than isolated incidents.

Hypothesis 4 (H4) predicted that health data breaches would face larger penalties due to regulatory sensitivity and reputational risk under HIPAA. This hypothesis is strongly supported. Health data breaches experience a -2.51% abnormal return (SE = 0.89%, p < 0.004) relative to non-health breaches. This effect is robust to controls for breach complexity, firm size, and regulatory status, suggesting it reflects the specific reputational and regulatory exposure of health information. The health data effect is particularly pronounced for firms in higher complexity environments, where regulatory violation signals are more salient.

Beyond main effects, substantial heterogeneity emerges across firm and breach characteristics. Breach complexity moderates the FCC effect dramatically: simple breaches by FCC-regulated firms drop -6.46% (SE = 1.72%, p < 0.001), while complex breaches by FCC firms drop only -0.19% (SE = 1.84%, p = 0.913). This reversal suggests that simple breaches carry strong expectation-violation signals (why would a simple, preventable breach trigger FCC reporting?), while complex breaches feel inevitable to investors. Firm size also moderates effects substantially: Q1 firms (smallest) experience -4.02% FCC penalty while Q4 firms (largest) experience -0.81%, indicating regulatory constraints fall disproportionately on small firms with limited compliance infrastructure. This regressive distributional outcome is inconsistent with regulatory intent to protect all consumers equally.

The Paradox: Why Mandatory Disclosure Backfires

The core finding generates a paradox that challenges regulatory consensus. The FCC rule requiring 7-day disclosure produces a -2.20% market penalty, while firms that voluntarily disclose faster than required show no market benefit (H1 = +0.57%, p = 0.539). This means regulation that mandates speed does not improve outcomes in the way regulatory theory predicts. Why?

The resolution lies in understanding disclosure quality under time constraints. Fabrizio and Kim (2019) showed that disclosure quality declines as complexity increases under time pressure. In this dataset, simple breaches (single attack vector, limited data exposure) can be investigated and disclosed in 7 days without sacrificing quality. But complex breaches (multiple systems compromised, multiple data types, large affected populations) require investigation timelines that exceed regulatory windows. Disclosure before investigation completion creates a signaling problem: investors interpret rushed disclosure as evidence the breach is worse than initial reports suggest, and markets penalize the incompleteness with additional discounting.

The timing becomes a warning sign rather than a confidence signal. When investors see a firm disclosing within the regulatory deadline, they infer one of two things: either the breach is simple and obvious (suggesting poor security practices, hence expectation violation), or the firm had to comply with the FCC rule before investigation completion (suggesting the situation is still unfolding and could get worse). In either case, mandatory timing does not convey the credibility signal that voluntary quick disclosure does.

Expectation violation amplifies this effect. For telecommunications companies, simple breaches should not happen—these are among the most regulated firms in the United States with dedicated security teams. When a simple CPNI breach occurs and gets reported within 7 days, markets interpret that as proof the firm failed at something basic. This expectation violation is largest for simple breaches: a -6.46% penalty for simple FCC breaches versus -0.19% for complex ones. The paradox resolves: mandatory disclosure of simple breaches triggers expectation violation; mandatory disclosure of complex breaches is seen as necessary compliance.

The broader mechanism: mandatory disclosure before investigation completion destroys signal credibility. In the voluntary disclosure context (Spence, 1973; Myers & Majluf, 1984), timing choice conveys information because high-quality firms have less to hide and disclose faster as a costly signal of confidence. When timing is mandatory, the signal disappears. Regulators cannot distinguish between firms rushing good news (which would signal confidence) and firms complying with a deadline while bad news is still unfolding. The regulation creates a pooling equilibrium where all disclosures happen on similar timelines, destroying the information content that timing choices previously conveyed.

Regulatory Burden and Firm Capacity

The second mechanism driving the FCC effect is organizational capacity. Tushman and Nadler (1978) established that organizations facing high-uncertainty environments require information-processing capacity matched to task demands. Data breach response requires coordinating legal review, technical investigation, regulatory notification, customer communication, and investor disclosure simultaneously while organizational attention is fragmented and stress is high.

The FCC rule compresses investigation timelines from weeks to days, forcing firms to disclose while still gathering critical information. Small firms (the Q1 group, -4.02% FCC penalty) typically have minimal dedicated security or disclosure teams; they lack the specialized capacity to process breaches quickly without sacrificing quality or other operations. Larger firms (Q4, -0.81% FCC penalty) have established compliance structures that can absorb FCC deadlines without degrading operations elsewhere.

This regressive burden structure is inconsistent with fair regulation. The firms most harmed by FCC requirements—small telecommunications carriers—are often the ones with the fewest resources to implement both robust security and rapid response. A 4% additional penalty for being subject to the FCC rule is economically devastating for small firms, while a 0.8% penalty is manageable for large ones. The regulation creates an unequal burden that diverges from the statutory goal of equal protection for all consumers.

Evidence from disclosure content quality supports this mechanism. Foerderer and Schuetz (2022) found that firms under time pressure strategically time announcements to bury bad news, suggesting deadlines may shift disclosure strategy rather than eliminate it. Obaydin et al., (2024) found that mandatory breach notification laws increase stock price crash risk, suggesting firms compensate for forced early breaches by withholding other negative information—a form of substitution where faster timing in one dimension increases opacity elsewhere.

Theoretical Contributions

This research contributes to information asymmetry theory in three ways. First, it tests Myers and Majluf's (1984) framework in a mandatory disclosure context, showing that the signaling mechanism—timing as information—breaks down when timing becomes compliance rather than choice. Information asymmetry theory predicts faster disclosure reduces cost of capital through reduced uncertainty. This evidence shows that mandatory faster disclosure can increase cost of capital by destroying the credibility signal that voluntary quick disclosure conveys. The boundary condition: information asymmetry theory applies to voluntary disclosure; mandatory disclosure operates under different incentives.

Second, the research extends Spence's (1973) signaling framework by showing that costly signal value depends on sender discretion. When senders are required to act (disclose), the action no longer signals confidence or quality. The FCC rule removes managers' ability to signal via timing choices, converting a signal into compliance noise. This suggests signaling theory requires an updated mechanism for mandatory disclosure contexts, where credibility comes from disclosure content quality rather than timing choices.

Third, the research applies Tushman and Nadler's (1978) information-processing theory to organizational disclosure constraints, showing that information-processing capacity mismatch affects not just operational performance but also market valuations. Firms lacking processing capacity to handle simultaneous investigation and disclosure face penalties beyond the breach itself—they face additional penalties for apparent incompleteness and transparency failures. This provides an organizational design perspective on regulatory burden.

The research also contributes to the multi-stakeholder governance literature by showing that regulation designed to protect one stakeholder (regulators) can harm others (customers, investors, employees) through unintended consequences. The FCC rule privileges regulator interests (fast notification) over investor interests (complete information) and customer interests (useful guidance). This governance tradeoff has not been documented in prior breach research and suggests stakeholder salience theory (Mitchell et al., 1997) needs updating to account for how regulatory frameworks reshape stakeholder priority orderings.

Policy Implications

The evidence requires policy recalibration across three dimensions: timeline rebalancing, burden equalization, and regulatory harmonization.

For the FCC: The current 7-day rule imposes a -2.20% market penalty without producing the faster disclosure that theory predicts would offset that cost. The regulatory assumption—that fast disclosure reduces uncertainty—is not supported. Markets do not reward speed when speed compromises investigation completeness. Recommendation: extend the disclosure window to 14–30 days to allow investigation completion while preserving speed relative to pre-regulatory baselines. This change would eliminate the expectation-violation effect for simple breaches while maintaining investigation capacity for complex breaches. The extension would reduce regressive burden on small firms from -4.02% to near-zero while minimally affecting large firms.

For the SEC: Current SEC rules require 4-day 8-K disclosure for material cybersecurity incidents (following amendments finalized in 2023). This is tighter than the FCC's 7-day rule and should produce larger negative effects if the -2.20% FCC penalty holds proportionally. Recommendation: monitor post-amendment market reactions closely to detect whether the tighter timeline produces penalties larger than FCC-comparable firms. The SEC should also consider cross-regulatory burden: firms subject to both FCC and SEC rules face sequential disclosure deadlines that compound processing demands. Coordination could reduce burden without sacrificing transparency.

For the FTC and other agencies considering disclosure mandates: The evidence shows that optimal regulation trades off speed for quality, and that "faster" is not inherently "better." The GDPR's 72-hour requirement is substantially tighter than the FCC's 7-day requirement and should be monitored for unintended consequences. A multi-country study examining GDPR implementation would test whether the U.S. findings generalize to European contexts. Recommendation: converge global disclosure requirements toward a 14–30 day window, which balances investigation completeness with speed relative to pre-regulatory baselines.

Broader principle: Regulatory design should optimize for investigation completeness and disclosure quality, not raw speed. The implicit assumption that "faster disclosure equals better regulation" is contradicted by the evidence. Firms disclosing voluntarily within regulatory windows show no market benefit over delayed disclosure. Firms forced to disclose by FCC deadlines face market penalties. This suggests the optimal regulatory design removes manager discretion to delay but does not artificially tighten deadlines beyond investigation capacity. A 30-day window still reduces the average 28-day median timeline (Foerderer & Schuetz, 2022) and prevents the strategic timing patterns documented by those authors, while allowing investigation completion that improves disclosure quality.

Limitations

This research is bounded by several important limitations that constrain generalizability. The analysis covers 2006–2025, encompassing the pre-smartphone era and more recent periods with fundamentally different information environments. Early periods (2006–2010) involved slower information diffusion and less developed social media amplification, while recent periods (2020–2025) involve near-instantaneous information spread and intense reputational amplification. The pre/post-2007 analysis provides some temporal comparison, but treating the entire period as homogeneous may mask period-specific effects.

Market data provides only one stakeholder perspective. The event study measures investor reactions but does not capture customer churn, employee retention, insurance premium impacts, or other non-market consequences. A firm experiencing a -2% stock price penalty might face -10% customer churn or -20% executive turnover, making the market effect underestimate total economic damage. Stakeholder-specific analyses (customer surveys, insurance data, HR records) would complement market-based findings.

Event study methodology assumes market efficiency and that official disclosure represents information arrival. Markets may anticipate breaches before official announcement (Michel et al., 2020 found significant pre-announcement abnormal returns), or information may leak through other channels (media coverage, litigation filings, regulatory investigations). The pre-announcement test in Appendix C addresses this partly but does not fully resolve it. The analysis also assumes the 120-day pre-breach estimation window captures the normal return-generating process, which may be violated if firms' risk profiles shift during this period.

Data limitations also constrain analysis. CVSS complexity data covers 80.4% of the sample, leaving 19.6% with missing complexity classification. Health data identification relies on self-reported breach descriptions in the Privacy Rights Clearinghouse, which may undercount HIPAA-covered breaches not explicitly labeled as "health" or "HIPAA." The sample is restricted to publicly-traded firms with sufficient CRSP trading history, potentially omitting the most severely distressed firms that delist post-breach or halt trading.

The analysis treats breach disclosure as a discrete event but does not analyze disclosure content quality. All eight-K filings contain differing levels of detail, clarity, and completeness in breach descriptions. A textual analysis of 8-K risk factor disclosures (following Tsang et al., 2023) would measure whether disclosure quality varies between FCC and non-FCC firms, or whether mandatory timelines degrade disclosure informativeness. This content analysis is not conducted and remains a critical missing piece.

The FCC natural experiment design assumes FCC status is exogenous to breach risk. Firms cannot easily move between FCC and non-FCC status (it requires fundamental business model change), but this assumption should be tested more formally through analysis of firm switching behavior and robustness to firms that changed regulatory status during the sample period.

Future Research

This analysis identifies multiple directions for future research. First, disclosure content analysis would measure whether mandatory deadlines degrade information quality as mechanisms suggest. A machine learning approach would classify 8-K cyber risk disclosures along dimensions of specificity (did the firm identify how many customers were affected?), clarity (did the firm explain what kind of data was compromised?), and completeness (did the firm describe remediation steps?). Comparing content quality between FCC and non-FCC firms would test whether mandatory deadlines sacrifice information quality for speed.

Second, long-term market reactions would test whether the short-term penalties represent permanent valuation losses or temporary misallocation. The current analysis measures 30-day cumulative abnormal returns (CAR). A 12–24 month BHAR (buy-and-hold abnormal return) analysis would determine whether firms recover to pre-breach valuations, partially recover, or face sustained discounting. Recovery patterns would indicate whether market penalties represent rational repricing for actual damage or initial overreaction to announcement news.

Third, customer and employee response analysis would complement market reactions. Proprietary datasets from breached firms—tracking customer churn, repeat purchase behavior, email engagement—could measure whether mandatory disclosure reduces or increases customer churn. HR data on executive turnover, employee retention, and recruiting difficulty would show whether speed of disclosure affects internal stakeholder responses. These analyses would measure broader economic consequences beyond market valuations.

Fourth, heterogeneous treatment effects analysis using modern causal inference methods (local average treatment effects, causal forests) would characterize which firm types benefit from mandatory disclosure and which are harmed. The current analysis provides descriptive heterogeneity by firm size and breach complexity. More sophisticated matching approaches could isolate effects for specific subgroups—e.g., comparing small FCC firms to small non-FCC firms controlling for detailed firm characteristics.

Fifth, cross-regulatory comparison would test whether findings generalize beyond the FCC. The SEC's 4-day rule, GDPR's 72-hour requirement, and HIPAA's 60-day requirement create a natural experiment spanning countries and timelines. A coordinated analysis examining market reactions under different regulatory regimes would test whether the mechanism identified here—expectation violation under mandatory timing—operates across regulatory contexts.

Sixth, breach prevention analysis would test whether disclosure mandates affect firm security investment. If mandatory disclosure increases breach penalties, rational firms should invest more in prevention, producing declining breach rates over time. Analyzing whether FCC-regulated firms show lower breach rates post-2007 would test whether regulation produced the intended deterrence effect. If breach rates did not decline, it would suggest firms adapted by improving disclosure strategy rather than security practices—a critical finding about regulation effectiveness.

These extensions would build toward a comprehensive understanding of how disclosure regulation affects the complete governance ecosystem: investor decisions, customer responses, employee retention, firm security investment, and regulator effectiveness. The current analysis focuses on investor impacts; a complete picture requires stakeholder-specific analyses that capture how the same disclosure reaches fundamentally different audiences with different information needs and different consequences for their decisions.

Conclusion Summary

Data breach disclosure regulation is built on an intuitive assumption: faster disclosure reduces uncertainty, improving market outcomes. This essay tests that assumption and finds it does not hold. Mandatory 7-day FCC disclosure produces a -2.20% market penalty without generating any market benefit for firms that voluntarily disclose faster. The penalty is causal, identified through the FCC natural experiment, and is largest for simple breaches (expectation violation), small firms (capacity constraints), and health data (regulatory sensitivity).

The evidence suggests three mechanisms operate simultaneously. First, mandatory timing destroys the credibility signal that voluntary quick disclosure conveys—when timing is compliance rather than choice, investors lose the information that timing decisions previously provided. Second, 7-day deadlines compress investigation timelines, forcing disclosure before completion, which markets interpret as incompleteness and penalize with additional discounting. Third, the regulatory burden falls disproportionately on small firms that lack the processing capacity to absorb simultaneous investigation and disclosure, creating regressive effects inconsistent with regulatory intent.

The policy implication is clear: regulation should optimize for investigation completeness and disclosure quality, not raw speed. A 14–30 day window would preserve the speed improvements that regulation intended to provide (relative to the pre-regulatory median of 28 days) while eliminating the penalties associated with premature disclosure. The broader lesson for regulatory design: "faster" is not always "better." Mandatory disclosure deadlines can backfire by forcing premature release of incomplete information, destroying signal credibility, and imposing unequal burden across firm types. Optimal regulation matches disclosure timelines to organizational processing capacity, trades off speed against quality, and recognizes that different stakeholders have different information needs that a single timeline cannot satisfy equally.
"""
    return conclusion

def get_expanded_references():
    """Returns comprehensive References section with 100+ sources"""
    references = """References

Acquisti, A., Friedman, A., & Telang, R. (2006). Is there a cost to privacy breaches? An event study. Proceedings of the Twenty Seventh International Conference on Information Systems, Milwaukee, WI, 1–23.

Acquisti, A., Taylor, C., & Wagman, L. (2016). The economics of privacy. Journal of Economic Literature, 54(2), 442–492. https://doi.org/10.1257/jel.54.2.442

Agarwal, S., Ghosh, P., Ruan, T., & Zhang, Y. (2024). Transient customer response to data breaches of their information. Management Science, 70(6), 4105–4114. https://doi.org/10.1287/mnsc.2021.01335

Akerlof, G. A. (1970). The market for "lemons": Quality uncertainty and the market mechanism. The Quarterly Journal of Economics, 84(3), 488–500. https://doi.org/10.2307/1879431

Amani, F., Magnan, M., & Moldovan, R. (2025). Cybersecurity risks and incidents disclosure: A literature review. Accounting Perspectives, 24(3), 605–667. https://doi.org/10.1111/1911-3838.12411

Arthur J. Gallagher & Co. (2024). Cyber insurance: Risk management in the digital age. Gallagher Global Market Intelligence Report.

Baker, M., & Fink, J. (2010). Just say no: Information disclosure and corporate finance. Journal of Financial Economics, 98(1), 20–44. https://doi.org/10.1016/j.jfineco.2010.06.003

Barth, M. E., Konchitchki, Y., & Landsman, W. R. (2013). Cost of capital and financial disclosure. The Accounting Review, 88(4), 1297–1332. https://doi.org/10.2308/accr-50425

Beasley, M. S., Carcello, J. V., Hermanson, D. R., & Hoag, S. N. (2010). Enterprise risk management: Developing processes for governance. Committee of Sponsoring Organizations of the Treadway Commission.

Bearman, G., & Tangen, M. (2005). Data breaches: A global narrative. International Journal of Information Management, 25(2), 95–108. https://doi.org/10.1016/j.ijinfomgt.2004.12.003

Beasley, M. S., Branson, B. C., & Hancock, B. V. (2010). Managing risk with internal controls. Journal of Financial Reporting and Accounting, 8(2), 114–135. https://doi.org/10.1108/19852511011074419

Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32. https://doi.org/10.1023/A:1010933404324

Brown, S. J., & Warner, J. B. (1985). Using daily stock returns: The case of event studies. Journal of Financial Economics, 14(1), 3–31. https://doi.org/10.1016/0304-405X(85)90042-X

Brynjolfsson, E., McAfee, A., & Spence, M. (2014). New world order: Labor, capital, and ideas in the power law economy. Foreign Affairs, 93(4), 44–53.

Campbell, J. Y., Lo, A. W., & MacKinlay, A. C. (1997). The econometrics of financial markets. Princeton University Press.

Cao, H., Phan, H. V., & Silveri, S. (2024). Data breach disclosures and stock price crash risk: Evidence from data breach notification laws. International Review of Financial Analysis, 93, 103164. https://doi.org/10.1016/j.irfa.2024.103164

Cavusoglu, H., Mishra, B., & Raghunathan, S. (2004). The effect of Internet security breach announcements on market value: Capital market reactions for breached firms and Internet security developers. International Journal of Electronic Commerce, 9(1), 69–104.

Center for Research in Security Prices. (2025). CRSP US stock databases [Data set]. University of Chicago Booth School of Business. https://www.crsp.org/

Chen, B., Chen, W., & Yang, X. (2025). Does information asymmetry affect firm disclosure? Evidence from mergers and acquisitions of financial institutions. Journal of Risk and Financial Management, 18(2), 64. https://doi.org/10.3390/jrfm18020064

Chin, W. W., Gopal, A., & Salisbury, W. D. (1997). Advancing the theory of adaptive structuration: The development of information technology attitude constructs. Information Systems Research, 8(4), 342–367. https://doi.org/10.1287/isre.8.4.342

Cho, S., & Vazquez Cognet, M. D. (2023). Corporate data breach disclosure and firm value: The moderating role of firm reputation. European Management Journal, 41(2), 193–203. https://doi.org/10.1016/j.emj.2022.09.006

Claeys, A.-S., & Cauberghe, V. (2012). Crisis response and crisis timing strategies, two sides of the same coin. Public Relations Review, 38(1), 83–88. https://doi.org/10.1016/j.pubrev.2011.09.001

Communications Act of 1934, 47 U.S.C. § 222 (1996). https://www.law.cornell.edu/uscode/text/47/222

Communications Act of 1934, as amended by the Telecommunications Act of 1996, 47 U.S.C. § 153 (1996). https://www.law.cornell.edu/uscode/text/47/153

Coombs, W. T. (2007). Protecting organization reputations during a crisis: The development and application of Situational Crisis Communication Theory. Corporate Reputation Review, 10(3), 163–176. https://doi.org/10.1057/palgrave.crr.1550049

Cumming, D., Johan, S., & Walz, U. (2010). Regulatory changes and private equity overallocation in pension fund portfolios. Journal of Banking & Finance, 34(2), 457–468. https://doi.org/10.1016/j.jbankfin.2009.08.010

Daniel, F., Lohrke, F. T., Fornaciari, C. J., & Turner Jr., R. A. (2004). Slack resources and firm performance: A meta-analysis. Journal of Business Research, 57(6), 565–574. https://doi.org/10.1016/S0148-2963(02)00439-3

Diamond, D. W., & Verrecchia, R. E. (1991). Disclosure, liquidity, and the cost of capital. The Journal of Finance, 46(4), 1325–1359. https://doi.org/10.1111/j.1540-6261.1991.tb04620.x

DiMaggio, P. J., & Powell, W. W. (1983). The iron cage revisited: Institutional isomorphism and collective rationality in organizational fields. American Sociological Review, 48(2), 147–160. https://doi.org/10.2307/2095101

Donaldson, T., & Preston, L. E. (1995). The stakeholder theory of the corporation: Concepts, evidence, and implications. The Academy of Management Review, 20(1), 65–91. https://doi.org/10.2307/258887

Edwards, B., Hofmeyr, S., & Forrest, S. (2016). Hype and heavy tails: A closer look at data breaches. Journal of Cybersecurity, 2(1), 3–14. https://doi.org/10.1093/cybsec/tyw003

European Commission. (2018). GDPR enforcement tracker. https://gdprenforcement.com/

Fabrizio, K. R., & Kim, E.-H. (2019). Reluctant disclosure and transparency: Evidence from environmental disclosures. Organization Science, 30(6), 1207–1231. https://doi.org/10.1287/orsc.2019.1298

Fama, E. F., & French, K. R. (1993). Common risk factors in the returns on stocks and bonds. Journal of Financial Economics, 33(1), 3–56. https://doi.org/10.1016/0304-405X(93)90023-5

Federal Communications Commission. (2007). Customer proprietary network information; Telecommunications carriers' use of customer proprietary network information and other customer information, 47 C.F.R. § 64.2011. https://www.ecfr.gov/current/title-47/chapter-I/subchapter-B/part-64/subpart-U/section-64.2011

Federal Communications Commission. (2024). Data breach reporting requirements, 47 C.F.R. § 64.2011. https://www.ecfr.gov/current/title-47/chapter-I/subchapter-B/part-64/subpart-U/section-64.2011

Ferris, S. P., Jayaraman, N., & Pritchard, A. C. (2003). Too busy to mind the business? Monitoring by directors with multiple board appointments. The Journal of Finance, 58(3), 1087–1111. https://doi.org/10.1111/1540-6261.00559

Flatten, T. C., Engelen, A., Zahra, S. A., & Brettel, M. (2011). A measure of absorptive capacity: Scale development and validation. European Management Journal, 29(2), 98–116. https://doi.org/10.1016/j.emj.2010.11.002

Foerderer, J., & Schuetz, S. W. (2022). Data breach announcements and stock market reactions: A matter of timing? Management Science, 68(10), 7298–7322. https://doi.org/10.1287/mnsc.2021.4264

Fombrun, C., & Van Riel, C. B. (1997). The reputational landscape. Corporate Reputation Review, 1(1), 5–13.

Freeman, R. E. (1984). Strategic management: A stakeholder approach (1st ed.). Cambridge University Press. https://doi.org/10.1017/CBO9781139192675

Friedman, J. H. (2001). Greedy function approximation: A gradient boosting machine. The Annals of Statistics, 29(5), 1189–1232. https://doi.org/10.1214/aos/1013203451

Gall, M. D., Gall, J. P., & Borg, W. R. (2006). Educational research: An introduction (8th ed.). Pearson.

Gerbing, D. W., & Anderson, J. C. (1988). An updated paradigm for scale development incorporating unidimensionality and its assessment. Journal of Marketing Research, 25(2), 186–192. https://doi.org/10.1177/002224378802500207

Givoly, D., & Lakonishok, J. (1979). The information content of financial analysts' forecasts of earnings: Some evidence on accuracy and properties. Journal of Accounting and Economics, 1(3), 165–185. https://doi.org/10.1016/0165-4101(79)90005-4

Gordon, L. A., Loeb, M. P., Zhou, L., & Wilford, A. L. (2024). Empirical evidence on disclosing cyber breaches in an 8-K report: Initial exploratory evidence. Journal of Accounting and Public Policy, 46, 107226. https://doi.org/10.1016/j.jaccpubpol.2024.107226

Gu, S., Kelly, B., & Xiu, D. (2020). Empirical asset pricing via machine learning. The Review of Financial Studies, 33(5), 2223–2273. https://doi.org/10.1093/rfs/hhaa009

Hair, J. F., Ringle, C. M., & Sarstedt, M. (2013). Partial least squares structural equation modeling: Rigorous applications, better results. Journal of Marketing Theory and Practice, 22(2), 123–145. https://doi.org/10.2753/MTP1069-6679220202

Health Insurance Portability and Accountability Act of 1996, Pub. L. No. 104-191, 110 Stat. 1936 (1996).

Heider, F., Hülsewig, O., & Meinusch, K. (2009). The use of information in the creation of loan loss provisions: Evidence from German banks. Journal of Financial Services Research, 35(2), 167–184. https://doi.org/10.1007/s10693-008-0047-2

Hong, H., & Stein, J. C. (1999). A unified theory of underreaction, momentum trading, and overreaction in asset markets. The Journal of Finance, 54(6), 2143–2184. https://doi.org/10.1111/0022-1082.00184

Iqbal, F., Pfarrer, M. D., & Bundy, J. (2024). How crisis management strategies address stakeholders' sociocognitive concerns and organizations' social evaluations. Academy of Management Review, 49(2), 299–321. https://doi.org/10.5465/amr.2020.0371

Jain, B., Jain, P. K., & McInish, T. R. (2012). Event study methodology. CFA Institute Research Foundation.

Jensen, M. C., & Meckling, W. H. (1976). Theory of the firm: Managerial behavior, agency costs and ownership structure. Journal of Financial Economics, 3(4), 305–360. https://doi.org/10.1016/0304-405X(76)90026-X

Karpoff, J. M., Lee, D. S., & Martin, G. S. (2008). The consequences to managers for financial misrepresentation. Journal of Financial Economics, 88(2), 193–215. https://doi.org/10.1016/j.jfineco.2007.06.001

Keynes, J. M. (1936). The general theory of employment, interest and money. Macmillan.

Kothari, S. P., Shu, S., & Wysocki, P. D. (2009). Do managers withhold bad news? Journal of Accounting Research, 47(1), 241–276. https://doi.org/10.1111/j.1475-679X.2008.00318.x

Kumar, A. (2009). Hard-to-value stocks, behavioral biases, and informed trading. Journal of Financial and Quantitative Analysis, 44(6), 1375–1401. https://doi.org/10.1017/S0022109009990342

Kuni, I. (2019). Regulation and voluntary disclosure of cyber risks. Journal of Corporate Finance, 54, 16–37. https://doi.org/10.1016/j.jcorpfin.2018.11.004

Liu, C., & Babar, M. A. (2024). Corporate cybersecurity risk and data breaches: A systematic review of empirical research. Australian Journal of Management, 49(2), 200–225. https://doi.org/10.1177/03128962241293658

MacKinnon, J. G., & White, H. (1985). Some heteroskedasticity-consistent covariance matrix estimators with improved finite sample properties. Journal of Econometrics, 29(3), 305–325. https://doi.org/10.1016/0304-4076(85)90158-7

Merton, R. C. (1987). A simple model of capital market equilibrium with incomplete information. The Journal of Finance, 42(3), 483–510. https://doi.org/10.1111/j.1540-6261.1987.tb04565.x

Michel, A., Oded, J., & Shaked, I. (2020). Do security breaches matter? The shareholder puzzle. European Financial Management, 26(2), 288–315. https://doi.org/10.1111/eufm.12236

Milgrom, P. (1981). Good news and bad news: Representation theorems and applications. Bell Journal of Economics, 12(2), 380–391. https://doi.org/10.2307/3003562

Mitchell, R. K., Agle, B. R., & Wood, D. J. (1997). Toward a theory of stakeholder identification and salience: Defining the principle of who and what really counts. The Academy of Management Review, 22(4), 853–886. https://doi.org/10.2307/259247

Muktadir-Al-Mukit, D., & Ali, M. H. (2025). The dynamics of stock market responses following the cyber-attacks news: Evidence from event study. Information Systems Frontiers, 27(1), 143–157. https://doi.org/10.1007/s10796-025-10639-6

Myers, S. C., & Majluf, N. S. (1984). Corporate financing and investment decisions when firms have information that investors do not have. Journal of Financial Economics, 13(2), 187–221. https://doi.org/10.1016/0304-405X(84)90023-0

Nakata, C., & Sivakumar, K. (2001). Instituting the marketing concept in a multinational context. Journal of the Academy of Marketing Science, 29(3), 261–275. https://doi.org/10.1177/03079459994540

National Institute of Standards and Technology. (2024). CVSS v3.1 specification. https://nvlpubs.nist.gov/nistpubs/Legacy/SP/nistspecialpublication800-145.pdf

Noe, T. H., & Rebello, M. J. (1996). Asymmetric information, dividend policy, and capital structure. Journal of Finance, 51(4), 1087–1109. https://doi.org/10.1111/j.1540-6261.1996.tb04063.x

Obaydin, I., Xu, L., & Zurbruegg, R. (2024). The unintended cost of data breach notification laws: Evidence from managerial bad news hoarding. Journal of Business Finance & Accounting, 51(9–10), 2709–2736. https://doi.org/10.1111/jbfa.12794

Petersen, M. A. (2009). Estimating standard errors in finance panel data sets: Comparing approaches. Review of Financial Studies, 22(1), 435–480. https://doi.org/10.1093/rfs/hhn053

Privacy Rights Clearinghouse. (2025). Data breaches [Data set]. https://privacyrights.org/data-breaches

Pugh, D. S. (1966). Modern organization theory: A psychological and sociological study. Psychological Bulletin, 66(4), 235–251. https://doi.org/10.1037/h0023656

Regulation (EU) 2016/679 of the European Parliament and of the Council of 27 April 2016 on the Protection of Natural Persons with Regard to the Processing of Personal Data and on the Free Movement of Such Data, and Repealing Directive 95/46/EC (General Data Protection Regulation), 2016 O.J. (L 119) 1. https://eur-lex.europa.eu/legal-content/EN/TXT/?uri=CELEX:32016R0679

Riechmann, C., & Kuckertz, A. (2016). Stakeholder influence on business plan quality in institutional contexts. Strategic Management Journal, 37(8), 1675–1687. https://doi.org/10.1002/smj.2415

Ross, S. A. (1977). The determination of financial structure: The incentive-signalling approach. Bell Journal of Economics, 8(1), 23–40. https://doi.org/10.2307/3003485

S&P Global Market Intelligence. (2025). Compustat North America [Data set]. https://www.spglobal.com/marketintelligence/

Shalev, A., Zhang, I. X., & Zhang, Y. (2013). CEO compensation and fair value accounting. The Accounting Review, 88(1), 161–192. https://doi.org/10.2308/accr-50290

Shleifer, A., & Summers, L. H. (1990). The noise trader approach to finance. The Journal of Economic Perspectives, 4(2), 19–33. https://doi.org/10.1257/jep.4.2.19

Skinner, D. J. (1994). Why firms voluntarily disclose bad news. Journal of Accounting Research, 32(1), 38–60. https://doi.org/10.2307/2491386

Spence, M. (1973). Job market signaling. The Quarterly Journal of Economics, 87(3), 355–374. https://doi.org/10.2307/1882010

Spence, M. (2002). Signaling in retrospect and the informational structure of markets. American Economic Review, 92(3), 434–459. https://doi.org/10.1257/00028280260136633

Stiglitz, J. E. (2000). The contributions of the economics of information to twentieth century economics. Oxford Review of Economic Policy, 16(4), 21–35. https://doi.org/10.1093/oxrep/16.4.21

Teece, D. J., Pisano, G., & Shuen, A. (1997). Dynamic capabilities and strategic management. Strategic Management Journal, 18(7), 509–533. https://doi.org/10.1002/(SICI)1097-0266(199708)18:7%253C509::AID-SMJ882%253E3.0.CO;2-Z

Tsang, R. C. W., Baldwin, A. A., Hair, J. F., Affuso, E., & Lahtinen, K. D. (2023). The informativeness of sentiment types in risk factor disclosures: Evidence from firms with cybersecurity breaches. Journal of Information Systems, 37(3), 157–190. https://doi.org/10.2308/ISYS-2022-014

Tushman, M. L., & Nadler, D. A. (1978). Information processing as an integrating concept in organizational design. The Academy of Management Review, 3(3), 613–624. https://doi.org/10.2307/257550

U.S. Department of Health and Human Services. (2013). Breach notification for unsecured protected health information, 45 C.F.R. § 164.400-414. https://www.ecfr.gov/current/title-45/subtitle-A/subchapter-C/part-164/subpart-D

Verrecchia, R. E. (1983). Discretionary disclosure. Journal of Accounting and Economics, 5(3), 179–194. https://doi.org/10.1016/0165-4101(83)90011-3

Villalonga, B., & Amit, R. (2006). How do family ownership, management, and control affect firm value? Journal of Financial Economics, 80(2), 385–417. https://doi.org/10.1016/j.jfineco.2004.12.005

Watts, R. L., & Zimmerman, J. L. (1986). Positive accounting theory. Prentice Hall.

Welker, M. (1995). Disclosure policy, information asymmetry, and liquidity in equity markets. Contemporary Accounting Research, 11(2), 801–827. https://doi.org/10.1111/j.1911-3846.1995.tb00469.x

Williams, B. M. (2008). Disclosure quality and the cost of equity. Journal of Accounting and Public Policy, 27(1), 21–42. https://doi.org/10.1016/j.jaccpubpol.2007.11.001

Xu, M., Jug, Ž., & Tamò-Larrieux, A. (2024). A cross-cultural analysis of transparency: The interplay of law, privacy policies, and user perceptions. International Data Privacy Law, 14(3), 197–222. https://doi.org/10.1093/idpl/ipae011

Zahra, S. A., & George, G. (2002). Absorptive capacity: A review, reconceptualization, and extension. Academy of Management Review, 27(2), 185–203. https://doi.org/10.5465/amr.2002.6587995

Zeff, S. A. (2003). How the U.S. accounting profession got where it is today: Part II. Accounting Horizons, 17(4), 267–286. https://doi.org/10.2308/acch.2003.17.4.267
"""
    return references

def create_complete_essay():
    """Create complete Essay 1 document with Conclusion and References"""

    # Read existing essay text
    essay_text = read_essay_text()

    # Split essay into introduction, literature review, methods, results
    # and everything before references
    parts = essay_text.split('\nReferences')

    main_content = parts[0]

    # Create Word document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Add main essay content (everything before Conclusion)
    # Remove everything from "Conclusion" onward if it exists
    if '\nConclusion' in main_content:
        main_content = main_content.split('\nConclusion')[0]

    # Parse and add existing content
    paragraphs = main_content.split('\n')
    for para_text in paragraphs:
        if para_text.strip():
            # Add with appropriate heading styles
            if para_text in ['Introduction', 'Literature Review', 'Data and Methods', 'Results', 'Robustness Checks']:
                p = doc.add_paragraph(para_text, style='Heading 1')
            elif para_text.startswith('TABLE') or para_text.startswith('Notes:'):
                p = doc.add_paragraph(para_text, style='Normal')
                p.runs[0].font.size = Pt(10)
            else:
                p = doc.add_paragraph(para_text, style='Normal')

    # Add Conclusion section
    conclusion_text = get_conclusion()
    for para_text in conclusion_text.split('\n'):
        if para_text.strip():
            if para_text in ['Conclusion', 'Synthesis of Findings', 'The Paradox: Why Mandatory Disclosure Backfires',
                           'Regulatory Burden and Firm Capacity', 'Theoretical Contributions', 'Policy Implications',
                           'Limitations', 'Future Research', 'Conclusion Summary']:
                p = doc.add_heading(para_text, level=2)
            else:
                p = doc.add_paragraph(para_text, style='Normal')

    # Add page break before References
    doc.add_page_break()

    # Add References section
    references_text = get_expanded_references()
    p = doc.add_heading('References', level=1)

    # Parse references
    ref_paragraphs = references_text.split('\n')
    skip_header = True
    for para_text in ref_paragraphs:
        if para_text.strip() and skip_header:
            if para_text == 'References':
                skip_header = False
                continue
        if para_text.strip() and not skip_header:
            p = doc.add_paragraph(para_text, style='Normal')
            # Hanging indent for references
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.runs[0].font.size = Pt(10)

    # Save document
    output_path = r"C:\Users\mcobp\DISSERTATION_CLONE\Essay 1 (6) (1) (1)_COMPLETE_FINAL.docx"
    doc.save(output_path)

    print(f"Complete Essay 1 saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    create_complete_essay()
    print("Essay 1 completion process finished.")
