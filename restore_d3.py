#!/usr/bin/env python3
"""
Restore Table D3 (Diagnostic Tests) from original to updated document.
"""

from docx import Document
from copy import deepcopy

# Load original (has all tables) and updated (missing D3)
original = Document('ESSAY2_FINAL_APPENDIX.docx')
updated = Document('ESSAY2_FINAL_APPENDIX_UPDATED.docx')

print("=" * 60)
print("RESTORING TABLE D3 (DIAGNOSTIC TESTS)")
print("=" * 60)

# Find D3 in original - it should be a table with "Test" as first cell
d3_table = None
for table_idx, table in enumerate(original.tables):
    if len(table.rows) > 0:
        first_cell = table.rows[0].cells[0].text if table.rows[0].cells else ""
        # D3 has diagnostic tests (Shapiro-Wilk, Breusch-Pagan, etc.)
        if 'Test' in first_cell and ('Shapiro' in str(table.rows) or 'Breusch' in str(table)):
            print(f"\nFound D3 at table {table_idx} in original")
            d3_table = table
            break

if not d3_table:
    # Alternative: look for table with specific diagnostic content
    for table_idx, table in enumerate(original.tables):
        table_text = ' '.join([cell.text for row in table.rows for cell in row.cells])
        if 'Shapiro-Wilk' in table_text and 'Breusch-Pagan' in table_text:
            print(f"\nFound D3 by content match at table {table_idx} in original")
            d3_table = table
            break

if d3_table:
    # Deep copy the table
    d3_copy = deepcopy(d3_table._element)

    # Find where to insert in updated document
    # Insert after the last current table (or before the end of document)
    updated_element = updated.element.body
    updated_element.append(d3_copy)

    # Save
    updated.save('ESSAY2_FINAL_APPENDIX_UPDATED.docx')
    print("D3 table restored to updated document")
    print("[OK] Document saved")
else:
    print("ERROR: Could not find D3 in original document")
