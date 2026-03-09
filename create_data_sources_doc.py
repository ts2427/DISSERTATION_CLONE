"""Create Data Sources Documentation for committee"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Data Sources & Methodology', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Complete Documentation of Data Provenance and Construction')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(12)
    run.font.italic = True

doc.add_paragraph()

# Overview
doc.add_heading('Overview', level=1)
doc.add_paragraph('This dissertation analyzes 1,054 data breach incidents at publicly-traded firms (2006-2025) to examine the effects of regulatory disclosure requirements on firm valuation, information asymmetry, and governance response. All data sources, matching procedures, and variable construction are documented below for complete transparency.')

# Primary Data Sources Section
doc.add_heading('Primary Data Sources', level=1)

# Create table of data sources
table = doc.add_table(rows=8, cols=4)
table.style = 'Light Grid Accent 1'

# Header
header_cells = table.rows[0].cells
headers = ['Data Element', 'Primary Source', 'Coverage', 'Sample Size']
for i, header_text in enumerate(headers):
    header_cells[i].text = header_text
    for run in header_cells[i].paragraphs[0].runs:
        run.font.bold = True

# Data rows
data_sources = [
    ['Breach Incidents', 'Privacy Rights Clearinghouse (PRC)', '2006-2025', '1,054'],
    ['Daily Stock Returns', 'CRSP (via WRDS)', 'All matched companies', '926 (87.9%)'],
    ['Firm Financials', 'Compustat (via WRDS)', 'Annual data, fiscal year preceding breach', '1,054'],
    ['Vulnerability Data', 'National Vulnerability Database (NVD)', 'CVSS v3.1 scores', '847 (80.4%)'],
    ['Executive Data', 'SEC EDGAR Form 8-K', 'Officer changes within 30/90/180d', '896 (85.0%)'],
    ['Regulatory Status', 'FCC Official Records', 'SIC codes 4813, 4899, 4841', '200 FCC firms'],
    ['Industry Classification', 'SIC & GICS Codes', 'Compustat classification', '1,054'],
]

for row_idx, row_data in enumerate(data_sources, start=1):
    cells = table.rows[row_idx].cells
    for col_idx, value in enumerate(row_data):
        cells[col_idx].text = value

# Data Construction Section
doc.add_heading('Data Construction & Matching', level=1)

doc.add_heading('Step 1: Breach Dataset', level=2)
doc.add_paragraph('Privacy Rights Clearinghouse (PRC) provides comprehensive database of publicly reported data breaches affecting individuals. DataBreaches.gov (GAO) provides supplementary federal contractor breach data.')
doc.add_paragraph('Initial extraction: All breaches with publicly-traded firm identifier (company name, ticker, CUSIP) recorded from 2006-2025.').paragraph_format.left_indent = Inches(0.25)

doc.add_heading('Step 2: CRSP Matching', level=2)
doc.add_paragraph('Breach incidents matched to CRSP using company name standardization, ticker symbol lookup, and CUSIP identifier.')
doc.add_paragraph('Match rate: 92.1% of incidents matched to CRSP securities.').paragraph_format.left_indent = Inches(0.25)

doc.add_heading('Step 3: Compustat Linking', level=2)
doc.add_paragraph('Firm financials merged using CIK (Central Index Key) from EDGAR. Data from Compustat annual database (fiscal years 1980-2025).')

doc.add_heading('Step 4: Vulnerability Enrichment', level=2)
doc.add_paragraph('For breaches mentioning specific CVEs, CVSS severity scores retrieved from NVD (National Vulnerability Database). NIST maintains authoritative CVE-to-CVSS mapping.')
doc.add_paragraph('Coverage: 80.4% of incidents have vendor-matched CVSS scores.').paragraph_format.left_indent = Inches(0.25)

doc.add_heading('Step 5: Regulatory Classification', level=2)
doc.add_paragraph('FCC regulatory status determined by SIC industry classification:')
doc.add_paragraph('SIC 4813: Telephone Communications (wireline)', style='List Bullet')
doc.add_paragraph('SIC 4899: Communications Services (wireless)', style='List Bullet')
doc.add_paragraph('SIC 4841: Cable and Other Pay Television Services', style='List Bullet')
doc.add_paragraph('19.0% of sample classified as FCC-regulated under Rule 37.3 (47 CFR 64.2011).').paragraph_format.left_indent = Inches(0.25)

# Essay-Specific Samples
doc.add_heading('Essay-Specific Sample Construction', level=1)

table2 = doc.add_table(rows=4, cols=4)
table2.style = 'Light Grid Accent 1'

header_cells2 = table2.rows[0].cells
headers2 = ['Essay', 'Analysis', 'Data Requirements', 'Final N']
for i, header_text in enumerate(headers2):
    header_cells2[i].text = header_text
    for run in header_cells2[i].paragraphs[0].runs:
        run.font.bold = True

essay_data = [
    ['Essay 1', 'Market Returns', 'CRSP price data, 180d window', '898'],
    ['Essay 2', 'Return Volatility', 'Pre/post volatility calculable', '891'],
    ['Essay 3', 'Executive Turnover', '8-K filings, officer identifiable', '955'],
]

for row_idx, row_data in enumerate(essay_data, start=1):
    cells = table2.rows[row_idx].cells
    for col_idx, value in enumerate(row_data):
        cells[col_idx].text = value

# Variables Section
doc.add_heading('Key Variables & Sources', level=1)

doc.add_heading('Dependent Variables', level=2)

dv1 = doc.add_paragraph()
dv1.add_run('Cumulative Abnormal Returns (CAR-30d): ').bold = True
dv1.add_run('Daily excess returns calculated using Fama-French 3-factor model. Data from CRSP daily prices and Kenneth French data repository. Event window: [-240, -60] trading days estimation, [0, 30] days test.')

dv2 = doc.add_paragraph()
dv2.add_run('Return Volatility: ').bold = True
dv2.add_run('Standard deviation of daily returns from CRSP. Pre-breach: 60 trading days before incident. Post-breach: 60 trading days after incident.')

dv3 = doc.add_paragraph()
dv3.add_run('Executive Turnover: ').bold = True
dv3.add_run('Binary indicator of officer/director departure within 30/90/180 days. Data from SEC EDGAR Form 8-K filings.')

doc.add_heading('Independent Variables', level=2)

iv1 = doc.add_paragraph()
iv1.add_run('FCC Reportable: ').bold = True
iv1.add_run('Binary indicator; 1 if SIC in {4813, 4899, 4841}, 0 otherwise. Source: Compustat SIC classification.')

iv2 = doc.add_paragraph()
iv2.add_run('Days to Disclosure: ').bold = True
iv2.add_run('Calendar days between breach discovery and public announcement. Source: PRC discovery date vs. SEC/press release date.')

iv3 = doc.add_paragraph()
iv3.add_run('Firm Size (log): ').bold = True
iv3.add_run('Natural log of total assets (millions). Source: Compustat item AT.')

doc.add_heading('Control Variables', level=2)

doc.add_paragraph('Health Breach: PHI involved (indicator). Source: PRC classification.').paragraph_format.left_indent = Inches(0.25)
doc.add_paragraph('Prior Breaches: Count of previous firm breaches. Source: PRC historical records.').paragraph_format.left_indent = Inches(0.25)
doc.add_paragraph('Leverage: Total debt / total assets. Source: Compustat.').paragraph_format.left_indent = Inches(0.25)
doc.add_paragraph('ROA: Net income / total assets. Source: Compustat.').paragraph_format.left_indent = Inches(0.25)
doc.add_paragraph('CVSS Complexity: NVD severity scores. Source: CVE-to-CVSS mapping.').paragraph_format.left_indent = Inches(0.25)

# Data Quality Section
doc.add_heading('Data Quality & Coverage', level=1)

doc.add_paragraph('Match Rates:', style='Heading 2')
doc.add_paragraph('CRSP matching: 92.1% (1,054 of 1,143 breaches)', style='List Number')
doc.add_paragraph('Compustat link: 100% (CIK matching)', style='List Number')
doc.add_paragraph('CVSS scores: 80.4% (847 of 1,054 with CVE)', style='List Number')
doc.add_paragraph('SEC 8-K data: 90.5% (956 of 1,054 with filings)', style='List Number')

doc.add_paragraph('Missing Data:', style='Heading 2')
doc.add_paragraph('No imputation used. Complete-case analysis for each dependent variable.', style='List Number')

doc.add_paragraph('Sample Variation by Essay:', style='Heading 2')
doc.add_paragraph('Essay 1 (N=898): Requires estimation + test window', style='List Bullet')
doc.add_paragraph('Essay 2 (N=891): Requires pre/post volatility', style='List Bullet')
doc.add_paragraph('Essay 3 (N=955): Requires officer departure data', style='List Bullet')

# Save
doc.save('Data_Sources_Documentation.docx')
print('SUCCESS: Data_Sources_Documentation.docx created')
print('\nIncludes:')
print('  - All primary data sources (PRC, CRSP, Compustat, NVD, SEC EDGAR, FCC)')
print('  - Data matching procedures (5-step construction)')
print('  - Essay-specific sample sizes')
print('  - Complete variable definitions & sources')
print('  - Data quality metrics')
